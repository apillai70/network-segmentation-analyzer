#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Complete Report Generator - Fortinet-Style Professional Reports
================================================================
Generates comprehensive network segmentation and threat surface reports
using enriched flows.csv data (18 columns with DNS validation, VMware detection,
and cross-app dependencies).

OUTPUT PACKAGE:
===============
For each application:
  1. Interactive HTML Diagram (with controls, zoom, pan)
  2. Mermaid (.mmd) source
  3. PNG Export (via mmdc CLI)
  4. SVG Export (vector graphics)
  5. Architecture DOCX (detailed documentation)
  6. Threat Surface DOCX (security analysis)
  7. Network Segmentation JSON (machine-readable)
  8. Dependency Graph HTML (inter-app connections)

Master Reports:
  9. Unified Network Analysis HTML (all apps combined)
  10. Threat Surface Analysis HTML (enterprise-wide)
  11. Segmentation Strategy DOCX (executive summary)
  12. Gap Analysis Report DOCX
  13. Enterprise Dashboard HTML (interactive overview)

Author: Enterprise Security Team
Version: 1.0 - Complete Report Suite
"""

import sys
import logging
import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import Dict, List
import subprocess
import json

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

logger = logging.getLogger(__name__)


class CompleteReportGenerator:
    """
    Generates complete report package from enriched flows.csv
    """

    def __init__(self, apps_dir: str = 'persistent_data/applications',
                 output_dir: str = 'outputs'):
        self.apps_dir = Path(apps_dir)
        self.output_dir = Path(output_dir)

        # Create output subdirectories
        self.diagrams_dir = self.output_dir / 'diagrams'
        self.html_dir = self.output_dir / 'html'
        self.docx_dir = self.output_dir / 'docx'
        self.png_dir = self.output_dir / 'png'
        self.svg_dir = self.output_dir / 'svg'
        self.mmd_dir = self.output_dir / 'mermaid'
        self.json_dir = self.output_dir / 'json'

        for dir_path in [self.diagrams_dir, self.html_dir, self.docx_dir,
                         self.png_dir, self.svg_dir, self.mmd_dir, self.json_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

        logger.info("CompleteReportGenerator initialized")
        logger.info(f"  Apps directory: {self.apps_dir}")
        logger.info(f"  Output directory: {self.output_dir}")

    def generate_for_app(self, app_id: str, flows_df: pd.DataFrame) -> Dict:
        """
        Generate complete report package for one application

        Args:
            app_id: Application ID
            flows_df: Enriched DataFrame (18 columns)

        Returns:
            Dict with paths to generated files
        """
        logger.info(f"[{app_id}] Generating complete report package...")

        outputs = {}

        # 1. Generate DOWNSTREAM diagram (what this app depends on)
        downstream_mermaid = self._generate_mermaid_diagram(app_id, flows_df)
        mmd_file = self.mmd_dir / f"{app_id}_downstream.mmd"
        with open(mmd_file, 'w', encoding='utf-8') as f:
            f.write(downstream_mermaid)
        outputs['downstream_mermaid'] = str(mmd_file)
        logger.info(f"  [OK] Downstream Mermaid: {mmd_file.name}")

        # 2. Generate UPSTREAM diagram (who depends on this app)
        upstream_mermaid = self._generate_upstream_diagram(app_id)
        upstream_mmd_file = self.mmd_dir / f"{app_id}_upstream.mmd"
        with open(upstream_mmd_file, 'w', encoding='utf-8') as f:
            f.write(upstream_mermaid)
        outputs['upstream_mermaid'] = str(upstream_mmd_file)
        logger.info(f"  [OK] Upstream Mermaid: {upstream_mmd_file.name}")

        # 3. Generate FULL FLOW diagram (Source Apps → Our App → Dest Apps)
        fullflow_mermaid = self._generate_fullflow_diagram(app_id, flows_df)
        fullflow_mmd_file = self.mmd_dir / f"{app_id}_fullflow.mmd"
        with open(fullflow_mmd_file, 'w', encoding='utf-8') as f:
            f.write(fullflow_mermaid)
        outputs['fullflow_mermaid'] = str(fullflow_mmd_file)
        logger.info(f"  [OK] Full Flow Mermaid: {fullflow_mmd_file.name}")

        # 3.5 Generate ARCHITECTURE diagram (Professional tier-based layout for Word docs)
        architecture_mermaid = self._generate_architecture_diagram(app_id, flows_df)
        architecture_mmd_file = self.mmd_dir / f"{app_id}_architecture.mmd"
        with open(architecture_mmd_file, 'w', encoding='utf-8') as f:
            f.write(architecture_mermaid)
        outputs['architecture_mermaid'] = str(architecture_mmd_file)
        logger.info(f"  [OK] Architecture Mermaid: {architecture_mmd_file.name}")

        # 4. Generate Interactive HTML with ALL FOUR diagrams
        html_content = self._generate_interactive_html(app_id, downstream_mermaid, upstream_mermaid, fullflow_mermaid, architecture_mermaid, flows_df)
        html_file = self.html_dir / f"{app_id}_application_diagram.html"
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        outputs['html'] = str(html_file)
        logger.info(f"  [OK] Interactive HTML: {html_file.name}")

        # 3. Generate PNG (if mmdc available)
        png_file = self.png_dir / f"{app_id}_architecture.png"
        if self._generate_png_from_mermaid(mmd_file, png_file):
            outputs['png'] = str(png_file)
            logger.info(f"  [OK] PNG diagram: {png_file.name}")
        else:
            logger.warning(f"  [SKIP] PNG generation (mmdc not available)")

        # 4. Generate SVG
        svg_file = self.svg_dir / f"{app_id}_architecture.svg"
        if self._generate_svg_from_mermaid(mmd_file, svg_file):
            outputs['svg'] = str(svg_file)
            logger.info(f"  [OK] SVG diagram: {svg_file.name}")

        # 5. Generate Architecture DOCX
        arch_docx = self.docx_dir / f"{app_id}_architecture.docx"
        self._generate_architecture_docx(app_id, flows_df, arch_docx)
        outputs['architecture_docx'] = str(arch_docx)
        logger.info(f"  [OK] Architecture DOCX: {arch_docx.name}")

        # 6. Generate Threat Surface DOCX
        threat_docx = self.docx_dir / f"{app_id}_threat_surface.docx"
        self._generate_threat_docx(app_id, flows_df, threat_docx)
        outputs['threat_docx'] = str(threat_docx)
        logger.info(f"  [OK] Threat Surface DOCX: {threat_docx.name}")

        # 7. Generate Network Segmentation JSON
        seg_json = self.json_dir / f"{app_id}_segmentation.json"
        self._generate_segmentation_json(app_id, flows_df, seg_json)
        outputs['segmentation_json'] = str(seg_json)
        logger.info(f"  [OK] Segmentation JSON: {seg_json.name}")

        logger.info(f"[{app_id}] Complete! Generated {len(outputs)} outputs")
        return outputs

    def _generate_mermaid_diagram(self, app_id: str, flows_df: pd.DataFrame) -> str:
        """Generate DOWNSTREAM diagram - CLASSIC AGGREGATED VIEW (what this app depends on)"""
        return self._generate_downstream_only_diagram(app_id, flows_df)

    def _detect_load_balancer(self, hostname: str) -> bool:
        """Detect if hostname is a load balancer"""
        lb_patterns = ['f5', 'lb', 'loadbalancer', 'load-balancer', 'alb', 'elb', 'nlb',
                       'bigip', 'netscaler', 'citrix', 'haproxy', 'nginx-lb', 'a10']
        hostname_lower = hostname.lower()
        return any(pattern in hostname_lower for pattern in lb_patterns)

    def _detect_server_type(self, hostname: str, ports: list = None) -> str:
        """
        Detect server type based on hostname and ports
        Returns: Web, Database, App, DNS, Mail, File, Domain Controller, Cache, Message Queue, or General
        """
        hostname_lower = hostname.lower() if hostname else ""

        # Hostname-based detection (most reliable)
        hostname_patterns = {
            'Web': ['web', 'www', 'nginx', 'apache', 'iis', 'httpd', 'frontend', 'ui'],
            'Database': ['db', 'sql', 'mysql', 'postgres', 'oracle', 'mongo', 'redis', 'cassandra', 'mariadb', 'mssql'],
            'App': ['app', 'application', 'api', 'backend', 'service', 'tomcat', 'jboss', 'wildfly'],
            'DNS': ['dns', 'nameserver', 'ns1', 'ns2', 'bind'],
            'Mail': ['mail', 'smtp', 'exchange', 'postfix', 'sendmail', 'mx'],
            'File': ['file', 'storage', 'nas', 'san', 'ftp', 'sftp', 'share'],
            'Domain Controller': ['dc', 'ad', 'domaincontroller', 'activedirectory', 'ldap'],
            'Cache': ['cache', 'memcached', 'varnish'],
            'Message Queue': ['mq', 'kafka', 'rabbitmq', 'activemq', 'queue'],
            'Monitoring': ['monitor', 'nagios', 'zabbix', 'prometheus', 'grafana'],
            'Backup': ['backup', 'bkp', 'veeam', 'commvault']
        }

        for server_type, patterns in hostname_patterns.items():
            if any(pattern in hostname_lower for pattern in patterns):
                return server_type

        # Port-based detection (fallback)
        if ports:
            port_mappings = {
                'Web': [80, 443, 8080, 8443, 8000, 8888, 9000, 3000],
                'Database': [1433, 3306, 5432, 1521, 27017, 6379, 9042, 5984],
                'Mail': [25, 587, 465, 110, 995, 143, 993],
                'DNS': [53],
                'File': [21, 22, 445, 139, 2049],
                'Domain Controller': [389, 636, 88, 464],
                'Cache': [11211, 6379],
                'Message Queue': [5672, 9092, 61616]
            }

            for server_type, type_ports in port_mappings.items():
                if any(port in ports for port in type_ports):
                    return server_type

        return 'General'

    def _get_edge_style(self, flow_count: int) -> str:
        """
        Return edge style based on flow count:
        - < 10 flows: dashed line (light traffic)
        - 10-100 flows: solid normal line (medium traffic)
        - > 100 flows: thick solid line (heavy traffic)
        """
        if flow_count < 10:
            return ".->"  # Dotted/dashed line
        elif flow_count <= 100:
            return "-->"  # Normal solid line
        else:
            return "==>"  # Thick solid line

    def _generate_architecture_diagram(self, app_id: str, flows_df: pd.DataFrame) -> str:
        """
        Generate PROFESSIONAL ARCHITECTURE DIAGRAM with server type grouping
        Layout: Upstream (Top) → [Web | App | DB | LB] (Center) → Downstream (Bottom)
        Uses VERTICAL layout (TD) to fit better on pages, limits shown items for scalability
        """
        mermaid = "graph TD\n"
        mermaid += f"    %% {app_id} Architecture Diagram (Vertical Layout)\n\n"

        # Classify source servers by type
        server_types = {}
        for _, row in flows_df[['Source IP', 'Source Hostname', 'Port']].drop_duplicates().iterrows():
            ip = row['Source IP']
            hostname = row.get('Source Hostname', ip)
            ports = flows_df[flows_df['Source IP'] == ip]['Port'].unique().tolist()
            server_type = self._detect_server_type(hostname, ports)

            if server_type not in server_types:
                server_types[server_type] = []
            server_types[server_type].append({'ip': ip, 'hostname': hostname})

        # Get upstream applications
        all_flows = []
        for app_dir in self.apps_dir.iterdir():
            if not app_dir.is_dir():
                continue
            flows_csv = app_dir / 'flows.csv'
            if flows_csv.exists():
                try:
                    df = pd.read_csv(flows_csv)
                    all_flows.append(df)
                except:
                    continue

        upstream_apps = []
        if all_flows:
            all_flows_df = pd.concat(all_flows, ignore_index=True)
            upstream_flows = all_flows_df[all_flows_df['Dest App'] == app_id]
            if len(upstream_flows) > 0:
                upstream_apps = sorted(upstream_flows['App'].unique())

        # Get downstream applications
        outbound_data = flows_df[flows_df['Flow Direction'] == 'outbound']
        downstream_apps = []
        if len(outbound_data) > 0:
            downstream_apps = sorted(outbound_data['Dest App'].unique())

        # UPSTREAM (Top)
        if upstream_apps:
            mermaid += "    subgraph UPSTREAM[\"⬆️ UPSTREAM SOURCES\"]\n"
            for idx, src_app in enumerate(upstream_apps[:5]):  # Limit to 5 for page fit
                app_flows = upstream_flows[upstream_flows['App'] == src_app]
                flow_count = len(app_flows)
                mermaid += f"        UP{idx}[\"{src_app}<br/>{flow_count} flows\"]:::upstream\n"
            if len(upstream_apps) > 5:
                mermaid += f"        UP_MORE[\"... +{len(upstream_apps) - 5} more apps\"]:::upstream\n"
            mermaid += "    end\n\n"

        # CENTER APPLICATION (Main box with server type groups)
        mermaid += f"    subgraph CENTER[\"{app_id} APPLICATION\"]\n"

        # Infer load balancers based on multiple servers
        has_web_lb = 'Web' in server_types and len(server_types['Web']) > 1
        has_app_lb = 'App' in server_types and len(server_types['App']) > 1

        # Network Load Balancer (for Web tier if multiple web servers)
        if has_web_lb:
            mermaid += f"        subgraph NLB[\"⚖️ NETWORK LOAD BALANCER (Inferred)\"]\n"
            mermaid += f"            NLB_NODE[\"External/Internal Gateway<br/>Distributes traffic to Web Tier\"]:::loadbalancer\n"
            mermaid += f"        end\n"

        # Web servers
        if 'Web' in server_types:
            web_count = len(server_types['Web'])
            mermaid += f"        subgraph WEB[\"🌐 WEB TIER\"]\n"
            if web_count <= 2:
                for idx, server in enumerate(server_types['Web']):
                    ip = server['ip']
                    hostname = server['hostname']
                    if hostname != ip and hostname != 'Unknown':
                        label = f"{ip}<br/>({hostname})"
                    else:
                        label = ip
                    mermaid += f"            W{idx}[\"{label}\"]:::web\n"
            else:
                # Show first 2 with "... + X more"
                for idx, server in enumerate(server_types['Web'][:2]):
                    ip = server['ip']
                    hostname = server['hostname']
                    if hostname != ip and hostname != 'Unknown':
                        label = f"{ip}<br/>({hostname})"
                    else:
                        label = ip
                    mermaid += f"            W{idx}[\"{label}\"]:::web\n"
                more_count = web_count - 2
                mermaid += f"            WMORE[\"... (+{more_count} more)\"]:::webmore\n"
                mermaid += f"            click WMORE showWebServers\n"
            mermaid += "        end\n"

        # Application Load Balancer (for App tier if multiple app servers)
        if has_app_lb:
            mermaid += f"        subgraph ALB[\"⚖️ APPLICATION LOAD BALANCER (Inferred)\"]\n"
            mermaid += f"            ALB_NODE[\"Internal Gateway/API Gateway<br/>Distributes traffic to App Tier\"]:::loadbalancer\n"
            mermaid += f"        end\n"

        # App servers
        if 'App' in server_types:
            app_count = len(server_types['App'])
            mermaid += f"        subgraph APP[\"⚙️ APPLICATION TIER\"]\n"
            if app_count <= 2:
                for idx, server in enumerate(server_types['App']):
                    ip = server['ip']
                    hostname = server['hostname']
                    if hostname != ip and hostname != 'Unknown':
                        label = f"{ip}<br/>({hostname})"
                    else:
                        label = ip
                    mermaid += f"            A{idx}[\"{label}\"]:::app\n"
            else:
                # Show first 2 with "... + X more"
                for idx, server in enumerate(server_types['App'][:2]):
                    ip = server['ip']
                    hostname = server['hostname']
                    if hostname != ip and hostname != 'Unknown':
                        label = f"{ip}<br/>({hostname})"
                    else:
                        label = ip
                    mermaid += f"            A{idx}[\"{label}\"]:::app\n"
                more_count = app_count - 2
                mermaid += f"            AMORE[\"... (+{more_count} more)\"]:::appmore\n"
                mermaid += f"            click AMORE showAppServers\n"
            mermaid += "        end\n"

        # Database servers
        if 'Database' in server_types:
            db_count = len(server_types['Database'])
            mermaid += f"        subgraph DB[\"🗄️ DATABASE TIER\"]\n"
            if db_count <= 2:
                for idx, server in enumerate(server_types['Database']):
                    ip = server['ip']
                    hostname = server['hostname']
                    if hostname != ip and hostname != 'Unknown':
                        label = f"{ip}<br/>({hostname})"
                    else:
                        label = ip
                    mermaid += f"            D{idx}[\"{label}\"]:::database\n"
            else:
                # Show first 2 with "... + X more"
                for idx, server in enumerate(server_types['Database'][:2]):
                    ip = server['ip']
                    hostname = server['hostname']
                    if hostname != ip and hostname != 'Unknown':
                        label = f"{ip}<br/>({hostname})"
                    else:
                        label = ip
                    mermaid += f"            D{idx}[\"{label}\"]:::database\n"
                more_count = db_count - 2
                mermaid += f"            DMORE[\"... (+{more_count} more)\"]:::dbmore\n"
                mermaid += f"            click DMORE showDatabaseServers\n"
            mermaid += "        end\n"

        # Load Balancers (if detected)
        lb_servers = []
        for _, row in flows_df[['Source IP', 'Source Hostname']].drop_duplicates().iterrows():
            if self._detect_load_balancer(row.get('Source Hostname', '')):
                lb_servers.append({'ip': row['Source IP'], 'hostname': row.get('Source Hostname', row['Source IP'])})

        if lb_servers:
            mermaid += f"        subgraph LB[\"⚖️ LOAD BALANCERS\"]\n"
            for idx, server in enumerate(lb_servers[:2]):
                hostname = server['hostname'] if server['hostname'] != server['ip'] else server['ip']
                mermaid += f"            LB{idx}[\"{hostname}<br/>{server['ip']}\"]:::loadbalancer\n"
            mermaid += "        end\n"

        # Other server types
        other_types = [t for t in server_types.keys() if t not in ['Web', 'App', 'Database']]
        if other_types:
            mermaid += f"        subgraph OTHER[\"🖥️ OTHER SERVICES\"]\n"
            for server_type in other_types[:2]:  # Limit to 2 types for page fit
                count = len(server_types[server_type])
                mermaid += f"            OTH_{server_type.replace(' ', '_')}[\"{server_type}<br/>{count} servers\"]:::other\n"
            mermaid += "        end\n"

        mermaid += "    end\n\n"

        # DOWNSTREAM (Bottom)
        if downstream_apps:
            mermaid += "    subgraph DOWNSTREAM[\"⬇️ DOWNSTREAM DESTINATIONS\"]\n"
            for idx, dest_app in enumerate(downstream_apps[:5]):  # Limit to 5 for page fit
                dest_data = outbound_data[outbound_data['Dest App'] == dest_app]
                flow_count = len(dest_data)
                server_count = dest_data['Dest IP'].nunique()
                mermaid += f"        DOWN{idx}[\"{dest_app}<br/>{server_count} servers | {flow_count} flows\"]:::downstream\n"
            if len(downstream_apps) > 5:
                mermaid += f"        DOWN_MORE[\"... +{len(downstream_apps) - 5} more apps\"]:::downstream\n"
            mermaid += "    end\n\n"

        # External traffic
        external_flows = len(flows_df[flows_df['Flow Direction'] == 'external'])
        if external_flows > 0:
            external_ips = flows_df[flows_df['Flow Direction'] == 'external']['Dest IP'].nunique()
            mermaid += f"    EXTERNAL[\"🌍 EXTERNAL / INTERNET<br/>{external_ips} destinations | {external_flows} flows\"]:::external\n\n"

        # Connections
        if upstream_apps:
            for idx in range(min(len(upstream_apps), 5)):
                mermaid += f"    UP{idx} --> CENTER\n"

        if downstream_apps:
            for idx in range(min(len(downstream_apps), 5)):
                mermaid += f"    CENTER --> DOWN{idx}\n"

        if external_flows > 0:
            mermaid += f"    CENTER --> EXTERNAL\n"

        # Styles
        mermaid += "\n    %% Styles (thinner strokes for professional appearance)\n"
        mermaid += "    classDef upstream fill:#E3F2FD,stroke:#1976D2,stroke-width:1.5px,color:#000\n"
        mermaid += "    classDef web fill:#4CAF50,stroke:#2E7D32,stroke-width:2px,color:#fff\n"
        mermaid += "    classDef webmore fill:#81C784,stroke:#2E7D32,stroke-width:1.5px,stroke-dasharray:5,color:#fff\n"
        mermaid += "    classDef app fill:#FF9800,stroke:#E65100,stroke-width:2px,color:#fff\n"
        mermaid += "    classDef appmore fill:#FFB74D,stroke:#E65100,stroke-width:1.5px,stroke-dasharray:5,color:#fff\n"
        mermaid += "    classDef database fill:#9C27B0,stroke:#6A1B9A,stroke-width:2px,color:#fff\n"
        mermaid += "    classDef dbmore fill:#BA68C8,stroke:#6A1B9A,stroke-width:1.5px,stroke-dasharray:5,color:#fff\n"
        mermaid += "    classDef loadbalancer fill:#FFF,stroke:#F44336,stroke-width:2px,color:#000\n"
        mermaid += "    classDef other fill:#607D8B,stroke:#37474F,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef downstream fill:#2196F3,stroke:#1565C0,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef external fill:#F44336,stroke:#C62828,stroke-width:1.5px,color:#fff\n"

        return mermaid

    def _generate_downstream_only_diagram(self, app_id: str, flows_df: pd.DataFrame) -> str:
        """Generate DOWNSTREAM ONLY: Our App → Destination Apps"""
        # Build CLASSIC aggregated diagram
        mermaid = "graph LR\n"
        mermaid += f"    %% DOWNSTREAM: Where {app_id} sends data\n\n"

        # Filter out infrastructure protocols (DNS, LDAP, Kerberos, NTP, DHCP, etc.)
        infrastructure_ports = [53, 389, 636, 88, 123, 67, 68, 135, 137, 138, 139, 445, 464]
        flows_df = flows_df[~flows_df['Port'].isin(infrastructure_ports)] if 'Port' in flows_df.columns else flows_df

        # THIS application as source (aggregated)
        total_flows = len(flows_df)
        internal_flows = len(flows_df[flows_df['Flow Direction'] == 'internal'])
        outbound_flows = len(flows_df[flows_df['Flow Direction'] == 'outbound'])
        external_flows = len(flows_df[flows_df['Flow Direction'] == 'external'])
        source_servers = flows_df['Source IP'].nunique()

        # Check for load balancers in source
        source_lbs = []
        source_regular_servers = []
        if 'Source Hostname' in flows_df.columns:
            for _, row in flows_df[['Source IP', 'Source Hostname']].drop_duplicates().iterrows():
                hostname = row['Source Hostname']
                ip = row['Source IP']
                if self._detect_load_balancer(hostname):
                    source_lbs.append((hostname, ip))
                else:
                    source_regular_servers.append((hostname, ip))

        # If we have load balancers, show LB → Servers pattern
        if source_lbs:
            # Show Load Balancer with IP (hostname) format
            lb_entries = []
            for name, ip in source_lbs[:2]:
                if name and name != ip:
                    lb_entries.append(f"{ip} ({name})")
                else:
                    lb_entries.append(ip)
            lb_str = '<br/>'.join(lb_entries)
            lb_label = f"<b>Load Balancer</b><br/>{lb_str}"
            mermaid += f"    LB[\"{lb_label}\"]:::loadbalancer\n"

            # Show servers behind LB with IP (hostname) format
            server_entries = []
            for hostname, ip in source_regular_servers[:2]:
                if hostname and hostname != ip:
                    server_entries.append(f"{ip} ({hostname})")
                else:
                    server_entries.append(ip)
            server_str = '<br/>'.join(server_entries)
            if len(source_regular_servers) > 2:
                server_str += f"<br/>... (+{len(source_regular_servers) - 2} more)"

            source_label = f"<b>{app_id} Servers</b><br/>{server_str}<br/>{len(source_regular_servers)} Servers"
            mermaid += f"    SOURCE[\"{source_label}\"]:::sourceapp\n"
            mermaid += f"    LB -->|distributes| SOURCE\n\n"
        else:
            # No LB, show regular source with IP (hostname) format
            source_ip_hostname_list = []
            for _, row in flows_df[['Source IP', 'Source Hostname']].drop_duplicates().iterrows():
                ip = row['Source IP']
                hostname = row.get('Source Hostname', ip)
                if hostname and hostname != ip and hostname != 'Unknown':
                    source_ip_hostname_list.append(f"{ip} ({hostname})")
                else:
                    source_ip_hostname_list.append(ip)

            # Show first 3 entries
            source_str = '<br/>'.join(source_ip_hostname_list[:3])
            if len(source_ip_hostname_list) > 3:
                source_str += f"<br/>... (+{len(source_ip_hostname_list) - 3} more)"

            source_label = f"<b>{app_id}</b><br/>{source_str}<br/>{source_servers} Servers | {total_flows} Flows"
            mermaid += f"    SOURCE[\"{source_label}\"]:::sourceapp\n\n"

        # Internal flows (if significant)
        if internal_flows > 0:
            internal_label = f"<b>Internal Traffic</b><br/>{internal_flows} Flows"
            mermaid += f"    INTERNAL[\"{internal_label}\"]:::internal\n"
            edge_style = self._get_edge_style(internal_flows)
            mermaid += f"    SOURCE {edge_style}|{internal_flows}| INTERNAL\n\n"

        # Outbound to other applications (AGGREGATED)
        outbound_data = flows_df[flows_df['Flow Direction'] == 'outbound']
        if len(outbound_data) > 0:
            dest_apps = outbound_data.groupby('Dest App').agg(
                Servers=('Dest IP', 'nunique'),
                Flows=('Source IP', 'count')
            ).reset_index()

            for idx, row in dest_apps.iterrows():
                dest_app = row['Dest App']
                server_count = row['Servers']
                flow_count = row['Flows']

                # Get IPs and hostnames for this destination app
                dest_app_data = outbound_data[outbound_data['Dest App'] == dest_app]

                # Build IP (hostname) format
                ip_hostname_list = []
                for _, row in dest_app_data[['Dest IP', 'Dest Hostname']].drop_duplicates().iterrows():
                    ip = row['Dest IP']
                    hostname = row.get('Dest Hostname', ip)
                    if hostname and hostname != ip and hostname != 'Unknown':
                        ip_hostname_list.append(f"{ip} ({hostname})")
                    else:
                        ip_hostname_list.append(ip)

                # Show first 2 entries
                ip_hostname_str = '<br/>'.join(ip_hostname_list[:2])
                if len(ip_hostname_list) > 2:
                    ip_hostname_str += f"<br/>... (+{len(ip_hostname_list) - 2} more)"

                # Check if VMware
                app_vmware = dest_app_data['Dest Is VMware'].sum() if 'Dest Is VMware' in dest_app_data.columns else 0

                dest_node = f"DEST_APP_{idx}"
                vmware_text = f" | {app_vmware} VMware" if app_vmware > 0 else ""
                dest_label = f"<b>{dest_app}</b><br/>{ip_hostname_str}<br/>{server_count} Servers | {flow_count} Flows{vmware_text}"

                if app_vmware > 0:
                    mermaid += f"    {dest_node}[(\"{dest_label}\")]:::vmware\n"
                else:
                    mermaid += f"    {dest_node}[\"{dest_label}\"]:::destapp\n"

                edge_style = self._get_edge_style(flow_count)
                mermaid += f"    SOURCE {edge_style}|{flow_count}| {dest_node}\n"

        # External/Unknown destinations (AGGREGATED)
        if external_flows > 0:
            external_data = flows_df[flows_df['Flow Direction'] == 'external']

            # Build IP (hostname) format for external
            ext_ip_hostname_list = []
            for _, row in external_data[['Dest IP', 'Dest Hostname']].drop_duplicates().iterrows():
                ip = row['Dest IP']
                hostname = row.get('Dest Hostname', ip)
                if hostname and hostname != ip and hostname != 'Unknown':
                    ext_ip_hostname_list.append(f"{ip} ({hostname})")
                else:
                    ext_ip_hostname_list.append(ip)

            # Show first 2 entries
            ext_str = '<br/>'.join(ext_ip_hostname_list[:2])
            if len(ext_ip_hostname_list) > 2:
                ext_str += f"<br/>... (+{len(ext_ip_hostname_list) - 2} more)"

            external_label = f"<b>External / Internet</b><br/>{ext_str}<br/>{external_data['Dest IP'].nunique()} Destinations | {external_flows} Flows"
            mermaid += f"\n    EXTERNAL[\"{external_label}\"]:::external\n"
            edge_style = self._get_edge_style(external_flows)
            mermaid += f"    SOURCE {edge_style}|{external_flows}| EXTERNAL\n"

        # Add styles (thinner strokes for professional appearance)
        mermaid += "\n    %% Styles\n"
        mermaid += "    classDef sourceapp fill:#4CAF50,stroke:#2E7D32,stroke-width:2px,color:#fff\n"
        mermaid += "    classDef destapp fill:#2196F3,stroke:#1565C0,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef vmware fill:#FF9800,stroke:#E65100,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef internal fill:#9C27B0,stroke:#6A1B9A,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef external fill:#F44336,stroke:#C62828,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef loadbalancer fill:#FFF,stroke:#F44336,stroke-width:2px,color:#000,rx:2,ry:2\n"

        return mermaid

    def _generate_fullflow_diagram(self, app_id: str, flows_df: pd.DataFrame) -> str:
        """Generate FULL FLOW: Source Apps → Our App → Destination Apps (End-to-End)"""
        # Load ALL flows from ALL applications to get upstream
        all_flows = []
        for app_dir in self.apps_dir.iterdir():
            if not app_dir.is_dir():
                continue
            flows_csv = app_dir / 'flows.csv'
            if flows_csv.exists():
                try:
                    df = pd.read_csv(flows_csv)
                    all_flows.append(df)
                except:
                    continue

        mermaid = "graph LR\n"
        mermaid += f"    %% FULL FLOW: Source Apps → {app_id} → Destination Apps\n\n"

        # UPSTREAM: Who sends TO us
        if all_flows:
            all_flows_df = pd.concat(all_flows, ignore_index=True)
            upstream_flows = all_flows_df[all_flows_df['Dest App'] == app_id]

            if len(upstream_flows) > 0:
                source_apps = sorted(upstream_flows['App'].unique())
                for idx, source_app in enumerate(source_apps):
                    app_flows = upstream_flows[upstream_flows['App'] == source_app]
                    server_count = app_flows['Source IP'].nunique()
                    flow_count = len(app_flows)
                    vmware_count = app_flows[app_flows.get('Source Is VMware', False) == True]['Source IP'].nunique() if 'Source Is VMware' in app_flows.columns else 0

                    # Build IP (hostname) format
                    src_ip_hostname_list = []
                    for _, row in app_flows[['Source IP', 'Source Hostname']].drop_duplicates().iterrows():
                        ip = row['Source IP']
                        hostname = row.get('Source Hostname', ip)
                        if hostname and hostname != ip and hostname != 'Unknown':
                            src_ip_hostname_list.append(f"{ip} ({hostname})")
                        else:
                            src_ip_hostname_list.append(ip)

                    # Show first 2 entries
                    src_str = '<br/>'.join(src_ip_hostname_list[:2])
                    if len(src_ip_hostname_list) > 2:
                        src_str += f"<br/>... (+{len(src_ip_hostname_list) - 2} more)"

                    app_node = f"SRC_APP_{idx}"
                    vmware_text = f" | {vmware_count} VMware" if vmware_count > 0 else ""
                    label = f"<b>{source_app}</b><br/>{src_str}<br/>{server_count} Servers | {flow_count} Flows{vmware_text}"

                    if vmware_count > 0:
                        mermaid += f"    {app_node}[(\"{label}\")]:::upstream\n"
                    else:
                        mermaid += f"    {app_node}[\"{label}\"]:::upstream\n"

        # CENTER: THIS application with IP (hostname)
        source_servers = flows_df['Source IP'].nunique()
        total_flows = len(flows_df)

        # Build IP (hostname) format for center
        center_ip_hostname_list = []
        for _, row in flows_df[['Source IP', 'Source Hostname']].drop_duplicates().iterrows():
            ip = row['Source IP']
            hostname = row.get('Source Hostname', ip)
            if hostname and hostname != ip and hostname != 'Unknown':
                center_ip_hostname_list.append(f"{ip} ({hostname})")
            else:
                center_ip_hostname_list.append(ip)

        # Show first 2 entries
        center_str = '<br/>'.join(center_ip_hostname_list[:2])
        if len(center_ip_hostname_list) > 2:
            center_str += f"<br/>... (+{len(center_ip_hostname_list) - 2} more)"

        center_label = f"<b>{app_id} (CENTER)</b><br/>{center_str}<br/>{source_servers} Servers | {total_flows} Flows"
        mermaid += f"\n    CENTER[\"{center_label}\"]:::centerapp\n\n"

        # Connect upstream TO center with flow-based edge styling
        if all_flows and len(upstream_flows) > 0:
            for idx, source_app in enumerate(source_apps):
                app_flows = upstream_flows[upstream_flows['App'] == source_app]
                flow_count = len(app_flows)
                edge_style = self._get_edge_style(flow_count)
                app_node = f"SRC_APP_{idx}"
                mermaid += f"    {app_node} {edge_style}|{flow_count}| CENTER\n"

        # DOWNSTREAM: Where we send
        outbound_data = flows_df[flows_df['Flow Direction'] == 'outbound']
        if len(outbound_data) > 0:
            dest_apps = outbound_data.groupby('Dest App').agg(
                Servers=('Dest IP', 'nunique'),
                Flows=('Source IP', 'count')
            ).reset_index()

            for idx, row in dest_apps.iterrows():
                dest_app = row['Dest App']
                server_count = row['Servers']
                flow_count = row['Flows']
                app_vmware = outbound_data[outbound_data['Dest App'] == dest_app]['Dest Is VMware'].sum() if 'Dest Is VMware' in outbound_data.columns else 0

                # Build IP (hostname) format for destination
                dest_app_data = outbound_data[outbound_data['Dest App'] == dest_app]
                dest_ip_hostname_list = []
                for _, row_data in dest_app_data[['Dest IP', 'Dest Hostname']].drop_duplicates().iterrows():
                    ip = row_data['Dest IP']
                    hostname = row_data.get('Dest Hostname', ip)
                    if hostname and hostname != ip and hostname != 'Unknown':
                        dest_ip_hostname_list.append(f"{ip} ({hostname})")
                    else:
                        dest_ip_hostname_list.append(ip)

                # Show first 2 entries
                dest_str = '<br/>'.join(dest_ip_hostname_list[:2])
                if len(dest_ip_hostname_list) > 2:
                    dest_str += f"<br/>... (+{len(dest_ip_hostname_list) - 2} more)"

                dest_node = f"DEST_APP_{idx}"
                vmware_text = f" | {app_vmware} VMware" if app_vmware > 0 else ""
                dest_label = f"<b>{dest_app}</b><br/>{dest_str}<br/>{server_count} Servers | {flow_count} Flows{vmware_text}"

                if app_vmware > 0:
                    mermaid += f"    {dest_node}[(\"{dest_label}\")]:::downstream\n"
                else:
                    mermaid += f"    {dest_node}[\"{dest_label}\"]:::downstream\n"

                edge_style = self._get_edge_style(flow_count)
                mermaid += f"    CENTER {edge_style}|{flow_count}| {dest_node}\n"

        # External destinations with IP (hostname)
        external_flows = len(flows_df[flows_df['Flow Direction'] == 'external'])
        if external_flows > 0:
            external_data = flows_df[flows_df['Flow Direction'] == 'external']

            # Build IP (hostname) format for external
            ext_ip_hostname_list = []
            for _, row in external_data[['Dest IP', 'Dest Hostname']].drop_duplicates().iterrows():
                ip = row['Dest IP']
                hostname = row.get('Dest Hostname', ip)
                if hostname and hostname != ip and hostname != 'Unknown':
                    ext_ip_hostname_list.append(f"{ip} ({hostname})")
                else:
                    ext_ip_hostname_list.append(ip)

            # Show first 2 entries
            ext_str = '<br/>'.join(ext_ip_hostname_list[:2])
            if len(ext_ip_hostname_list) > 2:
                ext_str += f"<br/>... (+{len(ext_ip_hostname_list) - 2} more)"

            external_ips = external_data['Dest IP'].nunique()
            external_label = f"<b>External / Internet</b><br/>{ext_str}<br/>{external_ips} Destinations | {external_flows} Flows"
            mermaid += f"    EXTERNAL[\"{external_label}\"]:::external\n"
            edge_style = self._get_edge_style(external_flows)
            mermaid += f"    CENTER {edge_style}|{external_flows}| EXTERNAL\n"

        # Add styles (thinner strokes for professional appearance)
        mermaid += "\n    %% Styles\n"
        mermaid += "    classDef centerapp fill:#4CAF50,stroke:#2E7D32,stroke-width:2px,color:#fff\n"
        mermaid += "    classDef upstream fill:#2196F3,stroke:#1565C0,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef downstream fill:#9C27B0,stroke:#6A1B9A,stroke-width:1.5px,color:#fff\n"
        mermaid += "    classDef external fill:#F44336,stroke:#C62828,stroke-width:1.5px,color:#fff\n"

        return mermaid

    def _generate_upstream_diagram(self, app_id: str) -> str:
        """Generate UPSTREAM diagram - CLASSIC AGGREGATED VIEW (who depends on THIS app)"""
        # Load ALL flows from ALL applications
        all_flows = []
        for app_dir in self.apps_dir.iterdir():
            if not app_dir.is_dir():
                continue
            flows_csv = app_dir / 'flows.csv'
            if flows_csv.exists():
                try:
                    df = pd.read_csv(flows_csv)
                    all_flows.append(df)
                except:
                    continue

        if not all_flows:
            return f"graph LR\n    NO_DATA[\"No upstream dependencies found\"]"

        all_flows_df = pd.concat(all_flows, ignore_index=True)

        # Find flows where THIS app is the DESTINATION (upstream = who sends TO us)
        upstream_flows = all_flows_df[all_flows_df['Dest App'] == app_id]

        if len(upstream_flows) == 0:
            return f"graph LR\n    NO_UPSTREAM[\"No upstream dependencies found<br/>No other applications send data to {app_id}\"]"

        # Build CLASSIC aggregated diagram
        mermaid = "graph LR\n"
        mermaid += f"    %% UPSTREAM: Who sends data TO {app_id}\n\n"

        # Group by source application and count
        source_apps = sorted(upstream_flows['App'].unique())

        for idx, source_app in enumerate(source_apps):
            app_flows = upstream_flows[upstream_flows['App'] == source_app]

            # Count stats
            server_count = app_flows['Source IP'].nunique()
            flow_count = len(app_flows)
            vmware_count = app_flows[app_flows.get('Source Is VMware', False) == True]['Source IP'].nunique() if 'Source Is VMware' in app_flows.columns else 0

            # Build IP (hostname) format for source IPs
            src_ip_hostname_list = []
            for _, row in app_flows[['Source IP', 'Source Hostname']].drop_duplicates().iterrows():
                ip = row['Source IP']
                hostname = row.get('Source Hostname', ip)
                if hostname and hostname != ip and hostname != 'Unknown':
                    src_ip_hostname_list.append(f"{ip} ({hostname})")
                else:
                    src_ip_hostname_list.append(ip)

            # Show first 2 entries
            src_str = '<br/>'.join(src_ip_hostname_list[:2])
            if len(src_ip_hostname_list) > 2:
                src_str += f"<br/>... (+{len(src_ip_hostname_list) - 2} more)"

            # Create aggregated node
            app_node = f"SRC_APP_{idx}"
            vmware_text = f" | {vmware_count} VMware" if vmware_count > 0 else ""
            label = f"<b>{source_app}</b><br/>{src_str}<br/>{server_count} Servers | {flow_count} Flows{vmware_text}"

            if vmware_count > 0:
                mermaid += f"    {app_node}[(\"{label}\")]:::vmware\n"
            else:
                mermaid += f"    {app_node}[\"{label}\"]:::sourceapp\n"

        # Add THIS application as target (aggregated) with IP (hostname)
        dest_server_count = upstream_flows['Dest IP'].nunique()
        dest_flow_count = len(upstream_flows)
        dest_vmware = upstream_flows[upstream_flows.get('Dest Is VMware', False) == True]['Dest IP'].nunique() if 'Dest Is VMware' in upstream_flows.columns else 0

        # Build IP (hostname) format for target IPs
        dest_ip_hostname_list = []
        for _, row in upstream_flows[['Dest IP', 'Dest Hostname']].drop_duplicates().iterrows():
            ip = row['Dest IP']
            hostname = row.get('Dest Hostname', ip)
            if hostname and hostname != ip and hostname != 'Unknown':
                dest_ip_hostname_list.append(f"{ip} ({hostname})")
            else:
                dest_ip_hostname_list.append(ip)

        # Show first 2 entries
        dest_str = '<br/>'.join(dest_ip_hostname_list[:2])
        if len(dest_ip_hostname_list) > 2:
            dest_str += f"<br/>... (+{len(dest_ip_hostname_list) - 2} more)"

        vmware_text = f" | {dest_vmware} VMware" if dest_vmware > 0 else ""
        target_label = f"<b>{app_id} (TARGET)</b><br/>{dest_str}<br/>{dest_server_count} Servers | {dest_flow_count} Flows{vmware_text}"

        if dest_vmware > 0:
            mermaid += f"    TARGET[(\"{target_label}\")]:::targetapp\n\n"
        else:
            mermaid += f"    TARGET[\"{target_label}\"]:::targetapp\n\n"

        # Connect source apps TO this app with flow counts
        for idx, source_app in enumerate(source_apps):
            app_flows = upstream_flows[upstream_flows['App'] == source_app]
            flow_count = len(app_flows)
            app_node = f"SRC_APP_{idx}"
            mermaid += f"    {app_node} -->|{flow_count}| TARGET\n"

        # Add styles
        mermaid += "\n    %% Styles\n"
        mermaid += "    classDef targetapp fill:#4CAF50,stroke:#2E7D32,stroke-width:4px,color:#fff\n"
        mermaid += "    classDef sourceapp fill:#2196F3,stroke:#1565C0,stroke-width:3px,color:#fff\n"
        mermaid += "    classDef vmware fill:#FF9800,stroke:#E65100,stroke-width:3px,color:#fff\n"

        return mermaid

    def _generate_interactive_html(self, app_id: str, downstream_mermaid: str, upstream_mermaid: str, fullflow_mermaid: str, architecture_mermaid: str, flows_df: pd.DataFrame) -> str:
        """Generate interactive HTML with toggle for Full Flow/Upstream/Downstream/Architecture views"""

        # Calculate statistics
        total_flows = len(flows_df)
        internal_flows = len(flows_df[flows_df['Flow Direction'] == 'internal'])
        outbound_flows = len(flows_df[flows_df['Flow Direction'] == 'outbound'])
        external_flows = len(flows_df[flows_df['Flow Direction'] == 'external'])
        vmware_sources = len(flows_df[flows_df['Source Is VMware'] == True])
        vmware_dests = len(flows_df[flows_df['Dest Is VMware'] == True])
        failed_dns = len(flows_df[flows_df['Dest DNS Status'] == 'NXDOMAIN'])
        dest_apps = flows_df[flows_df['Dest App'].notna()]['Dest App'].nunique()

        # Build detailed data for modal popups (ALL IPs/hostnames with server types)
        import json
        node_details = {}

        # Source node (this app) - with server type detection
        source_servers = []
        for _, row in flows_df[['Source IP', 'Source Hostname', 'Port']].drop_duplicates().iterrows():
            ip = row['Source IP']
            hostname = row.get('Source Hostname', ip)
            port = row.get('Port')

            # Detect server type
            ports = flows_df[flows_df['Source IP'] == ip]['Port'].unique().tolist() if 'Port' in flows_df.columns else []
            server_type = self._detect_server_type(hostname, ports)

            if hostname and hostname != ip and hostname != 'Unknown':
                server_entry = {
                    'ip': ip,
                    'hostname': hostname,
                    'type': server_type,
                    'display': f"{ip} ({hostname})"
                }
            else:
                server_entry = {
                    'ip': ip,
                    'hostname': ip,
                    'type': server_type,
                    'display': ip
                }
            source_servers.append(server_entry)

        node_details['SOURCE'] = {
            'title': f'{app_id} Source Servers',
            'items': source_servers
        }

        # Architecture tier data - organized by server type
        # Build server_types dictionary for architecture modal
        server_types_for_modal = {}
        for _, row in flows_df[['Source IP', 'Source Hostname', 'Port']].drop_duplicates().iterrows():
            ip = row['Source IP']
            hostname = row.get('Source Hostname', ip)
            ports = flows_df[flows_df['Source IP'] == ip]['Port'].unique().tolist()
            server_type = self._detect_server_type(hostname, ports)

            if server_type not in server_types_for_modal:
                server_types_for_modal[server_type] = []

            if hostname and hostname != ip and hostname != 'Unknown':
                server_entry = {
                    'ip': ip,
                    'hostname': hostname,
                    'type': server_type,
                    'display': f"{ip} ({hostname})"
                }
            else:
                server_entry = {
                    'ip': ip,
                    'hostname': ip,
                    'type': server_type,
                    'display': ip
                }
            server_types_for_modal[server_type].append(server_entry)

        # Add Web servers to node_details
        if 'Web' in server_types_for_modal:
            node_details['ARCH_WEB'] = {
                'title': f'🌐 {app_id} Web Tier Servers',
                'items': server_types_for_modal['Web']
            }

        # Add App servers to node_details
        if 'App' in server_types_for_modal:
            node_details['ARCH_APP'] = {
                'title': f'⚙️ {app_id} Application Tier Servers',
                'items': server_types_for_modal['App']
            }

        # Add Database servers to node_details
        if 'Database' in server_types_for_modal:
            node_details['ARCH_DB'] = {
                'title': f'🗄️ {app_id} Database Tier Servers',
                'items': server_types_for_modal['Database']
            }

        # Destination apps - with server type detection
        outbound_data = flows_df[flows_df['Flow Direction'] == 'outbound']
        if len(outbound_data) > 0:
            for dest_app in outbound_data['Dest App'].unique():
                dest_app_data = outbound_data[outbound_data['Dest App'] == dest_app]
                dest_servers = []
                for _, row in dest_app_data[['Dest IP', 'Dest Hostname', 'Port']].drop_duplicates().iterrows():
                    ip = row['Dest IP']
                    hostname = row.get('Dest Hostname', ip)
                    port = row.get('Port')

                    # Detect server type
                    ports = dest_app_data[dest_app_data['Dest IP'] == ip]['Port'].unique().tolist()
                    server_type = self._detect_server_type(hostname, ports)

                    if hostname and hostname != ip and hostname != 'Unknown':
                        server_entry = {
                            'ip': ip,
                            'hostname': hostname,
                            'type': server_type,
                            'display': f"{ip} ({hostname})"
                        }
                    else:
                        server_entry = {
                            'ip': ip,
                            'hostname': ip,
                            'type': server_type,
                            'display': ip
                        }
                    dest_servers.append(server_entry)

                node_details[f'DEST_{dest_app}'] = {
                    'title': f'{dest_app} Servers',
                    'items': dest_servers
                }

        # External destinations - with server type detection
        external_data = flows_df[flows_df['Flow Direction'] == 'external']
        if len(external_data) > 0:
            ext_servers = []
            for _, row in external_data[['Dest IP', 'Dest Hostname', 'Port']].drop_duplicates().iterrows():
                ip = row['Dest IP']
                hostname = row.get('Dest Hostname', ip)
                port = row.get('Port')

                # Detect server type
                ports = external_data[external_data['Dest IP'] == ip]['Port'].unique().tolist()
                server_type = self._detect_server_type(hostname, ports)

                if hostname and hostname != ip and hostname != 'Unknown':
                    server_entry = {
                        'ip': ip,
                        'hostname': hostname,
                        'type': server_type,
                        'display': f"{ip} ({hostname})"
                    }
                else:
                    server_entry = {
                        'ip': ip,
                        'hostname': ip,
                        'type': server_type,
                        'display': ip
                    }
                ext_servers.append(server_entry)

            node_details['EXTERNAL'] = {
                'title': 'External / Internet Destinations',
                'items': ext_servers
            }

        node_details_json = json.dumps(node_details, indent=2)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{app_id} - Application Architecture</title>
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }}

        .container {{
            max-width: 1400px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            overflow: hidden;
        }}

        .header {{
            background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
            color: white;
            padding: 30px;
            text-align: center;
        }}

        .header h1 {{
            font-size: 2.5em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }}

        .header .subtitle {{
            font-size: 1.2em;
            opacity: 0.9;
        }}

        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
            gap: 10px;
            padding: 20px 30px;
            background: #f8f9fa;
        }}

        .stat-card {{
            background: white;
            padding: 12px 15px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.08);
            text-align: center;
            transition: transform 0.2s;
        }}

        .stat-card:hover {{
            transform: translateY(-3px);
            box-shadow: 0 4px 15px rgba(0,0,0,0.12);
        }}

        .stat-card .value {{
            font-size: 1.8em;
            font-weight: bold;
            color: #2a5298;
            margin-bottom: 3px;
        }}

        .stat-card .label {{
            color: #666;
            font-size: 0.75em;
            line-height: 1.2;
        }}

        .controls {{
            padding: 20px 30px;
            background: #fff;
            border-bottom: 2px solid #e0e0e0;
            display: flex;
            gap: 15px;
            flex-wrap: wrap;
            align-items: center;
        }}

        .btn {{
            padding: 10px 20px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
            transition: all 0.3s;
            font-weight: 500;
        }}

        .btn-primary {{
            background: #2a5298;
            color: white;
        }}

        .btn-primary:hover {{
            background: #1e3c72;
            transform: scale(1.05);
        }}

        .btn-secondary {{
            background: #6c757d;
            color: white;
        }}

        .btn-secondary:hover {{
            background: #5a6268;
        }}

        .btn-toggle {{
            padding: 12px 24px;
            border: 2px solid #2a5298;
            background: white;
            color: #2a5298;
            font-weight: 600;
            border-radius: 5px;
            cursor: pointer;
            font-size: 15px;
            transition: all 0.3s;
        }}

        .btn-toggle:hover {{
            background: #e3f2fd;
        }}

        .btn-toggle.active {{
            background: #2a5298;
            color: white;
        }}

        .toggle-section {{
            padding: 20px 30px;
            background: #f8f9fa;
            border-bottom: 2px solid #e0e0e0;
            display: flex;
            gap: 10px;
            justify-content: center;
            flex-wrap: wrap;
        }}

        .diagram-container {{
            padding: 30px;
            background: white;
            overflow-x: auto;
            overflow-y: hidden;
            width: 100%;
        }}

        .diagram-view {{
            min-height: 600px;
            background: white;
            width: 100%;
            overflow: visible;
        }}

        .diagram-view .mermaid {{
            width: 100%;
            max-width: 100%;
            overflow: visible;
        }}

        .diagram-view .mermaid svg {{
            max-width: 100% !important;
            height: auto !important;
        }}

        .diagram-view.hidden {{
            display: none;
        }}

        /* Professional uniform box sizing for Mermaid */
        #diagram svg .node rect,
        #diagram svg .node circle,
        #diagram svg .node ellipse,
        #diagram svg .node polygon,
        #diagram svg .node path {{
            min-width: 180px !important;
            min-height: 80px !important;
        }}

        .mermaid svg {{
            max-width: 100%;
            height: auto;
        }}

        /* Uniform node sizing */
        .mermaid .node {{
            min-width: 200px !important;
            text-align: center !important;
        }}

        .mermaid .nodeLabel {{
            padding: 15px 20px !important;
            line-height: 1.4 !important;
        }}

        /* Uniform text sizing */
        .mermaid text {{
            font-size: 14px !important;
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif !important;
        }}

        .legend {{
            padding: 15px 30px;
            background: #f8f9fa;
            border-top: 2px solid #e0e0e0;
        }}

        .legend h3 {{
            color: #2a5298;
            margin-bottom: 10px;
            font-size: 1.1em;
        }}

        .legend-items {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 8px;
        }}

        .legend-item {{
            display: flex;
            align-items: center;
            gap: 8px;
            font-size: 0.85em;
        }}

        .legend-color {{
            width: 20px;
            height: 20px;
            border-radius: 4px;
            border: 2px solid #ddd;
            flex-shrink: 0;
        }}

        .footer {{
            padding: 20px 30px;
            background: #2a5298;
            color: white;
            text-align: center;
        }}

        @media print {{
            body {{
                background: white;
            }}
            .controls, .stats {{
                display: none;
            }}
        }}
    </style>
</head>
<body>
    <!-- Modal for showing all IPs/hostnames -->
    <div id="detailsModal" style="display: none; position: fixed; z-index: 10000; left: 0; top: 0; width: 100%; height: 100%; overflow: auto; background-color: rgba(0,0,0,0.6);">
        <div style="background-color: white; margin: 5% auto; padding: 0; border-radius: 10px; width: 80%; max-width: 800px; box-shadow: 0 5px 30px rgba(0,0,0,0.3);">
            <div style="background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); color: white; padding: 20px; border-radius: 10px 10px 0 0; display: flex; justify-content: space-between; align-items: center;">
                <h2 id="modalTitle" style="margin: 0; font-size: 1.5em;">Server Details</h2>
                <span onclick="closeModal()" style="cursor: pointer; font-size: 28px; font-weight: bold;">&times;</span>
            </div>
            <div id="modalContent" style="padding: 30px;">
                <!-- Content will be populated by JavaScript -->
            </div>
        </div>
    </div>

    <div class="container">
        <div class="header">
            <h1>{app_id}</h1>
            <div class="subtitle">Application Architecture & Dependencies</div>
            <div class="subtitle">Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</div>
        </div>

        <div class="stats">
            <div class="stat-card">
                <div class="value">{total_flows}</div>
                <div class="label">Total Flows</div>
            </div>
            <div class="stat-card">
                <div class="value">{internal_flows}</div>
                <div class="label">Internal Flows</div>
            </div>
            <div class="stat-card">
                <div class="value">{outbound_flows}</div>
                <div class="label">Outbound to Apps</div>
            </div>
            <div class="stat-card">
                <div class="value">{external_flows}</div>
                <div class="label">External Flows</div>
            </div>
            <div class="stat-card">
                <div class="value">{dest_apps}</div>
                <div class="label">Destination Apps</div>
            </div>
            <div class="stat-card">
                <div class="value">{vmware_sources + vmware_dests}</div>
                <div class="label">VMware Instances</div>
            </div>
            <div class="stat-card">
                <div class="value">{failed_dns}</div>
                <div class="label">Failed DNS (RED)</div>
            </div>
        </div>

        <div class="toggle-section">
            <button class="btn-toggle active" onclick="showView('architecture')" id="btn-architecture">
                🏗️ Architecture (Tier-Based)
            </button>
            <button class="btn-toggle" onclick="showView('fullflow')" id="btn-fullflow">
                🔄 Full Flow (End-to-End)
            </button>
            <button class="btn-toggle" onclick="showView('upstream')" id="btn-upstream">
                ⬆️ Upstream Only
            </button>
            <button class="btn-toggle" onclick="showView('downstream')" id="btn-downstream">
                ⬇️ Downstream Only
            </button>
        </div>

        <div class="controls">
            <button class="btn btn-primary" onclick="zoomIn()">🔍 Zoom In</button>
            <button class="btn btn-primary" onclick="zoomOut()">🔍 Zoom Out</button>
            <button class="btn btn-primary" onclick="resetZoom()">↺ Reset</button>
            <button class="btn btn-secondary" onclick="window.print()">🖨️ Print</button>
            <button class="btn btn-secondary" onclick="downloadSVG()">⬇️ Download SVG</button>
            <button class="btn btn-secondary" onclick="downloadPNG()">⬇️ Download PNG</button>
        </div>

        <div class="diagram-container">
            <div id="view-architecture" class="diagram-view" style="display: block;">
                <div style="text-align: center; padding: 20px 0; background: linear-gradient(135deg, #4CAF50 0%, #2E7D32 100%); border-radius: 8px; margin-bottom: 20px;">
                    <h3 style="color: #fff; margin: 0;">🏗️ ARCHITECTURE: {app_id} Vertical Tier-Based Layout (Optimized for Large Deployments)</h3>
                </div>
                <div class="mermaid">
{architecture_mermaid}
                </div>
            </div>

            <div id="view-fullflow" class="diagram-view" style="display: block;">
                <div style="text-align: center; padding: 20px 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 8px; margin-bottom: 20px;">
                    <h3 style="color: #fff; margin: 0;">🔄 FULL FLOW: Source Apps → {app_id} → Destination Apps</h3>
                </div>
                <div class="mermaid">
{fullflow_mermaid}
                </div>
            </div>

            <div id="view-upstream" class="diagram-view" style="display: block;">
                <div style="text-align: center; padding: 20px 0; background: #e3f2fd; border-radius: 8px; margin-bottom: 20px;">
                    <h3 style="color: #2a5298; margin: 0;">⬆️ UPSTREAM ONLY: Who Sends Data TO {app_id}</h3>
                </div>
                <div class="mermaid">
{upstream_mermaid}
                </div>
            </div>

            <div id="view-downstream" class="diagram-view" style="display: block;">
                <div style="text-align: center; padding: 20px 0; background: #fff3e0; border-radius: 8px; margin-bottom: 20px;">
                    <h3 style="color: #2a5298; margin: 0;">⬇️ DOWNSTREAM ONLY: Where {app_id} Sends Data</h3>
                </div>
                <div class="mermaid">
{downstream_mermaid}
                </div>
            </div>
        </div>

        <div class="legend">
            <h3>Legend</h3>
            <div class="legend-items">
                <div class="legend-item">
                    <div class="legend-color" style="background: #4CAF50; border-color: #2E7D32;"></div>
                    <span>Source Application</span>
                </div>
                <div class="legend-item">
                    <div class="legend-color" style="background: #E3F2FD; border-color: #1976D2;"></div>
                    <span>Normal Server</span>
                </div>
                <div class="legend-item">
                    <div class="legend-color" style="background: #FF9800; border-color: #E65100;"></div>
                    <span>VMware/ESXi Server</span>
                </div>
                <div class="legend-item">
                    <div class="legend-color" style="background: #FFF; border: 4px solid #F44336;"></div>
                    <span>Load Balancer (F5/ALB)</span>
                </div>
                <div class="legend-item">
                    <div class="legend-color" style="background: #F44336; border-color: #C62828;"></div>
                    <span>External/Internet</span>
                </div>
                <div class="legend-item">
                    <div class="legend-color" style="background: #F1F8E9; border-color: #689F38;"></div>
                    <span>Internal Communication</span>
                </div>
                <div class="legend-item">
                    <div class="legend-color" style="background: #FCE4EC; border-color: #C2185B;"></div>
                    <span>External/Internet</span>
                </div>
            </div>
        </div>

        <div class="footer">
            &copy; {datetime.now().year} Enterprise Security Team | Network Segmentation Analysis
        </div>
    </div>

    <script>
        mermaid.initialize({{
            startOnLoad: true,
            theme: 'default',
            securityLevel: 'loose',
            themeVariables: {{
                fontSize: '13px',
                fontFamily: 'Inter, -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif',
                primaryColor: '#fff',
                primaryTextColor: '#000',
                lineColor: '#333',
                edgeLabelBackground: '#ffffff'
            }},
            flowchart: {{
                nodeSpacing: 50,
                rankSpacing: 80,
                padding: 15,
                useMaxWidth: true,
                htmlLabels: true,
                curve: 'basis'
            }}
        }});

        let currentZoom = 1;
        let currentView = 'architecture';

        // After Mermaid renders, hide the non-active views
        window.addEventListener('load', function() {{
            setTimeout(function() {{
                document.getElementById('view-fullflow').style.display = 'none';
                document.getElementById('view-upstream').style.display = 'none';
                document.getElementById('view-downstream').style.display = 'none';
            }}, 500);
        }});

        function showView(view) {{
            // Hide all views
            document.getElementById('view-architecture').style.display = 'none';
            document.getElementById('view-fullflow').style.display = 'none';
            document.getElementById('view-upstream').style.display = 'none';
            document.getElementById('view-downstream').style.display = 'none';

            // Remove active from all buttons
            document.getElementById('btn-architecture').classList.remove('active');
            document.getElementById('btn-fullflow').classList.remove('active');
            document.getElementById('btn-upstream').classList.remove('active');
            document.getElementById('btn-downstream').classList.remove('active');

            // Show selected view
            document.getElementById('view-' + view).style.display = 'block';
            document.getElementById('btn-' + view).classList.add('active');

            currentView = view;
            resetZoom();
        }}

        function zoomIn() {{
            currentZoom += 0.1;
            const activeView = document.getElementById('view-' + currentView);
            if (activeView) {{
                activeView.style.transform = `scale(${{currentZoom}})`;
            }}
        }}

        function zoomOut() {{
            currentZoom = Math.max(0.5, currentZoom - 0.1);
            const activeView = document.getElementById('view-' + currentView);
            if (activeView) {{
                activeView.style.transform = `scale(${{currentZoom}})`;
            }}
        }}

        function resetZoom() {{
            currentZoom = 1;
            const activeView = document.getElementById('view-' + currentView);
            if (activeView) {{
                activeView.style.transform = 'scale(1)';
            }}
        }}

        function downloadSVG() {{
            const svg = document.querySelector('#diagram svg');
            if (svg) {{
                const svgData = new XMLSerializer().serializeToString(svg);
                const blob = new Blob([svgData], {{ type: 'image/svg+xml' }});
                const url = URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = '{app_id}_architecture.svg';
                a.click();
            }}
        }}

        function downloadPNG() {{
            const svg = document.querySelector('#diagram svg');
            if (svg) {{
                const canvas = document.createElement('canvas');
                const ctx = canvas.getContext('2d');
                const img = new Image();
                const svgData = new XMLSerializer().serializeToString(svg);
                const blob = new Blob([svgData], {{ type: 'image/svg+xml' }});
                const url = URL.createObjectURL(blob);

                img.onload = function() {{
                    canvas.width = img.width * 2;
                    canvas.height = img.height * 2;
                    ctx.drawImage(img, 0, 0, canvas.width, canvas.height);
                    canvas.toBlob(function(blob) {{
                        const pngUrl = URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.href = pngUrl;
                        a.download = '{app_id}_architecture.png';
                        a.click();
                    }});
                }};
                img.src = url;
            }}
        }}

        // Node details data for modal popups
        const nodeDetails = {node_details_json};

        // Add click handlers to all diagram nodes
        window.addEventListener('load', function() {{
            setTimeout(function() {{
                // Get all node elements from Mermaid diagrams
                const nodes = document.querySelectorAll('.node');
                nodes.forEach(function(node) {{
                    // Make nodes clickable if they have "+X more"
                    const label = node.querySelector('.nodeLabel');
                    if (label && label.textContent.includes('(+')) {{
                        node.style.cursor = 'pointer';
                        node.addEventListener('click', function(e) {{
                            e.stopPropagation();
                            showNodeDetails(node);
                        }});
                    }}
                }});
            }}, 1000);
        }});

        function showNodeDetails(node) {{
            const label = node.querySelector('.nodeLabel');
            if (!label) return;

            const text = label.textContent;

            // Try to determine which node this is
            let nodeKey = null;
            let details = null;

            // Check if it's SOURCE
            if (text.includes('{app_id}') && text.includes('Servers')) {{
                nodeKey = 'SOURCE';
                details = nodeDetails.SOURCE;
            }}
            // Check for destination apps
            else {{
                for (let key in nodeDetails) {{
                    if (key.startsWith('DEST_')) {{
                        const appName = key.replace('DEST_', '');
                        if (text.includes(appName)) {{
                            details = nodeDetails[key];
                            break;
                        }}
                    }}
                }}
            }}
            // Check for EXTERNAL
            if (!details && (text.includes('External') || text.includes('Internet'))) {{
                details = nodeDetails.EXTERNAL;
            }}

            if (details) {{
                showModal(details.title, details.items);
            }}
        }}

        function showModal(title, items) {{
            const modal = document.getElementById('detailsModal');
            const modalTitle = document.getElementById('modalTitle');
            const modalContent = document.getElementById('modalContent');

            modalTitle.textContent = title;

            // Server type icons/emojis
            const typeIcons = {{
                'Web': '🌐',
                'Database': '🗄️',
                'App': '⚙️',
                'DNS': '🔍',
                'Mail': '📧',
                'File': '📁',
                'Domain Controller': '👑',
                'Cache': '💾',
                'Message Queue': '📨',
                'Monitoring': '📊',
                'Backup': '💿',
                'General': '🖥️'
            }};

            // Build the list with server types
            let html = '<div style="max-height: 500px; overflow-y: auto;">';
            html += '<table style="width: 100%; border-collapse: collapse;">';
            html += '<tr style="background: #2a5298; color: white;">';
            html += '<th style="padding: 10px; text-align: left;">#</th>';
            html += '<th style="padding: 10px; text-align: left;">Type</th>';
            html += '<th style="padding: 10px; text-align: left;">IP (Hostname)</th>';
            html += '</tr>';

            items.forEach(function(item, idx) {{
                const bg = idx % 2 === 0 ? '#f8f9fa' : 'white';
                const serverType = item.type || 'General';
                const icon = typeIcons[serverType] || '🖥️';
                const display = item.display || item;

                html += `<tr style="background: ${{bg}};">`;
                html += `<td style="padding: 8px; border-bottom: 1px solid #ddd;">${{idx + 1}}</td>`;
                html += `<td style="padding: 8px; border-bottom: 1px solid #ddd;"><span style="font-size: 16px;">${{icon}}</span> ${{serverType}}</td>`;
                html += `<td style="padding: 8px; border-bottom: 1px solid #ddd; font-family: monospace;">${{display}}</td>`;
                html += `</tr>`;
            }});
            html += '</table></div>';

            // Count server types
            const typeCounts = {{}};
            items.forEach(function(item) {{
                const type = item.type || 'General';
                typeCounts[type] = (typeCounts[type] || 0) + 1;
            }});

            html += `<div style="margin-top: 15px; padding: 15px; background: #e3f2fd; border-radius: 5px;">`;
            html += `<strong>Total:</strong> ${{items.length}} servers<br/><br/>`;
            html += `<strong>Breakdown by Type:</strong><br/>`;
            for (let type in typeCounts) {{
                const icon = typeIcons[type] || '🖥️';
                html += `<span style="display: inline-block; margin: 5px 10px 5px 0;">${{icon}} ${{type}}: ${{typeCounts[type]}}</span>`;
            }}
            html += `</div>`;

            modalContent.innerHTML = html;
            modal.style.display = 'block';
        }}

        function closeModal() {{
            document.getElementById('detailsModal').style.display = 'none';
        }}

        // Architecture diagram click handlers
        function showWebServers() {{
            const data = nodeDetails['ARCH_WEB'];
            if (data) {{
                showModal(data.title, data.items);
            }}
        }}

        function showAppServers() {{
            const data = nodeDetails['ARCH_APP'];
            if (data) {{
                showModal(data.title, data.items);
            }}
        }}

        function showDatabaseServers() {{
            const data = nodeDetails['ARCH_DB'];
            if (data) {{
                showModal(data.title, data.items);
            }}
        }}

        // Close modal when clicking outside
        window.onclick = function(event) {{
            const modal = document.getElementById('detailsModal');
            if (event.target == modal) {{
                modal.style.display = 'none';
            }}
        }}
    </script>
</body>
</html>"""

        return html

    def _generate_png_from_mermaid(self, mmd_file: Path, png_file: Path) -> bool:
        """Generate PNG using mmdc CLI"""
        try:
            subprocess.run([
                'mmdc',
                '-i', str(mmd_file),
                '-o', str(png_file),
                '-b', 'white',
                '-w', '2400',
                '-H', '1800'
            ], check=True, capture_output=True)
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False

    def _generate_svg_from_mermaid(self, mmd_file: Path, svg_file: Path) -> bool:
        """Generate SVG using mmdc CLI"""
        try:
            subprocess.run([
                'mmdc',
                '-i', str(mmd_file),
                '-o', str(svg_file),
                '-b', 'white'
            ], check=True, capture_output=True)
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False

    def _generate_architecture_docx(self, app_id: str, flows_df: pd.DataFrame, output_file: Path):
        """Generate Architecture DOCX documentation"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f'{app_id} - Architecture Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph(f'Total Flows: {len(flows_df)}')
        doc.add_page_break()

        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(f'This document provides a comprehensive architecture overview of the {app_id} application, including its internal structure, dependencies, and network flows.')

        # Statistics
        doc.add_heading('Flow Statistics', 2)
        internal = len(flows_df[flows_df['Flow Direction'] == 'internal'])
        outbound = len(flows_df[flows_df['Flow Direction'] == 'outbound'])
        external = len(flows_df[flows_df['Flow Direction'] == 'external'])

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        table.cell(0, 0).text = 'Flow Type'
        table.cell(0, 1).text = 'Count'
        table.cell(1, 0).text = 'Internal'
        table.cell(1, 1).text = str(internal)
        table.cell(2, 0).text = 'Outbound (to other apps)'
        table.cell(2, 1).text = str(outbound)
        table.cell(3, 0).text = 'External (internet)'
        table.cell(3, 1).text = str(external)

        # Dependencies
        doc.add_heading('Application Dependencies', 2)
        dest_apps = flows_df[flows_df['Dest App'].notna()]['Dest App'].unique()
        if len(dest_apps) > 0:
            doc.add_paragraph(f'{app_id} depends on the following applications:')
            for dest_app in dest_apps:
                count = len(flows_df[flows_df['Dest App'] == dest_app])
                doc.add_paragraph(f'• {dest_app} ({count} flows)', style='List Bullet')
        else:
            doc.add_paragraph('No application dependencies detected.')

        # VMware Infrastructure
        doc.add_heading('VMware Infrastructure', 2)
        vmware_flows = flows_df[(flows_df['Source Is VMware'] == True) | (flows_df['Dest Is VMware'] == True)]
        if len(vmware_flows) > 0:
            doc.add_paragraph(f'Detected {len(vmware_flows)} VMware-related flows:')
            for _, row in vmware_flows.head(10).iterrows():
                if row['Source Is VMware']:
                    doc.add_paragraph(f'• {row["Source Hostname (Full)"]}', style='List Bullet')
        else:
            doc.add_paragraph('No VMware infrastructure detected.')

        doc.save(str(output_file))

    def _generate_threat_docx(self, app_id: str, flows_df: pd.DataFrame, output_file: Path):
        """Generate Threat Surface DOCX analysis"""
        from docx import Document

        doc = Document()
        doc.add_heading(f'{app_id} - Threat Surface Analysis', 0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_page_break()

        # Threat Summary
        doc.add_heading('Threat Summary', 1)

        # External Exposure
        external_flows = flows_df[flows_df['Flow Direction'] == 'external']
        doc.add_heading('External Exposure', 2)
        doc.add_paragraph(f'Total external flows: {len(external_flows)}')

        # Failed DNS (Unknown Servers)
        failed_dns = flows_df[flows_df['Dest DNS Status'] == 'NXDOMAIN']
        doc.add_heading('Unknown Destinations (Failed DNS)', 2)
        doc.add_paragraph(f'Found {len(failed_dns)} flows to unknown destinations:')
        for _, row in failed_dns.head(20).iterrows():
            doc.add_paragraph(f'• {row["Dest IP"]} (Port: {row["Port"]}, Protocol: {row["Protocol"]})', style='List Bullet')

        doc.save(str(output_file))

    def _generate_segmentation_json(self, app_id: str, flows_df: pd.DataFrame, output_file: Path):
        """Generate Network Segmentation JSON"""
        segmentation = {
            'application': app_id,
            'generated': datetime.now().isoformat(),
            'statistics': {
                'total_flows': len(flows_df),
                'internal_flows': len(flows_df[flows_df['Flow Direction'] == 'internal']),
                'outbound_flows': len(flows_df[flows_df['Flow Direction'] == 'outbound']),
                'external_flows': len(flows_df[flows_df['Flow Direction'] == 'external'])
            },
            'dependencies': flows_df[flows_df['Dest App'].notna()]['Dest App'].unique().tolist(),
            'vmware_instances': len(flows_df[(flows_df['Source Is VMware'] == True) | (flows_df['Dest Is VMware'] == True)]),
            'failed_dns_count': len(flows_df[flows_df['Dest DNS Status'] == 'NXDOMAIN'])
        }

        with open(output_file, 'w') as f:
            json.dump(segmentation, f, indent=2)

    def process_all_apps(self, batch_size: int = 10):
        """Process all applications in batches"""
        logger.info("="*80)
        logger.info("COMPLETE REPORT GENERATION - FORTINET STYLE")
        logger.info("="*80)

        # Get all app directories
        app_dirs = sorted([d for d in self.apps_dir.iterdir() if d.is_dir()])
        logger.info(f"Found {len(app_dirs)} applications")
        logger.info(f"Processing in batches of {batch_size}")
        logger.info("")

        total_outputs = {}
        processed = 0

        for app_dir in app_dirs:
            flows_csv = app_dir / 'flows.csv'
            if not flows_csv.exists():
                continue

            app_id = app_dir.name

            try:
                # Read enriched flows
                flows_df = pd.read_csv(flows_csv)

                # Generate all outputs
                outputs = self.generate_for_app(app_id, flows_df)
                total_outputs[app_id] = outputs

                processed += 1

                # Batch checkpoint
                if processed % batch_size == 0:
                    logger.info(f"[CHECKPOINT] Processed {processed}/{len(app_dirs)} apps")

            except Exception as e:
                logger.error(f"[{app_id}] Error: {e}")

        logger.info("")
        logger.info("="*80)
        logger.info(f"COMPLETE! Processed {processed} applications")
        logger.info("="*80)
        logger.info("")
        logger.info("Outputs generated:")
        logger.info(f"  HTML diagrams: {self.html_dir}")
        logger.info(f"  PNG exports: {self.png_dir}")
        logger.info(f"  SVG exports: {self.svg_dir}")
        logger.info(f"  DOCX reports: {self.docx_dir}")
        logger.info(f"  JSON data: {self.json_dir}")
        logger.info(f"  Mermaid source: {self.mmd_dir}")

        return total_outputs


def main():
    """Generate complete reports for all applications"""
    generator = CompleteReportGenerator()

    # Process all apps
    outputs = generator.process_all_apps(batch_size=10)

    logger.info("")
    logger.info(f"Total applications processed: {len(outputs)}")
    logger.info("")
    logger.info("Open any HTML file in your browser to view the interactive diagram!")


if __name__ == '__main__':
    main()

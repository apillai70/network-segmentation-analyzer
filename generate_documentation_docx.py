#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Documentation Consolidation Generator
======================================
Reads all markdown documentation files and generates a comprehensive Word document
with proper sections, table of contents, and formatting.

Author: Enterprise Security Team
Version: 1.0
"""

import logging
from pathlib import Path
from collections import defaultdict
import re

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    logger.error("python-docx not installed. Install with: pip install python-docx")
    exit(1)


class DocumentationConsolidator:
    """Consolidates markdown files into a comprehensive Word document"""

    # Document section categorization
    SECTION_CATEGORIES = {
        'Getting Started': [
            'README.md',
            'GETTING_STARTED.md',
            'QUICKSTART.md',
            'QUICKSTART_INCREMENTAL.md',
            'INSTALLATION_GUIDE.md',
            'SETUP_AND_RUN.md'
        ],
        'Core Features': [
            'README_COMPLETE.md',
            'COMPLETE_SYSTEM_GUIDE.md',
            'WHATS_NEW_V3.md',
            'PROJECT_STATUS_REPORT.md'
        ],
        'Server Classification & Diagrams': [
            'SERVER_CLASSIFICATION_SUMMARY.md',
            'MULTIFORMAT_DIAGRAM_GENERATION.md',
            'MULTIFORMAT_COMMIT_SUMMARY.md',
            'GRAPH_ANALYSIS_README.md',
            'DIAGRAM_LEGEND_VERIFICATION.md',
            'SVG_SOLUTION_CLARIFICATION.md'
        ],
        'Web Application': [
            'WEB_APP_README.md',
            'WEB_APP_SUMMARY.md',
            'INSTALL_WEB_APP.md',
            'QUICK_REFERENCE_WEB_APP.md',
            'FASTAPI_GUIDE.md',
            'UI_UPDATE_SUMMARY.md',
            'COMPACT_LAYOUT_SUMMARY.md'
        ],
        'Database & PostgreSQL': [
            'DATABASE_SETUP.md',
            'DATABASE_TEST_RESULTS.md',
            'SCHEMA_PROTECTION_SUMMARY.md',
            'IMPLEMENTATION_SUMMARY.md',
            'IMPLEMENTATION_STATUS.md'
        ],
        'Production Deployment': [
            'PRODUCTION_GUIDE.md',
            'DEPLOYMENT_GUIDE.md',
            'DEPLOYMENT_CHECKLIST.md',
            'FINAL_DEPLOYMENT_CHECKLIST.md',
            'CUSTOMER_DEPLOYMENT_GUIDE.md',
            'PRODUCTION_DEPLOYMENT_FIX.md',
            'PRODUCTION_FIX_SUMMARY.md',
            'PRODUCTION_FIX_IMMEDIATE.md',
            'PRODUCTION_INSTALL_FIX.md'
        ],
        'Processing & Analysis': [
            'BATCH_PROCESSING_GUIDE.md',
            'REPROCESSING_GUIDE.md',
            'INCREMENTAL_LEARNING_GUIDE.md',
            'MANUAL_TYPING_GUIDE.md',
            'ENABLE_MARKOV_PREDICTIONS.md',
            'ARCHETYPES_DETECTED.md'
        ],
        'DNS & Hostname Resolution': [
            'HOSTNAME_RESOLUTION_GUIDE.md',
            'DNS_LOOKUP_ENABLED.md',
            'DNS_VALIDATION_IMPLEMENTATION.md',
            'DNS_VALIDATION_REPORTS.md',
            'NONEXISTENT_DOMAIN_HANDLING.md'
        ],
        'PNG & Diagram Generation': [
            'PNG_SETUP_GUIDE.md',
            'PNG_GENERATION_OPTIONS.md',
            'README_WORD_DOCS.md',
            'TROUBLESHOOT_PNG_EPERM.md',
            'NODEENV_WRAPPER_GUIDE.md'
        ],
        'Troubleshooting & Fixes': [
            'CRITICAL_FIXES_NEEDED.md',
            'FIXES_APPLIED.md',
            'FIXES_COMPLETE_SUMMARY.md',
            'CLEANUP_GUIDE.md',
            'CLEANUP_AND_FRESH_START.md',
            'FIX_BASHRC_TYPO.md',
            'ENCODING_FIX_GUIDE.md',
            'README_CONFIDENCE_ISSUE.md'
        ],
        'Reference Guides': [
            'COMMAND_REFERENCE.md',
            'QUICK_REFERENCE_CARD.md',
            'README_USER_GUIDE.md',
            'OUTPUTS_GUIDE.md',
            'FILTERING_GUIDE.md',
            'SOLUTION_DESIGN_DOCS_GUIDE.md'
        ],
        'Session Notes & Updates': [
            'SESSION_SUMMARY.md',
            'COMMIT_SUMMARY.md',
            'CRITICAL_NOTES.md',
            'REQUIREMENTS_ROADMAP.md',
            'PYTHON_VERSION_NOTES.md'
        ]
    }

    def __init__(self):
        """Initialize documentation consolidator"""
        self.doc = Document()
        self.setup_styles()
        logger.info("DocumentationConsolidator initialized")

    def setup_styles(self):
        """Set up custom styles for the document"""
        styles = self.doc.styles

        # Title style
        try:
            title_style = styles['Title']
        except KeyError:
            title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(31, 73, 125)

        # Heading styles
        for level in range(1, 6):
            try:
                heading = styles[f'Heading {level}']
                heading.font.color.rgb = RGBColor(31, 73, 125)
            except:
                pass

    def add_cover_page(self):
        """Add professional cover page"""
        logger.info("Adding cover page...")

        # Title
        title = self.doc.add_heading('Network Segmentation Analyzer', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(36)
        title_run.font.color.rgb = RGBColor(31, 73, 125)

        # Subtitle
        subtitle = self.doc.add_paragraph('Complete Documentation Guide')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = RGBColor(68, 114, 196)
        subtitle_run.italic = True

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Feature highlights
        features = self.doc.add_paragraph()
        features.alignment = WD_ALIGN_PARAGRAPH.CENTER
        features_run = features.add_run(
            'Comprehensive System for Network Flow Analysis,\n'
            'Server Classification, and Security Assessment'
        )
        features_run.font.size = Pt(14)
        features_run.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Document metadata
        metadata = [
            ('Version', '3.0'),
            ('Document Type', 'Technical Documentation'),
            ('Generated', '2025-10-22'),
            ('Status', 'Production Ready'),
            ('Total Sections', str(len(self.SECTION_CATEGORIES)))
        ]

        table = self.doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Light Grid Accent 1'

        for idx, (key, value) in enumerate(metadata):
            row = table.rows[idx]
            row.cells[0].text = key
            row.cells[1].text = value
            # Bold the key
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Footer note
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(
            'This document consolidates 76+ markdown files into a comprehensive guide\n'
            'covering installation, configuration, usage, and troubleshooting.'
        )
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(150, 150, 150)

        # Page break
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents page"""
        logger.info("Adding table of contents...")

        heading = self.doc.add_heading('Table of Contents', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.doc.add_paragraph()

        # Add TOC entries for each section
        for section_name in self.SECTION_CATEGORIES.keys():
            toc_entry = self.doc.add_paragraph(f'{section_name}', style='List Bullet')
            toc_entry.runs[0].font.size = Pt(12)
            toc_entry.runs[0].font.bold = True

        self.doc.add_paragraph()

        # Add note about navigation
        note = self.doc.add_paragraph()
        note_run = note.add_run(
            'Note: Use the navigation pane in Word (View → Navigation Pane) '
            'to quickly jump between sections.'
        )
        note_run.font.size = Pt(10)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_page_break()

    def read_markdown_file(self, file_path: Path) -> str:
        """Read markdown file content

        Args:
            file_path: Path to markdown file

        Returns:
            File content as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.warning(f"Could not read {file_path.name}: {e}")
            return ""

    def convert_markdown_to_docx(self, content: str, file_name: str):
        """Convert markdown content to Word document format

        Args:
            content: Markdown content
            file_name: Name of source file (for reference)
        """
        if not content.strip():
            return

        # Add source file reference
        ref = self.doc.add_paragraph()
        ref_run = ref.add_run(f'Source: {file_name}')
        ref_run.font.size = Pt(9)
        ref_run.font.italic = True
        ref_run.font.color.rgb = RGBColor(150, 150, 150)

        lines = content.split('\n')
        in_code_block = False
        code_lines = []

        for line in lines:
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    if code_lines:
                        code_para = self.doc.add_paragraph('\n'.join(code_lines))
                        code_para.style = 'No Spacing'
                        code_run = code_para.runs[0]
                        code_run.font.name = 'Consolas'
                        code_run.font.size = Pt(9)
                        code_para.paragraph_format.left_indent = Inches(0.5)
                        code_para.paragraph_format.space_before = Pt(6)
                        code_para.paragraph_format.space_after = Pt(6)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if text:
                    self.doc.add_heading(text, level=min(level, 5))
                continue

            # Handle horizontal rules
            if line.strip() in ['---', '***', '___']:
                self.doc.add_paragraph('_' * 80)
                continue

            # Handle bullet lists
            if line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:]
                para = self.doc.add_paragraph(text, style='List Bullet')
                para.runs[0].font.size = Pt(11)
                continue

            # Handle numbered lists
            if re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                para = self.doc.add_paragraph(text, style='List Number')
                para.runs[0].font.size = Pt(11)
                continue

            # Handle bold/italic (basic markdown)
            if line.strip():
                para = self.doc.add_paragraph()

                # Simple bold **text**
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
                    run.font.size = Pt(11)

    def generate_document(self, output_path: str = 'Network_Segmentation_Analyzer_Complete_Documentation.docx'):
        """Generate complete documentation Word document

        Args:
            output_path: Output file path
        """
        logger.info("=" * 80)
        logger.info("GENERATING COMPREHENSIVE DOCUMENTATION")
        logger.info("=" * 80)

        # Add cover page
        self.add_cover_page()

        # Add table of contents
        self.add_table_of_contents()

        # Process each section
        total_files_processed = 0
        for section_name, file_list in self.SECTION_CATEGORIES.items():
            logger.info(f"\nProcessing section: {section_name}")

            # Add section heading
            self.doc.add_heading(section_name, level=1)
            self.doc.add_paragraph()

            # Add section description
            section_intro = self.doc.add_paragraph()
            section_intro_run = section_intro.add_run(
                f'This section contains {len(file_list)} documentation file(s) '
                f'related to {section_name.lower()}.'
            )
            section_intro_run.font.italic = True
            section_intro_run.font.size = Pt(10)
            section_intro_run.font.color.rgb = RGBColor(100, 100, 100)

            self.doc.add_paragraph()

            # Process each file in section
            for file_name in file_list:
                file_path = Path(file_name)

                if file_path.exists():
                    logger.info(f"  + {file_name}")

                    # Add subsection heading (file name without extension)
                    subsection_name = file_path.stem.replace('_', ' ').title()
                    self.doc.add_heading(subsection_name, level=2)

                    # Read and convert markdown content
                    content = self.read_markdown_file(file_path)
                    self.convert_markdown_to_docx(content, file_name)

                    total_files_processed += 1
                    self.doc.add_paragraph()
                else:
                    logger.warning(f"  - {file_name} (not found)")

            # Page break after each section
            self.doc.add_page_break()

        # Add appendix with uncategorized files
        logger.info("\nChecking for uncategorized files...")
        all_categorized = set()
        for file_list in self.SECTION_CATEGORIES.values():
            all_categorized.update(file_list)

        all_md_files = set(p.name for p in Path('.').glob('*.md'))
        uncategorized = all_md_files - all_categorized

        if uncategorized:
            logger.info(f"Found {len(uncategorized)} uncategorized files")
            self.doc.add_heading('Appendix: Additional Documentation', level=1)

            for file_name in sorted(uncategorized):
                logger.info(f"  + {file_name}")
                file_path = Path(file_name)

                subsection_name = file_path.stem.replace('_', ' ').title()
                self.doc.add_heading(subsection_name, level=2)

                content = self.read_markdown_file(file_path)
                self.convert_markdown_to_docx(content, file_name)

                total_files_processed += 1
                self.doc.add_paragraph()

        # Save document
        logger.info("\n" + "=" * 80)
        logger.info("SAVING DOCUMENT")
        logger.info("=" * 80)

        self.doc.save(output_path)

        logger.info(f"\n✓ Successfully generated: {output_path}")
        logger.info(f"✓ Total files processed: {total_files_processed}")
        logger.info(f"✓ Total sections: {len(self.SECTION_CATEGORIES)}")

        # Get file size
        output_file = Path(output_path)
        if output_file.exists():
            size_mb = output_file.stat().st_size / (1024 * 1024)
            logger.info(f"✓ Document size: {size_mb:.2f} MB")

        logger.info("\n" + "=" * 80)
        logger.info("DOCUMENTATION GENERATION COMPLETE")
        logger.info("=" * 80)


def main():
    """Main execution"""
    print("=" * 80)
    print("NETWORK SEGMENTATION ANALYZER - DOCUMENTATION CONSOLIDATION")
    print("=" * 80)
    print()
    print("This script will:")
    print("  1. Read all markdown documentation files (76+ files)")
    print("  2. Organize them into logical sections")
    print("  3. Generate a comprehensive Word document with:")
    print("     - Professional cover page")
    print("     - Table of contents")
    print("     - 11 main sections")
    print("     - Proper formatting and styling")
    print()
    print("=" * 80)
    print()

    # Create consolidator
    consolidator = DocumentationConsolidator()

    # Generate document
    output_path = 'Network_Segmentation_Analyzer_Complete_Documentation.docx'
    consolidator.generate_document(output_path)

    print()
    print("✅ SUCCESS!")
    print(f"📄 Open the document: {output_path}")
    print()
    print("Sections included:")
    for i, section_name in enumerate(consolidator.SECTION_CATEGORIES.keys(), 1):
        print(f"  {i}. {section_name}")
    print()
    print("Use the document for:")
    print("  - Comprehensive reference guide")
    print("  - Training new team members")
    print("  - Customer documentation")
    print("  - Audit and compliance")
    print()


if __name__ == '__main__':
    main()

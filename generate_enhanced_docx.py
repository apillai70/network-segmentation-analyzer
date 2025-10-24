"""
Enhanced DOCX Generator for Network Architecture and Threat Analysis
Generates professional Word documents with comprehensive security analysis
"""

import json
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level, color_rgb):
    """Add a colored heading"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb
    return heading

def add_threat_level_badge(paragraph, level):
    """Add a colored badge for threat level"""
    colors = {
        'HIGH': RGBColor(220, 53, 69),
        'MEDIUM': RGBColor(255, 193, 7),
        'LOW': RGBColor(40, 167, 69),
        'CRITICAL': RGBColor(139, 0, 0)
    }
    run = paragraph.add_run(f' [{level}] ')
    run.font.bold = True
    run.font.color.rgb = colors.get(level, RGBColor(128, 128, 128))

def generate_architecture_docx(app_id, flows_df, output_file):
    """Generate COMPREHENSIVE Architecture DOCX with server tiers and segmentation"""
    doc = Document()

    # ============================================================================
    # TITLE PAGE
    # ============================================================================
    title = doc.add_heading(f'{app_id}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(26, 82, 152)

    subtitle = doc.add_paragraph('Network Architecture & Segmentation Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(68, 68, 68)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n').font.size = Pt(11)
    meta.add_run(f'Total Network Flows: {len(flows_df)}\n').font.size = Pt(11)
    meta.add_run(f'Classification: Internal Use Only').font.size = Pt(10)

    doc.add_page_break()

    # ============================================================================
    # TABLE OF CONTENTS
    # ============================================================================
    add_colored_heading(doc, 'Table of Contents', 1, RGBColor(26, 82, 152))
    doc.add_paragraph('1. Executive Summary')
    doc.add_paragraph('2. Network Flow Statistics')
    doc.add_paragraph('3. Server Infrastructure by Tier')
    doc.add_paragraph('4. Application Dependencies')
    doc.add_paragraph('5. Network Segmentation Recommendations')
    doc.add_paragraph('6. Firewall Rules Requirements')
    doc.add_paragraph('7. VMware Infrastructure')

    doc.add_page_break()

    # ============================================================================
    # EXECUTIVE SUMMARY
    # ============================================================================
    add_colored_heading(doc, 'Executive Summary', 1, RGBColor(26, 82, 152))

    exec_para = doc.add_paragraph()
    exec_para.add_run(f'This document provides a comprehensive architecture analysis of the {app_id} application, ')
    exec_para.add_run('including its internal server structure, network dependencies, and security recommendations.')

    # Key Metrics Box
    doc.add_paragraph()
    metrics_table = doc.add_table(rows=1, cols=4)
    metrics_table.style = 'Light Grid Accent 1'

    cells = metrics_table.rows[0].cells
    cells[0].text = f'{len(flows_df)}\nTotal Flows'
    cells[1].text = f'{flows_df["Source IP"].nunique()}\nUnique Servers'
    cells[2].text = f'{flows_df["Port"].nunique()}\nUnique Ports'
    cells[3].text = f'{flows_df["Protocol"].nunique()}\nProtocols'

    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============================================================================
    # NETWORK FLOW STATISTICS
    # ============================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Network Flow Statistics', 1, RGBColor(26, 82, 152))

    internal = len(flows_df[flows_df['Flow Direction'] == 'internal'])
    outbound = len(flows_df[flows_df['Flow Direction'] == 'outbound'])
    external = len(flows_df[flows_df['Flow Direction'] == 'external'])

    stats_table = doc.add_table(rows=5, cols=3)
    stats_table.style = 'Medium Shading 1 Accent 1'

    # Header
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = 'Flow Type'
    hdr_cells[1].text = 'Count'
    hdr_cells[2].text = 'Percentage'

    total = len(flows_df)
    rows_data = [
        ('Internal Flows', internal, f'{(internal/total*100):.1f}%' if total > 0 else '0%'),
        ('Outbound to Apps', outbound, f'{(outbound/total*100):.1f}%' if total > 0 else '0%'),
        ('External/Internet', external, f'{(external/total*100):.1f}%' if total > 0 else '0%'),
        ('TOTAL', total, '100%')
    ]

    for idx, (flow_type, count, pct) in enumerate(rows_data, 1):
        row_cells = stats_table.rows[idx].cells
        row_cells[0].text = flow_type
        row_cells[1].text = str(count)
        row_cells[2].text = pct

    # ============================================================================
    # SERVER INFRASTRUCTURE BY TIER
    # ============================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Server Infrastructure by Tier', 1, RGBColor(26, 82, 152))

    doc.add_paragraph('This section details the server infrastructure organized by functional tier (Web, Application, Database, etc.).')

    # Classify servers by type
    from generate_complete_reports import CompleteReportGenerator
    generator = CompleteReportGenerator()

    server_types = {}
    for _, row in flows_df[['Source IP', 'Source Hostname', 'Port']].drop_duplicates().iterrows():
        ip = row['Source IP']
        hostname = row.get('Source Hostname', ip)
        ports = flows_df[flows_df['Source IP'] == ip]['Port'].unique().tolist()
        server_type = generator._detect_server_type(hostname, ports)

        if server_type not in server_types:
            server_types[server_type] = []
        server_types[server_type].append({'ip': ip, 'hostname': hostname})

    # Create server inventory table for each tier
    tier_colors = {
        'Web': RGBColor(76, 175, 80),
        'App': RGBColor(255, 152, 0),
        'Database': RGBColor(156, 39, 176),
        'Load Balancer': RGBColor(244, 67, 54)
    }

    for server_type in sorted(server_types.keys()):
        servers = server_types[server_type]

        doc.add_heading(f'{server_type} Tier ({len(servers)} servers)', 2)

        # Server table
        srv_table = doc.add_table(rows=len(servers)+1, cols=3)
        srv_table.style = 'Light List Accent 1'

        # Headers
        hdr = srv_table.rows[0].cells
        hdr[0].text = '#'
        hdr[1].text = 'IP Address'
        hdr[2].text = 'Hostname'

        for idx, server in enumerate(servers, 1):
            row = srv_table.rows[idx].cells
            row[0].text = str(idx)
            row[1].text = server['ip']
            row[2].text = server['hostname'] if server['hostname'] != server['ip'] else 'Unknown'

        doc.add_paragraph()

    # ============================================================================
    # APPLICATION DEPENDENCIES
    # ============================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Application Dependencies', 1, RGBColor(26, 82, 152))

    dest_apps = flows_df[flows_df['Dest App'].notna()]['Dest App'].unique()
    if len(dest_apps) > 0:
        doc.add_paragraph(f'{app_id} has dependencies on {len(dest_apps)} external applications:')

        dep_table = doc.add_table(rows=len(dest_apps)+1, cols=4)
        dep_table.style = 'Light Grid Accent 1'

        hdr = dep_table.rows[0].cells
        hdr[0].text = 'Application'
        hdr[1].text = 'Flows'
        hdr[2].text = 'Unique IPs'
        hdr[3].text = 'Protocols'

        for idx, dest_app in enumerate(sorted(dest_apps), 1):
            app_flows = flows_df[flows_df['Dest App'] == dest_app]
            row = dep_table.rows[idx].cells
            row[0].text = dest_app
            row[1].text = str(len(app_flows))
            row[2].text = str(app_flows['Dest IP'].nunique())
            row[3].text = ', '.join(app_flows['Protocol'].unique()[:3])
    else:
        doc.add_paragraph('[OK] No external application dependencies detected.')

    # ============================================================================
    # NETWORK SEGMENTATION RECOMMENDATIONS
    # ============================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Network Segmentation Recommendations', 1, RGBColor(26, 82, 152))

    doc.add_paragraph('Based on the network flow analysis, the following segmentation strategy is recommended:')

    doc.add_heading('Recommended Segmentation Zones', 2)

    zones = []
    if 'Web' in server_types:
        zones.append(('DMZ / Web Tier', 'Public-facing web servers', 'HIGH', server_types['Web']))
    if 'App' in server_types:
        zones.append(('Application Tier', 'Business logic servers', 'MEDIUM', server_types['App']))
    if 'Database' in server_types:
        zones.append(('Data Tier', 'Database servers', 'CRITICAL', server_types['Database']))

    for zone_name, description, criticality, servers in zones:
        doc.add_paragraph()
        zone_para = doc.add_paragraph()
        zone_para.add_run(f'{zone_name}').font.bold = True
        add_threat_level_badge(zone_para, criticality)
        doc.add_paragraph(f'Description: {description}')
        doc.add_paragraph(f'Servers: {len(servers)} nodes')

    # ============================================================================
    # SAVE DOCUMENT
    # ============================================================================
    doc.save(str(output_file))
    print(f'[OK] Generated: {output_file}')

# Main execution
if __name__ == '__main__':
    import sys
    from pathlib import Path

    # Test with ANSIBLE app
    app_dir = Path('persistent_data/applications/ANSIBLE')
    flows_csv = app_dir / 'flows.csv'

    if flows_csv.exists():
        flows_df = pd.read_csv(flows_csv)
        output_file = Path('outputs/docx/ANSIBLE_architecture_enhanced.docx')
        generate_architecture_docx('ANSIBLE', flows_df, output_file)
    else:
        print('Error: flows.csv not found')

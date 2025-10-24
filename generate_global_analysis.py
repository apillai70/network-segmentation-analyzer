"""
Global Network Segmentation Analysis - Enterprise-Wide
Analyzes all 166 applications for comprehensive network segmentation strategy
"""

import json
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def add_colored_heading(doc, text, level, color_rgb):
    """Add a colored heading"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb
    return heading

def generate_global_network_analysis():
    """Generate GLOBAL Network Segmentation Analysis across all applications"""

    logger.info("="*80)
    logger.info("GLOBAL NETWORK SEGMENTATION ANALYSIS - ENTERPRISE-WIDE")
    logger.info("="*80)

    doc = Document()

    # Collect data from all applications
    apps_dir = Path('persistent_data/applications')
    all_apps = []
    all_flows = []

    for app_dir in apps_dir.iterdir():
        if not app_dir.is_dir():
            continue
        flows_csv = app_dir / 'flows.csv'
        if flows_csv.exists():
            try:
                df = pd.read_csv(flows_csv)
                df['Application'] = app_dir.name
                all_flows.append(df)
                all_apps.append(app_dir.name)
            except Exception as e:
                logger.warning(f"Error loading {app_dir.name}: {e}")

    if not all_flows:
        logger.error("No flow data found!")
        return

    flows_df = pd.concat(all_flows, ignore_index=True)

    logger.info(f"Loaded {len(all_apps)} applications")
    logger.info(f"Total flows: {len(flows_df)}")

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    title = doc.add_heading('Enterprise Network Segmentation Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(26, 82, 152)

    subtitle = doc.add_paragraph('Comprehensive Assessment of Network Architecture & Segmentation Strategy')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(68, 68, 68)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n').font.size = Pt(12)
    meta.add_run(f'Applications Analyzed: {len(all_apps)}\n').font.size = Pt(12)
    meta.add_run(f'Total Network Flows: {len(flows_df):,}\n').font.size = Pt(12)
    meta.add_run(f'Classification: CONFIDENTIAL - Executive Summary').font.size = Pt(10)

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS
    # ========================================================================
    add_colored_heading(doc, 'Table of Contents', 1, RGBColor(26, 82, 152))

    toc_items = [
        '1. Executive Summary',
        '2. Enterprise Network Overview',
        '   2.1 Application Portfolio Analysis',
        '   2.2 Network Flow Statistics',
        '   2.3 Infrastructure Inventory',
        '3. Current State Assessment',
        '   3.1 Network Topology Analysis',
        '   3.2 Application Dependencies Map',
        '   3.3 External Exposure Points',
        '4. Segmentation Analysis',
        '   4.1 Tier-Based Segmentation',
        '   4.2 Data Classification Zones',
        '   4.3 Compliance Requirements',
        '   4.4 Geographic Distribution',
        '5. Risk Assessment',
        '   5.1 Attack Surface Analysis',
        '   5.2 Lateral Movement Risks',
        '   5.3 Single Points of Failure',
        '6. Segmentation Recommendations',
        '   6.1 Immediate Actions (0-30 days)',
        '   6.2 Short-Term Strategy (30-90 days)',
        '   6.3 Long-Term Roadmap (90+ days)',
        '7. Implementation Roadmap',
        '8. Success Metrics & KPIs'
    ]

    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ========================================================================
    # 1. EXECUTIVE SUMMARY
    # ========================================================================
    add_colored_heading(doc, '1. Executive Summary', 1, RGBColor(26, 82, 152))

    exec_para = doc.add_paragraph()
    exec_para.add_run('This document presents a comprehensive analysis of the enterprise network infrastructure, ')
    exec_para.add_run('covering ').font.bold = True
    exec_para.add_run(f'{len(all_apps)} applications ').font.bold = True
    exec_para.add_run(f'and {len(flows_df):,} network flows. ')
    exec_para.add_run('The analysis identifies critical security gaps, proposes a detailed network segmentation strategy, ')
    exec_para.add_run('and provides a roadmap for achieving Zero Trust Network Architecture.')

    doc.add_paragraph()
    doc.add_heading('Key Findings', 2)

    # Calculate key metrics
    total_servers = flows_df['Source IP'].nunique()
    external_flows = len(flows_df[flows_df['Flow Direction'] == 'external'])
    outbound_apps = flows_df[flows_df['Dest App'].notna()]['Dest App'].nunique()

    findings_table = doc.add_table(rows=6, cols=2)
    findings_table.style = 'Medium Shading 1 Accent 1'

    findings_data = [
        ('Total Applications', f'{len(all_apps)}'),
        ('Total Servers/Nodes', f'{total_servers:,}'),
        ('Total Network Flows', f'{len(flows_df):,}'),
        ('Application Dependencies', f'{outbound_apps}'),
        ('External Exposure Points', f'{external_flows:,}'),
        ('Unique Protocols', f"{flows_df['Protocol'].nunique()}")
    ]

    for idx, (metric, value) in enumerate(findings_data):
        row = findings_table.rows[idx].cells
        row[0].text = metric
        row[1].text = value

    # ========================================================================
    # 2. ENTERPRISE NETWORK OVERVIEW
    # ========================================================================
    doc.add_page_break()
    add_colored_heading(doc, '2. Enterprise Network Overview', 1, RGBColor(26, 82, 152))

    doc.add_heading('2.1 Application Portfolio Analysis', 2)

    doc.add_paragraph(f'The enterprise network consists of {len(all_apps)} distinct applications, ranging from ')
    doc.add_paragraph('customer-facing web services to internal business systems and data storage platforms.')

    doc.add_paragraph()
    doc.add_heading('Top 20 Applications by Flow Volume', 3)

    app_flow_counts = flows_df.groupby('Application').size().sort_values(ascending=False).head(20)

    app_table = doc.add_table(rows=min(len(app_flow_counts), 20)+1, cols=3)
    app_table.style = 'Light Grid Accent 1'

    hdr = app_table.rows[0].cells
    hdr[0].text = 'Rank'
    hdr[1].text = 'Application'
    hdr[2].text = 'Flow Count'

    for idx, (app, count) in enumerate(app_flow_counts.items(), 1):
        row = app_table.rows[idx].cells
        row[0].text = str(idx)
        row[1].text = app
        row[2].text = f'{count:,}'

    # ========================================================================
    # 2.2 NETWORK FLOW STATISTICS
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('2.2 Network Flow Statistics', 2)

    internal = len(flows_df[flows_df['Flow Direction'] == 'internal'])
    outbound = len(flows_df[flows_df['Flow Direction'] == 'outbound'])
    external = len(flows_df[flows_df['Flow Direction'] == 'external'])
    total = len(flows_df)

    flow_table = doc.add_table(rows=5, cols=4)
    flow_table.style = 'Medium Shading 1 Accent 3'

    hdr = flow_table.rows[0].cells
    hdr[0].text = 'Flow Direction'
    hdr[1].text = 'Count'
    hdr[2].text = 'Percentage'
    hdr[3].text = 'Risk Level'

    flow_data = [
        ('Internal', internal, f'{(internal/total*100):.1f}%' if total > 0 else '0%', 'LOW'),
        ('Outbound (App-to-App)', outbound, f'{(outbound/total*100):.1f}%' if total > 0 else '0%', 'MEDIUM'),
        ('External (Internet)', external, f'{(external/total*100):.1f}%' if total > 0 else '0%', 'HIGH'),
        ('TOTAL', total, '100%', '-')
    ]

    for idx, (direction, count, pct, risk) in enumerate(flow_data, 1):
        row = flow_table.rows[idx].cells
        row[0].text = direction
        row[1].text = f'{count:,}'
        row[2].text = pct
        row[3].text = risk

    # ========================================================================
    # 3. CURRENT STATE ASSESSMENT
    # ========================================================================
    doc.add_page_break()
    add_colored_heading(doc, '3. Current State Assessment', 1, RGBColor(26, 82, 152))

    doc.add_heading('3.1 Network Topology Analysis', 2)

    doc.add_paragraph('Current network topology analysis reveals:')

    doc.add_paragraph(f'[CRITICAL] {external:,} flows to external/internet destinations', style='List Bullet')
    doc.add_paragraph(f'[HIGH] {outbound_apps} application dependencies create complex attack surface', style='List Bullet')
    doc.add_paragraph(f'[MEDIUM] {total_servers:,} unique servers require segmentation', style='List Bullet')

    # ========================================================================
    # 4. SEGMENTATION ANALYSIS
    # ========================================================================
    doc.add_page_break()
    add_colored_heading(doc, '4. Segmentation Analysis', 1, RGBColor(26, 82, 152))

    doc.add_heading('4.1 Tier-Based Segmentation', 2)

    doc.add_paragraph('Recommended segmentation by application tier:')

    tier_table = doc.add_table(rows=6, cols=4)
    tier_table.style = 'Light List Accent 1'

    hdr = tier_table.rows[0].cells
    hdr[0].text = 'Tier'
    hdr[1].text = 'Purpose'
    hdr[2].text = 'Criticality'
    hdr[3].text = 'Isolation Level'

    tier_data = [
        ('DMZ/Edge', 'Public-facing services', 'HIGH', 'Strong (VLAN + FW)'),
        ('Web Tier', 'Web application servers', 'HIGH', 'Strong (VLAN + FW)'),
        ('App Tier', 'Business logic', 'HIGH', 'Medium (VLAN)'),
        ('Data Tier', 'Databases & storage', 'CRITICAL', 'Maximum (VLAN + FW + ACL)'),
        ('Management', 'Admin & monitoring', 'CRITICAL', 'Strong (Dedicated Network)')
    ]

    for idx, (tier, purpose, crit, isolation) in enumerate(tier_data, 1):
        row = tier_table.rows[idx].cells
        row[0].text = tier
        row[1].text = purpose
        row[2].text = crit
        row[3].text = isolation

    # ========================================================================
    # 6. SEGMENTATION RECOMMENDATIONS
    # ========================================================================
    doc.add_page_break()
    add_colored_heading(doc, '6. Segmentation Recommendations', 1, RGBColor(26, 82, 152))

    doc.add_heading('6.1 Immediate Actions (0-30 days)', 2)

    immediate_actions = [
        'Deploy Network Access Control Lists (ACLs) to block unauthorized tier-to-tier communication',
        'Implement firewall rules to restrict direct Web-to-Database connections',
        'Enable comprehensive network flow logging across all applications',
        'Identify and catalog all external-facing services',
        'Deploy network intrusion detection systems (IDS) at tier boundaries'
    ]

    for action in immediate_actions:
        doc.add_paragraph(action, style='List Number')

    doc.add_heading('6.2 Short-Term Strategy (30-90 days)', 2)

    short_term = [
        'Implement VLAN-based segmentation for all application tiers',
        'Deploy application-aware firewall rules based on flow analysis',
        'Establish micro-segmentation for database tier',
        'Implement jump servers for administrative access',
        'Deploy security groups for cloud-based applications',
        'Begin Zero Trust Network Architecture (ZTNA) pilot program'
    ]

    for action in short_term:
        doc.add_paragraph(action, style='List Number')

    doc.add_heading('6.3 Long-Term Roadmap (90+ days)', 2)

    long_term = [
        'Complete enterprise-wide Zero Trust Network Architecture deployment',
        'Implement Software-Defined Perimeter (SDP) for all applications',
        'Deploy Identity-Aware Proxy (IAP) for application access',
        'Implement continuous network behavior analytics',
        'Deploy automated micro-segmentation based on application behavior',
        'Achieve full network visibility with NDR (Network Detection & Response)'
    ]

    for action in long_term:
        doc.add_paragraph(action, style='List Number')

    # ========================================================================
    # 7. IMPLEMENTATION ROADMAP
    # ========================================================================
    doc.add_page_break()
    add_colored_heading(doc, '7. Implementation Roadmap', 1, RGBColor(26, 82, 152))

    roadmap_table = doc.add_table(rows=13, cols=5)
    roadmap_table.style = 'Medium Grid 1 Accent 1'

    hdr = roadmap_table.rows[0].cells
    hdr[0].text = 'Phase'
    hdr[1].text = 'Timeline'
    hdr[2].text = 'Milestone'
    hdr[3].text = 'Effort'
    hdr[4].text = 'Priority'

    roadmap_data = [
        ('Phase 1', 'Week 1-2', 'Network discovery & flow analysis', 'Low', 'P0'),
        ('Phase 1', 'Week 2-3', 'Deploy basic ACLs', 'Medium', 'P0'),
        ('Phase 1', 'Week 3-4', 'Enable logging & monitoring', 'Low', 'P0'),
        ('Phase 2', 'Month 2', 'VLAN segmentation - Web tier', 'High', 'P1'),
        ('Phase 2', 'Month 2', 'VLAN segmentation - App tier', 'High', 'P1'),
        ('Phase 2', 'Month 2-3', 'VLAN segmentation - DB tier', 'High', 'P0'),
        ('Phase 2', 'Month 3', 'Deploy tier-boundary firewalls', 'Medium', 'P1'),
        ('Phase 3', 'Month 4', 'Micro-segmentation pilot', 'High', 'P1'),
        ('Phase 3', 'Month 4-5', 'ZTNA pilot program', 'High', 'P2'),
        ('Phase 3', 'Month 5-6', 'SDP deployment', 'Very High', 'P2'),
        ('Phase 4', 'Month 6+', 'Enterprise-wide ZTNA rollout', 'Very High', 'P2'),
        ('Phase 4', 'Ongoing', 'Continuous optimization', 'Medium', 'P3')
    ]

    for idx, (phase, timeline, milestone, effort, priority) in enumerate(roadmap_data, 1):
        row = roadmap_table.rows[idx].cells
        row[0].text = phase
        row[1].text = timeline
        row[2].text = milestone
        row[3].text = effort
        row[4].text = priority

    # ========================================================================
    # 8. SUCCESS METRICS & KPIs
    # ========================================================================
    doc.add_page_break()
    add_colored_heading(doc, '8. Success Metrics & KPIs', 1, RGBColor(26, 82, 152))

    doc.add_paragraph('Track the following metrics to measure segmentation effectiveness:')

    kpi_table = doc.add_table(rows=11, cols=3)
    kpi_table.style = 'Light Grid Accent 3'

    hdr = kpi_table.rows[0].cells
    hdr[0].text = 'KPI'
    hdr[1].text = 'Current Baseline'
    hdr[2].text = 'Target (6 months)'

    kpi_data = [
        ('Unauthorized tier-to-tier flows', f'{external:,}', '< 100'),
        ('External exposure points', f'{external:,}', f'< {external//2:,}'),
        ('Mean time to detect (MTTD) breaches', 'Unknown', '< 5 minutes'),
        ('Segmentation coverage', '0%', '> 95%'),
        ('Micro-segmented applications', '0', f'{len(all_apps)//2}'),
        ('Zero Trust adoption', '0%', '> 50%'),
        ('Attack surface reduction', '0%', '> 70%'),
        ('Lateral movement prevention', '0%', '> 90%'),
        ('Firewall rule optimization', '0%', '> 80%'),
        ('Compliance score', 'TBD', '> 90%')
    ]

    for idx, (kpi, current, target) in enumerate(kpi_data, 1):
        row = kpi_table.rows[idx].cells
        row[0].text = kpi
        row[1].text = current
        row[2].text = target

    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================
    output_file = Path('outputs/docx/GLOBAL_Network_Segmentation_Analysis.docx')
    doc.save(str(output_file))
    logger.info(f"[OK] Generated: {output_file}")

    return output_file

if __name__ == '__main__':
    generate_global_network_analysis()

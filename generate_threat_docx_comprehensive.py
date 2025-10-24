"""
Comprehensive Threat Analysis & Segmentation Strategy Generator
Weaves threats into architecture and provides extensive micro-segmentation recommendations
"""

import json
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_threat_analysis():
    """Load threat surface analysis JSON"""
    threat_file = Path('outputs/threat_analysis/threat_surface_analysis.json')
    if threat_file.exists():
        with open(threat_file, 'r') as f:
            return json.load(f)
    return None

def add_colored_heading(doc, text, level, color_rgb):
    """Add a colored heading"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb
    return heading

def add_threat_badge(paragraph, level):
    """Add a colored badge for threat level"""
    colors = {
        'CRITICAL': RGBColor(139, 0, 0),
        'HIGH': RGBColor(220, 53, 69),
        'MEDIUM': RGBColor(255, 193, 7),
        'LOW': RGBColor(40, 167, 69)
    }
    run = paragraph.add_run(f' [{level}] ')
    run.font.bold = True
    run.font.color.rgb = colors.get(level, RGBColor(128, 128, 128))

def generate_threat_analysis_docx(app_id, flows_df, output_file):
    """Generate COMPREHENSIVE Threat Analysis & Segmentation Strategy Document"""

    doc = Document()

    # ===========================================================================
    # TITLE PAGE
    # ===========================================================================
    title = doc.add_heading(f'{app_id}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(139, 0, 0)

    subtitle = doc.add_paragraph('Threat Surface Analysis & Network Segmentation Strategy')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(68, 68, 68)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n').font.size = Pt(11)
    meta.add_run(f'Classification: CONFIDENTIAL - Security Analysis\n').font.size = Pt(10)
    meta.add_run(f'Distribution: Security Team, Architecture Team, Compliance').font.size = Pt(9)

    doc.add_page_break()

    # ===========================================================================
    # TABLE OF CONTENTS
    # ===========================================================================
    add_colored_heading(doc, 'Table of Contents', 1, RGBColor(139, 0, 0))
    doc.add_paragraph('1. Executive Threat Summary')
    doc.add_paragraph('2. Attack Surface Analysis')
    doc.add_paragraph('3. Attack Paths & Lateral Movement Risks')
    doc.add_paragraph('4. Tier-Based Threat Assessment')
    doc.add_paragraph('5. Comprehensive Segmentation Strategy')
    doc.add_paragraph('   5.1 Tier-Based Segmentation')
    doc.add_paragraph('   5.2 Data Privacy Zones')
    doc.add_paragraph('   5.3 External Customer Impact Zones')
    doc.add_paragraph('   5.4 Single Point of Failure Analysis')
    doc.add_paragraph('   5.5 Compliance-Based Segmentation')
    doc.add_paragraph('   5.6 Micro-Segmentation Recommendations')
    doc.add_paragraph('6. Firewall Rules & Access Control Lists')
    doc.add_paragraph('7. Remediation Recommendations')

    doc.add_page_break()

    # ===========================================================================
    # EXECUTIVE THREAT SUMMARY
    # ===========================================================================
    add_colored_heading(doc, 'Executive Threat Summary', 1, RGBColor(139, 0, 0))

    # Load threat analysis data
    threat_data = load_threat_analysis()

    if threat_data and 'summary' in threat_data:
        summary = threat_data['summary']

        doc.add_paragraph(f'This document presents a comprehensive security analysis of the {app_id} application, ')
        doc.add_paragraph('identifying attack surfaces, threat vectors, and providing detailed micro-segmentation strategies.')

        # Key Threat Metrics
        doc.add_paragraph()
        doc.add_heading('Critical Security Metrics', 2)

        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Medium Shading 1 Accent 2'

        data = [
            ('Total Attack Paths Identified', str(summary.get('total_attack_paths', 'N/A'))),
            ('Exposed Network Nodes', str(summary.get('exposed_nodes', 'N/A'))),
            ('Critical Assets at Risk', str(summary.get('critical_assets_at_risk', 'N/A'))),
            ('Average Attack Distance (hops)', f"{summary.get('avg_attack_distance', 0):.2f}"),
            ('High Threat Nodes', str(summary.get('high_threat_nodes', 'N/A')))
        ]

        for idx, (metric, value) in enumerate(data):
            row = metrics_table.rows[idx].cells
            row[0].text = metric
            row[1].text = value

    # ===========================================================================
    # ATTACK SURFACE ANALYSIS
    # ===========================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Attack Surface Analysis', 1, RGBColor(139, 0, 0))

    doc.add_heading('External Exposure', 2)
    external_flows = flows_df[flows_df['Flow Direction'] == 'external']

    if len(external_flows) > 0:
        exp_para = doc.add_paragraph()
        exp_para.add_run(f'CRITICAL FINDING: ')
        exp_para.runs[0].font.bold = True
        exp_para.runs[0].font.color.rgb = RGBColor(139, 0, 0)
        exp_para.add_run(f'{len(external_flows)} flows to external/internet destinations detected.')

        # External destinations table
        ext_table = doc.add_table(rows=min(len(external_flows), 20)+1, cols=4)
        ext_table.style = 'Light Grid Accent 2'

        hdr = ext_table.rows[0].cells
        hdr[0].text = 'Destination IP'
        hdr[1].text = 'Port'
        hdr[2].text = 'Protocol'
        hdr[3].text = 'Risk Level'

        for idx, (_, row) in enumerate(external_flows.head(20).iterrows(), 1):
            cells = ext_table.rows[idx].cells
            cells[0].text = str(row.get('Dest IP', 'Unknown'))
            cells[1].text = str(row.get('Port', 'Unknown'))
            cells[2].text = str(row.get('Protocol', 'Unknown'))

            # Determine risk based on port
            port = row.get('Port', 0)
            if port in [80, 443, 8080, 8443]:
                cells[3].text = 'MEDIUM'
            elif port in [22, 3389, 23, 21]:
                cells[3].text = 'HIGH'
            else:
                cells[3].text = 'LOW'
    else:
        doc.add_paragraph('[OK] No external internet exposure detected.')

    # ===========================================================================
    # ATTACK PATHS & LATERAL MOVEMENT
    # ===========================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Attack Paths & Lateral Movement Risks', 1, RGBColor(139, 0, 0))

    if threat_data and 'attack_paths' in threat_data:
        attack_paths = threat_data['attack_paths']

        doc.add_paragraph(f'Identified {len(attack_paths)} potential attack paths that could allow lateral movement from Web tier to sensitive Database tier.')

        doc.add_heading('High-Risk Attack Paths (Top 10)', 2)

        # Filter HIGH risk paths
        high_risk = [p for p in attack_paths if p.get('risk_level') == 'HIGH'][:10]

        if high_risk:
            for idx, path in enumerate(high_risk, 1):
                attack_para = doc.add_paragraph()
                attack_para.add_run(f'Attack Path #{idx}: ')
                attack_para.runs[0].font.bold = True
                add_threat_badge(attack_para, path.get('risk_level', 'UNKNOWN'))

                doc.add_paragraph(f"  Source: {path.get('source')} ({path.get('source_tier')})", style='List Bullet')
                doc.add_paragraph(f"  Target: {path.get('target')} ({path.get('target_tier')})", style='List Bullet')
                doc.add_paragraph(f"  Path Length: {path.get('path_length')} hops", style='List Bullet')
                doc.add_paragraph(f"  Attack Vector: {path.get('attack_vector', 'Unknown')}", style='List Bullet')
                doc.add_paragraph()

    # ===========================================================================
    # COMPREHENSIVE SEGMENTATION STRATEGY
    # ===========================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Comprehensive Segmentation Strategy', 1, RGBColor(26, 82, 152))

    doc.add_paragraph('This section provides extensive segmentation and micro-segmentation recommendations based on multiple security criteria.')

    # ===========================================================================
    # 5.1 TIER-BASED SEGMENTATION
    # ===========================================================================
    doc.add_heading('5.1 Tier-Based Segmentation', 2)

    tier_strategy = [
        {
            'zone': 'DMZ / Web Tier',
            'purpose': 'Public-facing web servers',
            'criticality': 'HIGH',
            'allowed_inbound': 'Internet (80/443), Load Balancer',
            'allowed_outbound': 'Application Tier only',
            'deny': 'Direct database access, SSH from internet'
        },
        {
            'zone': 'Application Tier',
            'purpose': 'Business logic and API servers',
            'criticality': 'HIGH',
            'allowed_inbound': 'Web Tier only',
            'allowed_outbound': 'Database Tier, External APIs',
            'deny': 'Internet access (except whitelisted APIs)'
        },
        {
            'zone': 'Database Tier',
            'criticality': 'CRITICAL',
            'purpose': 'Data persistence and storage',
            'allowed_inbound': 'Application Tier only',
            'allowed_outbound': 'None (data should not egress)',
            'deny': 'All other access, especially from Web Tier'
        }
    ]

    for zone_info in tier_strategy:
        zone_para = doc.add_paragraph()
        zone_para.add_run(f"{zone_info['zone']}").font.bold = True
        add_threat_badge(zone_para, zone_info['criticality'])

        zone_table = doc.add_table(rows=4, cols=2)
        zone_table.style = 'Light List Accent 1'

        zone_table.rows[0].cells[0].text = 'Allowed Inbound'
        zone_table.rows[0].cells[1].text = zone_info['allowed_inbound']

        zone_table.rows[1].cells[0].text = 'Allowed Outbound'
        zone_table.rows[1].cells[1].text = zone_info['allowed_outbound']

        zone_table.rows[2].cells[0].text = 'Explicit Denies'
        zone_table.rows[2].cells[1].text = zone_info['deny']

        zone_table.rows[3].cells[0].text = 'Purpose'
        zone_table.rows[3].cells[1].text = zone_info['purpose']

        doc.add_paragraph()

    # ===========================================================================
    # 5.2 DATA PRIVACY ZONES
    # ===========================================================================
    doc.add_page_break()
    doc.add_heading('5.2 Data Privacy Zones', 2)

    doc.add_paragraph('Segmentation based on data classification and privacy requirements:')

    privacy_zones = [
        ('PII Zone', 'Personally Identifiable Information', 'CRITICAL', 'GDPR, CCPA'),
        ('PHI Zone', 'Protected Health Information', 'CRITICAL', 'HIPAA'),
        ('PCI Zone', 'Payment Card Industry Data', 'CRITICAL', 'PCI-DSS'),
        ('Confidential Zone', 'Business confidential data', 'HIGH', 'Internal Policy'),
        ('Public Zone', 'Public-facing non-sensitive data', 'LOW', 'None')
    ]

    privacy_table = doc.add_table(rows=len(privacy_zones)+1, cols=4)
    privacy_table.style = 'Medium Shading 1 Accent 1'

    hdr = privacy_table.rows[0].cells
    hdr[0].text = 'Zone'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Criticality'
    hdr[3].text = 'Compliance'

    for idx, (zone, data_type, crit, compliance) in enumerate(privacy_zones, 1):
        row = privacy_table.rows[idx].cells
        row[0].text = zone
        row[1].text = data_type
        row[2].text = crit
        row[3].text = compliance

    # ===========================================================================
    # 5.3 EXTERNAL CUSTOMER IMPACT ZONES
    # ===========================================================================
    doc.add_page_break()
    doc.add_heading('5.3 External Customer Impact Zones', 2)

    doc.add_paragraph('Segmentation based on customer-facing impact and business criticality:')

    impact_zones = [
        ('Tier 0 - Customer Critical', 'Direct customer-facing services', 'CRITICAL', 'RTO: 15 min, RPO: 0'),
        ('Tier 1 - Customer Support', 'Customer support systems', 'HIGH', 'RTO: 1 hour, RPO: 15 min'),
        ('Tier 2 - Internal Operations', 'Internal business systems', 'MEDIUM', 'RTO: 4 hours, RPO: 1 hour'),
        ('Tier 3 - Non-Critical', 'Development/test systems', 'LOW', 'RTO: 24 hours, RPO: 4 hours')
    ]

    impact_table = doc.add_table(rows=len(impact_zones)+1, cols=4)
    impact_table.style = 'Medium Shading 1 Accent 3'

    hdr = impact_table.rows[0].cells
    hdr[0].text = 'Impact Tier'
    hdr[1].text = 'Description'
    hdr[2].text = 'Priority'
    hdr[3].text = 'Recovery Objectives'

    for idx, (tier, desc, priority, rto_rpo) in enumerate(impact_zones, 1):
        row = impact_table.rows[idx].cells
        row[0].text = tier
        row[1].text = desc
        row[2].text = priority
        row[3].text = rto_rpo

    # ===========================================================================
    # 5.4 SINGLE POINT OF FAILURE ANALYSIS
    # ===========================================================================
    doc.add_page_break()
    doc.add_heading('5.4 Single Point of Failure Analysis', 2)

    doc.add_paragraph('Identified single points of failure requiring redundancy and segmentation:')

    # Analyze servers by tier to identify SPOFs
    from generate_complete_reports import CompleteReportGenerator
    generator = CompleteReportGenerator()

    server_types = {}
    for _, row in flows_df[['Source IP', 'Source Hostname', 'Port']].drop_duplicates().iterrows():
        ip = row['Source IP']
        hostname = row.get('Source Hostname', ip)
        ports = flows_df[flows_df['Source IP'] == ip]['Port'].unique().tolist()
        server_type = generator._detect_server_type(hostname, ports)

        if server_type not in server_types:
            server_types[server_type] = []
        server_types[server_type].append({'ip': ip, 'hostname': hostname})

    spof_table = doc.add_table(rows=len(server_types)+1, cols=4)
    spof_table.style = 'Medium Shading 1 Accent 2'

    hdr = spof_table.rows[0].cells
    hdr[0].text = 'Tier'
    hdr[1].text = 'Server Count'
    hdr[2].text = 'SPOF Risk'
    hdr[3].text = 'Recommendation'

    for idx, (tier, servers) in enumerate(sorted(server_types.items()), 1):
        row = spof_table.rows[idx].cells
        row[0].text = tier
        row[1].text = str(len(servers))

        if len(servers) == 1:
            row[2].text = 'CRITICAL'
            row[3].text = 'Add redundancy immediately, implement HA'
        elif len(servers) == 2:
            row[2].text = 'HIGH'
            row[3].text = 'Add third node, configure load balancing'
        else:
            row[2].text = 'LOW'
            row[3].text = 'Verify load balancer configuration'

    # ===========================================================================
    # 5.5 COMPLIANCE-BASED SEGMENTATION
    # ===========================================================================
    doc.add_page_break()
    doc.add_heading('5.5 Compliance-Based Segmentation', 2)

    doc.add_paragraph('Segmentation strategies aligned with regulatory compliance requirements:')

    compliance_reqs = [
        ('PCI-DSS', 'Payment Card Data', 'Segment cardholder data environment (CDE) from rest of network, implement strong access controls'),
        ('HIPAA', 'Healthcare Data', 'Segregate PHI systems, implement audit logging, encrypt data in transit and at rest'),
        ('SOX', 'Financial Systems', 'Separate financial reporting systems, implement change controls and audit trails'),
        ('GDPR', 'EU Personal Data', 'Data residency controls, right to deletion capabilities, breach notification systems'),
        ('NIST 800-53', 'Federal Systems', 'Boundary protection, least privilege access, continuous monitoring')
    ]

    for standard, data_type, requirement in compliance_reqs:
        comp_para = doc.add_paragraph()
        comp_para.add_run(f'{standard} - {data_type}').font.bold = True
        doc.add_paragraph(f'Requirements: {requirement}', style='List Bullet')
        doc.add_paragraph()

    # ===========================================================================
    # 5.6 MICRO-SEGMENTATION RECOMMENDATIONS
    # ===========================================================================
    doc.add_page_break()
    doc.add_heading('5.6 Micro-Segmentation Recommendations', 2)

    doc.add_paragraph('Detailed micro-segmentation strategies for enhanced security:')

    micro_seg_strategies = [
        ('Application-Level Segmentation', 'Isolate each application into its own VLAN/subnet'),
        ('Process-Level Segmentation', 'Use containers/namespaces to isolate individual processes'),
        ('User-Based Segmentation', 'Segment based on user roles (admin, developer, end-user)'),
        ('Time-Based Segmentation', 'Restrict access during off-hours, maintenance windows'),
        ('Geo-Based Segmentation', 'Segment based on geographic location, data residency'),
        ('Device-Based Segmentation', 'Separate mobile, desktop, server, IoT devices'),
        ('Zero Trust Segmentation', 'Implement identity-based, least-privilege micro-perimeters')
    ]

    for strategy, description in micro_seg_strategies:
        doc.add_paragraph(f'{strategy}', style='List Number')
        doc.add_paragraph(f'   {description}')

    # ===========================================================================
    # REMEDIATION RECOMMENDATIONS
    # ===========================================================================
    doc.add_page_break()
    add_colored_heading(doc, 'Remediation Recommendations', 1, RGBColor(26, 82, 152))

    doc.add_heading('Priority 1 - Immediate Actions', 2)
    doc.add_paragraph('1. Block all direct Web-to-Database connections', style='List Number')
    doc.add_paragraph('2. Implement network-based intrusion detection (IDS/IPS)', style='List Number')
    doc.add_paragraph('3. Enable logging on all firewall rules and network flows', style='List Number')

    doc.add_heading('Priority 2 - Short Term (30 days)', 2)
    doc.add_paragraph('1. Deploy micro-segmentation for database tier', style='List Number')
    doc.add_paragraph('2. Implement application-aware firewall rules', style='List Number')
    doc.add_paragraph('3. Conduct penetration testing on attack paths', style='List Number')

    doc.add_heading('Priority 3 - Long Term (90 days)', 2)
    doc.add_paragraph('1. Migrate to Zero Trust Architecture', style='List Number')
    doc.add_paragraph('2. Implement Software-Defined Perimeter (SDP)', style='List Number')
    doc.add_paragraph('3. Deploy behavioral analytics and anomaly detection', style='List Number')

    # ===========================================================================
    # SAVE DOCUMENT
    # ===========================================================================
    doc.save(str(output_file))
    print(f'[OK] Generated comprehensive threat analysis: {output_file}')

# Main execution
if __name__ == '__main__':
    app_dir = Path('persistent_data/applications/ANSIBLE')
    flows_csv = app_dir / 'flows.csv'

    if flows_csv.exists():
        flows_df = pd.read_csv(flows_csv)
        output_file = Path('outputs/docx/ANSIBLE_threat_analysis_comprehensive.docx')
        generate_threat_analysis_docx('ANSIBLE', flows_df, output_file)
    else:
        print('Error: flows.csv not found')

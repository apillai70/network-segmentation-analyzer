"""
Comprehensive Solution Design Document Generator
=================================================
Generates detailed Solution Design documents for each application with:
- Application overview and metadata
- Architecture diagrams (embedded PNG)
- Mermaid diagram code
- Security zones and segmentation
- Data flows and dependencies
- Risk assessment and compliance
- Security recommendations

Author: Enterprise Architecture Team
Version: 1.0
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json

logger = logging.getLogger(__name__)


class ComprehensiveSolutionDocument:
    """Generates comprehensive solution design documents for applications"""

    def __init__(self, app_name: str, app_data: Optional[Dict] = None):
        """Initialize document for specific application

        Args:
            app_name: Application name/identifier
            app_data: Application topology and analysis data
        """
        self.app_name = app_name
        self.app_data = app_data or {}
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles

        # Customize heading styles
        try:
            h1_style = styles['Heading 1']
            h1_style.font.size = Pt(18)
            h1_style.font.color.rgb = RGBColor(0, 51, 102)
            h1_style.font.bold = True

            h2_style = styles['Heading 2']
            h2_style.font.size = Pt(14)
            h2_style.font.color.rgb = RGBColor(0, 102, 204)

            h3_style = styles['Heading 3']
            h3_style.font.size = Pt(12)
            h3_style.font.color.rgb = RGBColor(51, 102, 153)
        except:
            pass

    def generate_document(
        self,
        png_path: Optional[str],
        mermaid_path: Optional[str],
        output_path: str
    ):
        """Generate comprehensive solution design document

        Args:
            png_path: Path to PNG diagram file
            mermaid_path: Path to Mermaid diagram file
            output_path: Output path for Word document
        """
        logger.info(f"Generating comprehensive document for: {self.app_name}")

        # Cover page
        self._add_cover_page()
        self.doc.add_page_break()

        # Document control and revision history
        self._add_document_control()
        self.doc.add_page_break()

        # Table of Contents placeholder
        self._add_toc_placeholder()
        self.doc.add_page_break()

        # Executive Summary
        self._add_executive_summary()
        self.doc.add_page_break()

        # Application Overview
        self._add_application_overview()
        self.doc.add_page_break()

        # Architecture Design
        self._add_architecture_design(png_path, mermaid_path)
        self.doc.add_page_break()

        # Network Segmentation
        self._add_network_segmentation()
        self.doc.add_page_break()

        # Data Flows and Dependencies
        self._add_data_flows()
        self.doc.add_page_break()

        # Security Considerations
        self._add_security_considerations()
        self.doc.add_page_break()

        # Compliance and Risk
        self._add_compliance_risk()
        self.doc.add_page_break()

        # Recommendations
        self._add_recommendations()

        # Appendix
        self.doc.add_page_break()
        self._add_appendix(mermaid_path)

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output_file))
        logger.info(f"Success: Document saved: {output_path}")

    def _add_cover_page(self):
        """Add professional cover page"""
        # Main title
        title = self.doc.add_heading(
            f'Network Segmentation\nSolutions Architecture Document - {self.app_name}',
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        for _ in range(5):
            self.doc.add_paragraph()

        # Document metadata
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run('Version: 1.0\n').bold = True
        info_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
        info_para.add_run('Classification: Internal\n')
        info_para.add_run(f'Application: {self.app_name}\n')

        app_type = self.app_data.get('app_type', 'Unknown')
        zone = self.app_data.get('security_zone', 'Unknown')
        info_para.add_run(f'Type: {app_type.title()}\n')
        info_para.add_run(f'Security Zone: {zone}\n')

        for _ in range(3):
            self.doc.add_paragraph()

        # Footer
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Prepared by: Enterprise Architecture Team\n')
        footer.add_run('Enterprise Architecture & Network Security\n')
        footer.add_run('Auto-generated by Network Segmentation Analyzer v3.0\n')

    def _add_document_control(self):
        """Add document control and revision history"""
        self.doc.add_heading('Document Control', level=1)

        # Document information table
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'

        rows_data = [
            ('Document Title', f'Network Segmentation Solution Design - {self.app_name}'),
            ('Document ID', f'NSSD-{self.app_name}-001'),
            ('Version', '1.0'),
            ('Date', datetime.now().strftime("%Y-%m-%d")),
            ('Author', 'Network Segmentation Analyzer'),
            ('Status', 'Draft'),
            ('Classification', 'Internal')
        ]

        for idx, (label, value) in enumerate(rows_data):
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[idx].cells[1].text = value

        self.doc.add_paragraph()

        # Revision history
        self.doc.add_heading('Revision History', level=2)

        rev_table = self.doc.add_table(rows=2, cols=4)
        rev_table.style = 'Light Grid Accent 1'

        # Header
        headers = ['Version', 'Date', 'Author', 'Description']
        for idx, header in enumerate(headers):
            rev_table.rows[0].cells[idx].text = header
            rev_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

        # First entry
        rev_table.rows[1].cells[0].text = '1.0'
        rev_table.rows[1].cells[1].text = datetime.now().strftime("%Y-%m-%d")
        rev_table.rows[1].cells[2].text = 'Auto-generated'
        rev_table.rows[1].cells[3].text = 'Initial solution design document'

    def _add_toc_placeholder(self):
        """Add table of contents placeholder"""
        self.doc.add_heading('Table of Contents', level=1)

        toc_para = self.doc.add_paragraph()
        toc_para.add_run('[Insert Table of Contents here]\n\n')
        toc_para.add_run('In Microsoft Word:\n')
        toc_para.add_run('1. Click here\n')
        toc_para.add_run('2. Go to References > Table of Contents\n')
        toc_para.add_run('3. Select "Automatic Table 1" or "Automatic Table 2"\n')

    def _add_executive_summary(self):
        """Add executive summary"""
        self.doc.add_heading('Executive Summary', level=1)

        app_type = self.app_data.get('app_type', 'application service').replace('_', ' ')
        zone = self.app_data.get('security_zone', 'APP_TIER')
        confidence = self.app_data.get('confidence', 0.0) * 100
        risk = self.app_data.get('risk_level', 'MEDIUM')

        summary_text = f"""
This document presents the comprehensive solution design for the {self.app_name} application,
focusing on network segmentation and security architecture. The application has been classified
as a {app_type} and is designated for deployment in the {zone}.

Through automated analysis and topology discovery, this application has been analyzed with
{confidence:.0f}% confidence. The solution design ensures proper network segmentation, security
controls, and compliance with enterprise security policies.

Key Highlights:
• Application Type: {app_type.title()}
• Security Zone: {zone}
• Risk Level: {risk}
• Analysis Confidence: {confidence:.0f}%
• Security Posture: Aligned with enterprise security framework

This document provides detailed architecture diagrams, data flow analysis, security zone
assignments, and recommendations for secure deployment and operation.
"""

        self.doc.add_paragraph(summary_text.strip())

        # Key metrics
        self.doc.add_paragraph()
        self.doc.add_heading('Key Metrics', level=2)

        dependencies = self.app_data.get('predicted_dependencies', [])
        num_deps = len([d for d in dependencies if d.get('source') != 'type_inference'])

        metrics_table = self.doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Light Grid Accent 1'

        metrics = [
            ('Total Dependencies', str(num_deps)),
            ('Compliance Requirements', ', '.join(self.app_data.get('compliance_requirements', ['Standard']))),
            ('Risk Rating', risk),
            ('Confidence Score', f"{confidence:.0f}%"),
            ('Analysis Method', self.app_data.get('analysis_method', 'Automated'))
        ]

        for idx, (metric, value) in enumerate(metrics):
            metrics_table.rows[idx].cells[0].text = metric
            metrics_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            metrics_table.rows[idx].cells[1].text = value

    def _add_application_overview(self):
        """Add detailed application overview"""
        self.doc.add_heading('Application Overview', level=1)

        # Application description
        self.doc.add_heading('Description', level=2)

        primary_function = self.app_data.get('primary_function', 'Application service for enterprise operations')
        self.doc.add_paragraph(primary_function)

        # Application characteristics
        self.doc.add_heading('Characteristics', level=2)

        app_type = self.app_data.get('app_type', 'unknown').replace('_', ' ')
        zone = self.app_data.get('security_zone', 'APP_TIER')

        char_text = f"""
The {self.app_name} application exhibits the following characteristics:

• Type: {app_type.title()}
• Primary Zone: {zone}
• Architecture Pattern: Multi-tier distributed system
• Communication Patterns: Service-to-service, API-based
"""

        self.doc.add_paragraph(char_text.strip())

        # Additional characteristics
        characteristics = self.app_data.get('characteristics', [])
        if characteristics:
            self.doc.add_paragraph("\nAdditional Characteristics:")
            for char in characteristics:
                self.doc.add_paragraph(f"• {char}", style='List Bullet')

        # Technology stack
        self.doc.add_heading('Technology Stack', level=2)

        tech_stack = self.app_data.get('tech_stack', {})

        tech_table = self.doc.add_table(rows=4, cols=2)
        tech_table.style = 'Light Grid Accent 1'

        tech_items = [
            ('Language', tech_stack.get('language', 'To be determined')),
            ('Framework', tech_stack.get('framework', 'To be determined')),
            ('Database', tech_stack.get('database', 'To be determined')),
            ('Deployment', tech_stack.get('deployment', 'Enterprise container platform'))
        ]

        for idx, (tech, value) in enumerate(tech_items):
            tech_table.rows[idx].cells[0].text = tech
            tech_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            tech_table.rows[idx].cells[1].text = value

    def _add_architecture_design(self, png_path: Optional[str], mermaid_path: Optional[str]):
        """Add architecture design with diagrams"""
        self.doc.add_heading('Architecture Design', level=1)

        # Architecture overview
        self.doc.add_heading('Architecture Overview', level=2)

        arch_text = f"""
The {self.app_name} application follows a modern multi-tier architecture pattern designed for
scalability, security, and maintainability. The architecture is organized into distinct tiers,
each with specific responsibilities and security boundaries.

The following diagram illustrates the complete application architecture including internal tiers,
downstream dependencies, and infrastructure components.
"""

        self.doc.add_paragraph(arch_text.strip())

        self.doc.add_paragraph()

        # Embed PNG diagram
        self.doc.add_heading('Application Architecture Diagram', level=2)

        if png_path and Path(png_path).exists():
            try:
                self.doc.add_picture(str(png_path), width=Inches(6.5))
                logger.info(f"  PNG embedded: {Path(png_path).name}")
            except Exception as e:
                logger.error(f"  Failed to embed PNG: {e}")
                self.doc.add_paragraph(f'[Diagram image not found: {png_path}]')
        else:
            self.doc.add_paragraph(f'[Diagram image not found: {png_path}]')

        self.doc.add_paragraph()

        # Diagram legend
        self.doc.add_heading('Diagram Legend', level=3)

        legend_items = [
            'Application Container (subgraph) = Internal multi-tier architecture',
            'Web Tier = Frontend/presentation layer handling user requests',
            'App Tier = Business logic and application processing',
            'Data Tier = Persistent data storage and database services',
            'Cache Tier = High-speed caching layer (Redis, Memcache)',
            'Messaging Tier = Asynchronous messaging and event processing',
            'Management Tier = Monitoring, logging, and operational tools',
            'Downstream Applications (Circles) = External applications/services called by this app',
            'Infrastructure (Rectangles) = Databases, caches, queues, and other infrastructure',
            'Thick lines (===) = Application-to-application communication',
            'Regular lines (---) = Infrastructure dependencies',
            'Colors = Security zones and tier classifications'
        ]

        for item in legend_items:
            self.doc.add_paragraph(item, style='List Bullet')

        # Tier descriptions
        self.doc.add_heading('Tier Descriptions', level=2)

        tier_descriptions = {
            'Web Tier': 'Handles incoming HTTP/HTTPS requests, load balancing, SSL termination, and serves as the entry point for user traffic.',
            'App Tier': 'Contains core business logic, API endpoints, service orchestration, and application processing.',
            'Data Tier': 'Manages persistent data storage, database operations, and data consistency.',
            'Cache Tier': 'Provides high-speed data caching using Redis, Memcache, or similar technologies to reduce database load.',
            'Messaging Tier': 'Handles asynchronous messaging, event processing, and inter-service communication via Kafka, RabbitMQ, or similar.',
            'Management Tier': 'Provides monitoring, logging, metrics collection, health checks, and operational management.'
        }

        for tier, description in tier_descriptions.items():
            self.doc.add_heading(tier, level=3)
            self.doc.add_paragraph(description)

    def _add_network_segmentation(self):
        """Add network segmentation details"""
        self.doc.add_heading('Network Segmentation', level=1)

        zone = self.app_data.get('security_zone', 'APP_TIER')

        # Security zone assignment
        self.doc.add_heading('Security Zone Assignment', level=2)

        zone_text = f"""
Based on the application analysis, {self.app_name} has been assigned to the {zone}.
This zone provides appropriate security controls and network isolation for this application type.
"""

        self.doc.add_paragraph(zone_text.strip())

        self.doc.add_paragraph()

        # Zone characteristics
        zone_descriptions = {
            'WEB_TIER': {
                'description': 'Public-facing tier handling external user requests',
                'security_level': 'High',
                'access': 'Internet-facing with strict ingress controls',
                'controls': 'WAF, DDoS protection, rate limiting, TLS termination'
            },
            'APP_TIER': {
                'description': 'Application logic tier for business processing',
                'security_level': 'High',
                'access': 'Accessible from Web Tier only',
                'controls': 'Service mesh, authentication, authorization, encryption'
            },
            'DATA_TIER': {
                'description': 'Data persistence and database tier',
                'security_level': 'Critical',
                'access': 'Accessible from App Tier only',
                'controls': 'Database firewall, encryption at rest, audit logging'
            },
            'CACHE_TIER': {
                'description': 'High-speed caching layer',
                'security_level': 'Medium',
                'access': 'Accessible from App Tier',
                'controls': 'Authentication, encryption in transit'
            },
            'MESSAGING_TIER': {
                'description': 'Asynchronous messaging and event processing',
                'security_level': 'Medium',
                'access': 'Accessible from App Tier and authorized services',
                'controls': 'Authentication, authorization, message encryption'
            },
            'MANAGEMENT_TIER': {
                'description': 'Monitoring, logging, and operations',
                'security_level': 'High',
                'access': 'Restricted to operations team and management tools',
                'controls': 'Strong authentication, audit logging, access controls'
            }
        }

        if zone in zone_descriptions:
            zone_info = zone_descriptions[zone]

            self.doc.add_heading('Zone Characteristics', level=3)

            zone_table = self.doc.add_table(rows=4, cols=2)
            zone_table.style = 'Light Grid Accent 1'

            zone_data = [
                ('Description', zone_info['description']),
                ('Security Level', zone_info['security_level']),
                ('Access Pattern', zone_info['access']),
                ('Security Controls', zone_info['controls'])
            ]

            for idx, (label, value) in enumerate(zone_data):
                zone_table.rows[idx].cells[0].text = label
                zone_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
                zone_table.rows[idx].cells[1].text = value

        # Segmentation rules
        self.doc.add_paragraph()
        self.doc.add_heading('Segmentation Rules', level=2)

        rules_text = """
The following network segmentation rules should be implemented:

1. Micro-segmentation: Implement granular firewall rules between tiers
2. Zero Trust: Verify all connections regardless of source network
3. Least Privilege: Allow only necessary traffic between tiers
4. Encryption: All inter-tier traffic must be encrypted (TLS 1.2+)
5. Monitoring: Log and monitor all cross-tier communications
6. Lateral Movement Prevention: Block unauthorized east-west traffic
"""

        self.doc.add_paragraph(rules_text.strip())

    def _add_data_flows(self):
        """Add data flows and dependencies"""
        self.doc.add_heading('Data Flows and Dependencies', level=1)

        # Dependencies overview
        self.doc.add_heading('Dependency Overview', level=2)

        dependencies = self.app_data.get('predicted_dependencies', [])
        observed_deps = [d for d in dependencies if d.get('source') == 'network_observation']
        inferred_deps = [d for d in dependencies if d.get('source') != 'network_observation']

        overview_text = f"""
The {self.app_name} application has {len(dependencies)} identified dependencies:
• {len(observed_deps)} observed from network traffic analysis
• {len(inferred_deps)} inferred from application type and patterns

These dependencies represent both infrastructure components (databases, caches, queues) and
downstream applications that {self.app_name} communicates with.
"""

        self.doc.add_paragraph(overview_text.strip())

        self.doc.add_paragraph()

        # Observed dependencies
        if observed_deps:
            self.doc.add_heading('Observed Dependencies', level=2)
            self.doc.add_paragraph('The following dependencies were observed in network traffic:')
            self.doc.add_paragraph()

            obs_table = self.doc.add_table(rows=len(observed_deps)+1, cols=4)
            obs_table.style = 'Light Grid Accent 1'

            # Header
            headers = ['Target', 'Type', 'Purpose', 'Confidence']
            for idx, header in enumerate(headers):
                obs_table.rows[0].cells[idx].text = header
                obs_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

            # Data rows
            for idx, dep in enumerate(observed_deps, 1):
                obs_table.rows[idx].cells[0].text = dep.get('name', 'Unknown*')
                obs_table.rows[idx].cells[1].text = dep.get('type', 'unknown')
                obs_table.rows[idx].cells[2].text = dep.get('purpose', '')
                obs_table.rows[idx].cells[3].text = f"{dep.get('confidence', 0)*100:.0f}%"

        # Inferred dependencies
        if inferred_deps:
            self.doc.add_paragraph()
            self.doc.add_heading('Inferred Dependencies', level=2)
            self.doc.add_paragraph('The following dependencies are inferred from application type:')
            self.doc.add_paragraph()

            inf_table = self.doc.add_table(rows=len(inferred_deps)+1, cols=4)
            inf_table.style = 'Light Grid Accent 1'

            # Header
            for idx, header in enumerate(['Type', 'Service', 'Purpose', 'Confidence']):
                inf_table.rows[0].cells[idx].text = header
                inf_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

            # Data rows
            for idx, dep in enumerate(inferred_deps, 1):
                inf_table.rows[idx].cells[0].text = dep.get('type', 'unknown')
                inf_table.rows[idx].cells[1].text = dep.get('name', 'Unknown*')
                inf_table.rows[idx].cells[2].text = dep.get('purpose', '')
                inf_table.rows[idx].cells[3].text = f"{dep.get('confidence', 0)*100:.0f}%"

        # Data flow patterns
        self.doc.add_paragraph()
        self.doc.add_heading('Data Flow Patterns', level=2)

        flow_text = """
Typical data flow patterns for this application:

1. Ingress Traffic: Requests enter through the Web Tier
2. Processing: App Tier processes business logic
3. Data Access: App Tier queries databases and caches
4. Async Operations: Messaging tier handles events
5. Monitoring: Management tier collects metrics and logs
6. Egress Traffic: Calls to downstream applications
"""

        self.doc.add_paragraph(flow_text.strip())

    def _add_security_considerations(self):
        """Add security considerations"""
        self.doc.add_heading('Security Considerations', level=1)

        risk_level = self.app_data.get('risk_level', 'MEDIUM')

        # Security posture
        self.doc.add_heading('Security Posture', level=2)

        security_text = f"""
The {self.app_name} application has been assessed with a {risk_level} risk rating.
The following security controls should be implemented to maintain a strong security posture.
"""

        self.doc.add_paragraph(security_text.strip())

        # Security controls
        self.doc.add_heading('Required Security Controls', level=2)

        controls = [
            'Authentication and Authorization: Implement strong authentication (OAuth 2.0, SAML) and fine-grained authorization',
            'Encryption in Transit: All communication must use TLS 1.2 or higher',
            'Encryption at Rest: Sensitive data must be encrypted at rest using AES-256',
            'Network Segmentation: Implement micro-segmentation with strict firewall rules',
            'Access Control: Apply principle of least privilege for all service accounts',
            'Monitoring and Logging: Comprehensive logging of all security-relevant events',
            'Vulnerability Management: Regular security scans and patch management',
            'Incident Response: Define incident response procedures and playbooks',
            'Data Protection: Implement DLP controls for sensitive data',
            'API Security: API gateway with rate limiting, throttling, and validation'
        ]

        for control in controls:
            self.doc.add_paragraph(control, style='List Bullet')

        # Security architecture
        self.doc.add_heading('Security Architecture Principles', level=2)

        principles = [
            'Defense in Depth: Multiple layers of security controls',
            'Zero Trust: Never trust, always verify',
            'Least Privilege: Minimum access required for operation',
            'Secure by Default: Security controls enabled from deployment',
            'Privacy by Design: Privacy considerations in all design decisions',
            'Continuous Monitoring: Real-time security monitoring and alerting'
        ]

        for principle in principles:
            self.doc.add_paragraph(principle, style='List Bullet')

    def _add_compliance_risk(self):
        """Add compliance and risk assessment"""
        self.doc.add_heading('Compliance and Risk Assessment', level=1)

        # Compliance requirements
        self.doc.add_heading('Compliance Requirements', level=2)

        compliance_reqs = self.app_data.get('compliance_requirements', ['SOC2'])

        compliance_text = f"""
The {self.app_name} application must comply with the following regulatory and compliance frameworks:

{', '.join(compliance_reqs)}

The solution design ensures compliance through appropriate security controls, audit logging,
and data protection measures.
"""

        self.doc.add_paragraph(compliance_text.strip())

        # Risk assessment
        self.doc.add_heading('Risk Assessment', level=2)

        risk_level = self.app_data.get('risk_level', 'MEDIUM')

        risk_table = self.doc.add_table(rows=4, cols=2)
        risk_table.style = 'Light Grid Accent 1'

        risk_data = [
            ('Overall Risk Rating', risk_level),
            ('Confidentiality Risk', risk_level),
            ('Integrity Risk', risk_level),
            ('Availability Risk', risk_level)
        ]

        for idx, (risk_type, rating) in enumerate(risk_data):
            risk_table.rows[idx].cells[0].text = risk_type
            risk_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            risk_table.rows[idx].cells[1].text = rating

        # Risk mitigation
        self.doc.add_paragraph()
        self.doc.add_heading('Risk Mitigation Strategies', level=2)

        mitigations = [
            'Network Segmentation: Isolate application tiers to limit blast radius',
            'Access Controls: Implement strong authentication and authorization',
            'Encryption: Protect data in transit and at rest',
            'Monitoring: Continuous security monitoring and threat detection',
            'Incident Response: Defined procedures for security incidents',
            'Backup and Recovery: Regular backups with tested recovery procedures',
            'Vulnerability Management: Regular security assessments and patching',
            'Security Training: Regular security awareness training for team'
        ]

        for mitigation in mitigations:
            self.doc.add_paragraph(mitigation, style='List Bullet')

    def _add_recommendations(self):
        """Add recommendations"""
        self.doc.add_heading('Recommendations', level=1)

        # Implementation recommendations
        self.doc.add_heading('Implementation Recommendations', level=2)

        impl_recs = [
            'Deploy using infrastructure-as-code (Terraform, CloudFormation) for consistency',
            'Implement CI/CD pipeline with security gates and automated testing',
            'Use container orchestration (Kubernetes) for scalability and resilience',
            'Deploy service mesh (Istio, Linkerd) for advanced traffic management',
            'Implement centralized logging and monitoring (ELK, Splunk, Datadog)',
            'Use secrets management solution (Vault, AWS Secrets Manager)',
            'Configure auto-scaling based on traffic patterns and resource utilization',
            'Implement blue-green or canary deployment strategies',
            'Set up disaster recovery with defined RTO and RPO'
        ]

        for rec in impl_recs:
            self.doc.add_paragraph(rec, style='List Bullet')

        # Security hardening
        self.doc.add_heading('Security Hardening Recommendations', level=2)

        security_recs = [
            'Enable multi-factor authentication for all administrative access',
            'Implement network policies to restrict pod-to-pod communication',
            'Use security scanning in CI/CD pipeline (SAST, DAST, SCA)',
            'Configure security information and event management (SIEM)',
            'Implement runtime application self-protection (RASP)',
            'Use Web Application Firewall (WAF) for web-facing components',
            'Configure DDoS protection for internet-facing services',
            'Implement API rate limiting and throttling',
            'Enable audit logging for all security-relevant operations',
            'Perform regular penetration testing and security assessments'
        ]

        for rec in security_recs:
            self.doc.add_paragraph(rec, style='List Bullet')

        # Operational recommendations
        self.doc.add_heading('Operational Recommendations', level=2)

        ops_recs = [
            'Establish SLAs for availability, performance, and response time',
            'Define SLIs and SLOs for service level objectives',
            'Implement comprehensive monitoring with alerting and escalation',
            'Create runbooks for common operational procedures',
            'Establish change management process with approval workflows',
            'Perform regular disaster recovery drills',
            'Implement capacity planning based on growth projections',
            'Establish on-call rotation and incident response procedures'
        ]

        for rec in ops_recs:
            self.doc.add_paragraph(rec, style='List Bullet')

    def _add_appendix(self, mermaid_path: Optional[str]):
        """Add appendix with technical details"""
        self.doc.add_heading('Appendix', level=1)

        # Mermaid diagram code
        self.doc.add_heading('A. Mermaid Diagram Code', level=2)

        self.doc.add_paragraph(
            'The following Mermaid code can be used to regenerate or modify the architecture diagram:'
        )

        self.doc.add_paragraph()

        if mermaid_path and Path(mermaid_path).exists():
            try:
                with open(mermaid_path, 'r', encoding='utf-8') as f:
                    mermaid_code = f.read()

                # Add code with monospace font
                code_para = self.doc.add_paragraph()
                run = code_para.add_run(mermaid_code)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)

                logger.info(f"  Mermaid code embedded from: {Path(mermaid_path).name}")
            except Exception as e:
                logger.error(f"  Failed to read Mermaid file: {e}")
                self.doc.add_paragraph(f'[Mermaid code not found: {mermaid_path}]')
        else:
            self.doc.add_paragraph(f'[Mermaid code not found: {mermaid_path}]')

        # Glossary
        self.doc.add_paragraph()
        self.doc.add_heading('B. Glossary', level=2)

        glossary = {
            'API': 'Application Programming Interface',
            'App Tier': 'Application tier containing business logic',
            'CI/CD': 'Continuous Integration/Continuous Deployment',
            'Data Tier': 'Database and data persistence tier',
            'DDoS': 'Distributed Denial of Service',
            'Micro-segmentation': 'Granular network segmentation strategy',
            'SIEM': 'Security Information and Event Management',
            'TLS': 'Transport Layer Security',
            'WAF': 'Web Application Firewall',
            'Web Tier': 'Frontend/presentation tier',
            'Zero Trust': 'Security model requiring verification of all connections'
        }

        for term, definition in sorted(glossary.items()):
            para = self.doc.add_paragraph()
            para.add_run(f'{term}: ').bold = True
            para.add_run(definition)

        # References
        self.doc.add_paragraph()
        self.doc.add_heading('C. References', level=2)

        references = [
            'Enterprise Security Architecture Framework v2.0',
            'Network Segmentation Best Practices',
            'Zero Trust Architecture Guidelines',
            'NIST Cybersecurity Framework',
            'Cloud Security Alliance Controls',
            'Enterprise Compliance Requirements'
        ]

        for ref in references:
            self.doc.add_paragraph(ref, style='List Bullet')


def generate_comprehensive_solution_document(
    app_name: str,
    app_data: Dict,
    png_path: str,
    mermaid_path: str,
    output_path: str
):
    """Generate comprehensive solution design document

    Args:
        app_name: Application name
        app_data: Application topology and analysis data
        png_path: Path to PNG diagram
        mermaid_path: Path to Mermaid diagram
        output_path: Output path for Word document

    Returns:
        Path to generated document
    """
    doc_gen = ComprehensiveSolutionDocument(app_name, app_data)
    doc_gen.generate_document(png_path, mermaid_path, output_path)

    logger.info(f"Success: Comprehensive solution document generated: {output_path}")
    return output_path

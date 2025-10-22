#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Diagram Format Generator
========================
Generates diagrams in multiple formats: PNG, SVG, and DOCX from Mermaid content.

Features:
- PNG generation via Mermaid.ink API (4800px width for high resolution)
- SVG generation via Mermaid.ink API (infinite zoom capability)
- DOCX generation with embedded diagrams
- Automatic fallback to local mmdc if API fails
- Retry logic for API errors

Author: Enterprise Security Team
Version: 1.0
"""

import logging
import base64
import urllib.request
import urllib.parse
import time
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class DiagramFormatGenerator:
    """Generates diagrams in PNG, SVG, and DOCX formats"""

    def __init__(self, max_retries: int = 3, api_timeout: int = 30):
        """Initialize diagram format generator

        Args:
            max_retries: Maximum number of retries for API calls
            api_timeout: Timeout for API requests in seconds
        """
        self.max_retries = max_retries
        self.api_timeout = api_timeout
        logger.info(f"DiagramFormatGenerator initialized (max_retries={max_retries}, timeout={api_timeout}s)")

    def generate_png(self, mermaid_content: str, output_path: Path, width: int = 4800) -> bool:
        """Generate PNG diagram from Mermaid content

        Args:
            mermaid_content: Mermaid diagram content (without markdown fences)
            output_path: Output PNG file path
            width: PNG width in pixels (default: 4800 for high resolution)

        Returns:
            True if successful, False otherwise
        """
        logger.info(f"Generating PNG: {output_path.name}")

        # Strip markdown fences if present
        content = self._strip_markdown_fences(mermaid_content)

        # Try API first
        if self._generate_via_api(content, output_path, 'png', width):
            logger.info(f"  [OK] PNG generated via API: {output_path.name}")
            return True

        # Fallback to local mmdc
        logger.warning("  API failed, trying local mmdc...")
        if self._generate_via_mmdc(content, output_path, 'png', width):
            logger.info(f"  [OK] PNG generated via mmdc: {output_path.name}")
            return True

        logger.error(f"  [FAILED] Could not generate PNG: {output_path.name}")
        return False

    def generate_svg(self, mermaid_content: str, output_path: Path) -> bool:
        """Generate SVG diagram from Mermaid content

        Args:
            mermaid_content: Mermaid diagram content (without markdown fences)
            output_path: Output SVG file path

        Returns:
            True if successful, False otherwise
        """
        logger.info(f"Generating SVG: {output_path.name}")

        # Strip markdown fences if present
        content = self._strip_markdown_fences(mermaid_content)

        # Try API first
        if self._generate_via_api(content, output_path, 'svg'):
            logger.info(f"  [OK] SVG generated via API: {output_path.name}")
            return True

        # Fallback to local mmdc
        logger.warning("  API failed, trying local mmdc...")
        if self._generate_via_mmdc(content, output_path, 'svg'):
            logger.info(f"  [OK] SVG generated via mmdc: {output_path.name}")
            return True

        logger.error(f"  [FAILED] Could not generate SVG: {output_path.name}")
        return False

    def generate_docx(self, app_name: str, png_path: Path, output_path: Path,
                     classification_summary: Optional[str] = None) -> bool:
        """Generate Word document with embedded PNG diagram

        Args:
            app_name: Application name
            png_path: Path to PNG diagram to embed
            output_path: Output DOCX file path
            classification_summary: Optional text summary of server classification

        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return False

        logger.info(f"Generating DOCX: {output_path.name}")

        try:
            # Create document
            doc = Document()

            # Title
            title = doc.add_heading(f'Network Diagram: {app_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            doc.add_paragraph(f'Generated: {time.strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()

            # Add classification summary if provided
            if classification_summary:
                doc.add_heading('Server Classification Summary', level=1)
                doc.add_paragraph(classification_summary)
                doc.add_paragraph()

            # Add diagram section
            doc.add_heading('Architecture Diagram', level=1)

            # Add note about diagram
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(
                'Note: This diagram shows the network architecture with server classification. '
                'For best quality, SVG files are available for manual import (infinite zoom capability).'
            )
            note_run.font.size = Pt(10)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()

            # Embed PNG diagram (check if file exists)
            if png_path.exists():
                # Add diagram with 8.5 inch width (fits standard page with margins)
                doc.add_picture(str(png_path), width=Inches(6.5))
                logger.info(f"  Embedded PNG: {png_path.name}")
            else:
                logger.warning(f"  PNG not found: {png_path.name}")
                doc.add_paragraph(f'[PNG diagram not available: {png_path.name}]')

            # Add instructions for better quality
            doc.add_page_break()
            doc.add_heading('For Best Diagram Quality', level=1)

            instructions = [
                ('Automated (Good Quality)', [
                    'Word documents automatically include high-resolution PNG diagrams (4800px width).',
                    'This provides excellent quality for most use cases.'
                ]),
                ('Manual (Perfect Quality - Recommended for Presentations)', [
                    f'For infinite zoom capability:',
                    f'1. Locate SVG file: {png_path.with_suffix(".svg").name}',
                    '2. In Word: Insert → Pictures → Select SVG file',
                    '3. Replace existing PNG diagram',
                    '4. Result: Perfect quality at any zoom level'
                ]),
                ('Use HTML Diagrams', [
                    f'Open: {png_path.with_suffix(".html").name}',
                    '- Perfect zoom in web browser',
                    '- Interactive pan/zoom controls',
                    '- Ideal for presentations via browser'
                ])
            ]

            for section_title, items in instructions:
                doc.add_heading(section_title, level=2)
                for item in items:
                    doc.add_paragraph(item, style='List Bullet' if not item.endswith(':') else None)

            # Save document
            doc.save(str(output_path))
            logger.info(f"  [OK] DOCX generated: {output_path.name}")
            return True

        except Exception as e:
            logger.error(f"  [FAILED] DOCX generation error: {e}")
            return False

    def _strip_markdown_fences(self, content: str) -> str:
        """Strip markdown code fences from Mermaid content

        Args:
            content: Mermaid content (may include ```mermaid fences)

        Returns:
            Clean Mermaid content without fences
        """
        if '```mermaid' in content:
            lines = content.split('\n')
            graph_lines = []
            in_graph = False

            for line in lines:
                stripped = line.strip()
                if stripped.startswith('```mermaid'):
                    in_graph = True
                    continue
                elif stripped == '```':
                    in_graph = False
                    break
                elif in_graph:
                    graph_lines.append(line)

            return '\n'.join(graph_lines).strip()
        else:
            return content.strip()

    def _generate_via_api(self, content: str, output_path: Path,
                         format_type: str, width: int = 4800) -> bool:
        """Generate diagram via Mermaid.ink API

        Args:
            content: Clean Mermaid content (without fences)
            output_path: Output file path
            format_type: 'png' or 'svg'
            width: PNG width (only for PNG format)

        Returns:
            True if successful, False otherwise
        """
        retry_count = 0
        success = False

        while retry_count < self.max_retries and not success:
            try:
                # Encode diagram for URL (base64)
                encoded = base64.urlsafe_b64encode(content.encode('utf-8')).decode('ascii')

                # Build URL
                if format_type == 'png':
                    url = f"https://mermaid.ink/img/{encoded}?type=png&width={width}"
                else:  # svg
                    url = f"https://mermaid.ink/svg/{encoded}"

                # Download file
                req = urllib.request.Request(url)
                req.add_header('User-Agent', 'NetworkSegmentationAnalyzer/1.0')

                with urllib.request.urlopen(req, timeout=self.api_timeout) as response:
                    if response.status == 200:
                        data = response.read()

                        # Validate content
                        if format_type == 'png':
                            is_valid = data.startswith(b'\x89PNG') and len(data) > 10240
                        else:  # svg
                            is_valid = b'<svg' in data and len(data) > 100

                        if is_valid:
                            with open(output_path, 'wb') as f:
                                f.write(data)
                            success = True
                        else:
                            logger.warning(f"  Invalid {format_type.upper()} data received (retry {retry_count + 1})")
                            retry_count += 1
                            time.sleep(3)
                    else:
                        logger.warning(f"  API returned status {response.status} (retry {retry_count + 1})")
                        retry_count += 1
                        time.sleep(2)

                # Be nice to the API
                if success:
                    time.sleep(1.5)

            except Exception as e:
                error_msg = str(e)
                if '503' in error_msg or 'Service Unavailable' in error_msg:
                    logger.warning(f"  API unavailable (503), retry {retry_count + 1}")
                    retry_count += 1
                    time.sleep(3)
                else:
                    logger.warning(f"  API error: {error_msg}")
                    retry_count = self.max_retries
                    break

        return success

    def _generate_via_mmdc(self, content: str, output_path: Path,
                          format_type: str, width: int = 4800) -> bool:
        """Generate diagram via local mmdc (Mermaid CLI)

        Args:
            content: Clean Mermaid content (without fences)
            output_path: Output file path
            format_type: 'png' or 'svg'
            width: PNG width (only for PNG format)

        Returns:
            True if successful, False otherwise
        """
        # Write content to temporary file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            # Run mmdc
            cmd = ['mmdc', '-i', tmp_path, '-o', str(output_path)]
            if format_type == 'png':
                cmd.extend(['-w', str(width), '-b', 'transparent'])

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )

            # Clean up temp file
            Path(tmp_path).unlink()

            return result.returncode == 0

        except FileNotFoundError:
            # mmdc not found
            logger.warning("  mmdc not installed (install with: npm install -g @mermaid-js/mermaid-cli)")
            Path(tmp_path).unlink()
            return False
        except Exception as e:
            # Any other error
            logger.warning(f"  mmdc error: {e}")
            try:
                Path(tmp_path).unlink()
            except:
                pass
            return False


# Convenience functions
def generate_all_formats(mermaid_content: str, base_path: Path,
                        app_name: Optional[str] = None,
                        classification_summary: Optional[str] = None) -> Tuple[bool, bool, bool]:
    """Generate PNG, SVG, and DOCX from Mermaid content

    Args:
        mermaid_content: Mermaid diagram content
        base_path: Base path (without extension) for output files
        app_name: Application name (for DOCX title)
        classification_summary: Optional classification summary for DOCX

    Returns:
        Tuple of (png_success, svg_success, docx_success)
    """
    generator = DiagramFormatGenerator()

    # Generate PNG
    png_path = base_path.with_suffix('.png')
    png_success = generator.generate_png(mermaid_content, png_path)

    # Generate SVG
    svg_path = base_path.with_suffix('.svg')
    svg_success = generator.generate_svg(mermaid_content, svg_path)

    # Generate DOCX (only if PNG was successful)
    docx_path = base_path.with_suffix('.docx')
    docx_success = False
    if png_success:
        app_name = app_name or base_path.stem
        docx_success = generator.generate_docx(app_name, png_path, docx_path, classification_summary)
    else:
        logger.warning(f"Skipping DOCX generation (PNG required): {docx_path.name}")

    return png_success, svg_success, docx_success


if __name__ == '__main__':
    # Test/demo code
    print("Diagram Format Generator")
    print("=" * 80)
    print("Usage:")
    print("  from src.diagram_format_generator import generate_all_formats")
    print("  generate_all_formats(mermaid_content, Path('output/diagram'), 'MyApp')")
    print()
    print("Capabilities:")
    print("  - PNG: 4800px high resolution via Mermaid.ink API")
    print("  - SVG: Infinite zoom vector graphics")
    print("  - DOCX: Word documents with embedded diagrams")
    print("  - Automatic fallback to local mmdc if API fails")

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DNS Validation Reporter
========================
Generates comprehensive reports of DNS validation findings across all applications

Features:
- DNS mismatch detection and reporting
- Multiple IP tracking (VM + ESXi)
- NXDOMAIN identification
- Word document reports
- CSV exports for analysis
- JSON summary data

Author: Enterprise Security Team
Version: 1.0
"""

import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Set, Optional
from collections import defaultdict
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DNSValidationReporter:
    """
    Generate DNS validation reports from hostname resolver validation data

    Analyzes DNS validation results across all applications and generates:
    - Word document with findings and recommendations
    - CSV export for detailed analysis
    - JSON summary for programmatic access
    """

    def __init__(self, output_dir: str = './outputs_final/dns_reports'):
        """
        Initialize DNS validation reporter

        Args:
            output_dir: Directory for output reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Validation findings
        self.all_validations: List[Dict] = []
        self.mismatches: List[Dict] = []
        self.multiple_ips: List[Dict] = []
        self.nxdomain: List[Dict] = []
        self.failed: List[Dict] = []

        # Statistics
        self.stats = {
            'total_ips_validated': 0,
            'total_valid': 0,
            'total_valid_multiple': 0,
            'total_mismatches': 0,
            'total_nxdomain': 0,
            'total_failed': 0,
            'applications_analyzed': 0,
            'timestamp': datetime.now().isoformat()
        }

        logger.info("DNS Validation Reporter initialized")
        logger.info(f"  Output directory: {self.output_dir}")

    def add_validation_data(self, app_id: str, validation_metadata: Dict[str, Dict]):
        """
        Add validation data from a single application

        Args:
            app_id: Application identifier
            validation_metadata: Dictionary of IP -> validation results
        """
        self.stats['applications_analyzed'] += 1

        for ip, validation in validation_metadata.items():
            # Add app_id to validation record
            validation['app_id'] = app_id
            validation['ip'] = ip

            self.all_validations.append(validation)

            # Categorize by status
            status = validation.get('status', 'error')

            if status == 'valid':
                self.stats['total_valid'] += 1
            elif status == 'valid_multiple_ips':
                self.stats['total_valid_multiple'] += 1
                self.multiple_ips.append(validation)
            elif status == 'mismatch':
                self.stats['total_mismatches'] += 1
                self.mismatches.append(validation)
            elif status == 'nxdomain':
                self.stats['total_nxdomain'] += 1
                self.nxdomain.append(validation)
            else:
                self.stats['total_failed'] += 1
                self.failed.append(validation)

        self.stats['total_ips_validated'] = len(self.all_validations)

    def generate_word_report(self, output_path: Optional[str] = None) -> str:
        """
        Generate comprehensive Word document report

        Args:
            output_path: Optional custom output path

        Returns:
            Path to generated report
        """
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.output_dir / f'DNS_Validation_Report_{timestamp}.docx'

        logger.info(f"Generating Word report: {output_path}")

        doc = Document()

        # Title
        title = doc.add_heading('DNS Validation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(
            f'Network Segmentation Analyzer\n'
            f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Executive Summary
        self._add_executive_summary(doc)

        # Summary Statistics
        self._add_summary_statistics(doc)

        # DNS Mismatches
        if self.mismatches:
            self._add_mismatch_section(doc)

        # Multiple IPs (VM + ESXi)
        if self.multiple_ips:
            self._add_multiple_ip_section(doc)

        # NXDOMAIN Issues
        if self.nxdomain:
            self._add_nxdomain_section(doc)

        # Failed Validations
        if self.failed:
            self._add_failed_section(doc)

        # Recommendations
        self._add_recommendations(doc)

        # Save document
        doc.save(output_path)
        logger.info(f"  Word report saved: {output_path}")

        return str(output_path)

    def _add_executive_summary(self, doc):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"This report provides a comprehensive analysis of DNS validation findings "
            f"across {self.stats['applications_analyzed']} applications. "
        )
        para.add_run(
            f"The validation process performed bidirectional DNS checks (forward and reverse) "
            f"on {self.stats['total_ips_validated']} unique IP addresses to identify "
            f"configuration issues, multiple IP scenarios, and DNS mismatches."
        )

        doc.add_paragraph()

        # Key findings
        findings_para = doc.add_paragraph()
        findings_para.add_run("Key Findings:\n").bold = True

        if self.stats['total_mismatches'] > 0:
            findings_para.add_run(
                f"• {self.stats['total_mismatches']} DNS mismatches detected "
                f"(forward and reverse DNS do not match)\n"
            )

        if self.stats['total_valid_multiple'] > 0:
            findings_para.add_run(
                f"• {self.stats['total_valid_multiple']} hosts with multiple IP addresses "
                f"(VM + ESXi scenarios)\n"
            )

        if self.stats['total_nxdomain'] > 0:
            findings_para.add_run(
                f"• {self.stats['total_nxdomain']} non-existent domain names (NXDOMAIN)\n"
            )

        if self.stats['total_valid'] > 0:
            findings_para.add_run(
                f"• {self.stats['total_valid']} IP addresses with valid DNS configuration\n"
            )

        doc.add_paragraph()

    def _add_summary_statistics(self, doc):
        """Add summary statistics table"""
        doc.add_heading('Summary Statistics', 1)

        # Create table
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Count'

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        rows_data = [
            ('Applications Analyzed', self.stats['applications_analyzed']),
            ('Total IPs Validated', self.stats['total_ips_validated']),
            ('Valid DNS (Perfect Match)', self.stats['total_valid']),
            ('Valid DNS (Multiple IPs)', self.stats['total_valid_multiple']),
            ('DNS Mismatches', self.stats['total_mismatches']),
            ('NXDOMAIN', self.stats['total_nxdomain']),
            ('Validation Failures', self.stats['total_failed']),
        ]

        for idx, (label, value) in enumerate(rows_data, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = label
            row_cells[1].text = str(value)

        doc.add_paragraph()

    def _add_mismatch_section(self, doc):
        """Add DNS mismatch findings section"""
        doc.add_heading('DNS Mismatches', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"Found {len(self.mismatches)} DNS mismatches where forward and reverse DNS do not match. "
        )
        para.add_run(
            "This can indicate DNS configuration errors, stale DNS records, or security issues."
        )

        doc.add_paragraph()

        # Create table
        if len(self.mismatches) <= 100:  # Limit table size
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Application'
            header_cells[1].text = 'IP Address'
            header_cells[2].text = 'Reverse DNS'
            header_cells[3].text = 'Forward IP'
            header_cells[4].text = 'Issue'

            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Data rows
            for mismatch in self.mismatches[:100]:  # Limit to first 100
                row_cells = table.add_row().cells
                row_cells[0].text = mismatch.get('app_id', 'N/A')
                row_cells[1].text = mismatch.get('ip', 'N/A')
                row_cells[2].text = mismatch.get('reverse_hostname', 'N/A') or 'N/A'
                row_cells[3].text = mismatch.get('forward_ip', 'N/A') or 'N/A'
                row_cells[4].text = mismatch.get('mismatch', 'Unknown')
        else:
            para = doc.add_paragraph()
            para.add_run(f"Note: {len(self.mismatches)} mismatches found. ")
            para.add_run("See CSV export for complete list.").italic = True

        doc.add_paragraph()

    def _add_multiple_ip_section(self, doc):
        """Add multiple IP findings section"""
        doc.add_heading('Multiple IP Addresses (VM + ESXi)', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"Found {len(self.multiple_ips)} hostnames with multiple IP addresses. "
        )
        para.add_run(
            "This is common in virtualized environments where both the VM and ESXi host "
            "share the same hostname."
        )

        doc.add_paragraph()

        # Create table
        if len(self.multiple_ips) <= 50:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Application'
            header_cells[1].text = 'Hostname'
            header_cells[2].text = 'Primary IP'
            header_cells[3].text = 'All IPs'

            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Data rows
            for multi_ip in self.multiple_ips[:50]:
                row_cells = table.add_row().cells
                row_cells[0].text = multi_ip.get('app_id', 'N/A')
                row_cells[1].text = multi_ip.get('reverse_hostname', 'N/A') or 'N/A'
                row_cells[2].text = multi_ip.get('ip', 'N/A')

                all_ips = multi_ip.get('forward_ips', [])
                row_cells[3].text = ', '.join(all_ips) if all_ips else 'N/A'
        else:
            para = doc.add_paragraph()
            para.add_run(f"Note: {len(self.multiple_ips)} hosts with multiple IPs. ")
            para.add_run("See CSV export for complete list.").italic = True

        doc.add_paragraph()

    def _add_nxdomain_section(self, doc):
        """Add NXDOMAIN findings section"""
        doc.add_heading('Non-Existent Domains (NXDOMAIN)', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"Found {len(self.nxdomain)} IP addresses with no reverse DNS records (NXDOMAIN). "
        )
        para.add_run(
            "This can indicate temporary addresses, decommissioned systems, or network devices "
            "without DNS registration."
        )

        doc.add_paragraph()

        # Group by application
        nxdomain_by_app = defaultdict(list)
        for nxd in self.nxdomain:
            nxdomain_by_app[nxd.get('app_id', 'Unknown')].append(nxd.get('ip', 'N/A'))

        # List by application
        for app_id, ips in sorted(nxdomain_by_app.items())[:20]:  # Limit to 20 apps
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(f"{app_id}: ").bold = True
            para.add_run(f"{len(ips)} IPs ({', '.join(ips[:5])}{'...' if len(ips) > 5 else ''})")

        doc.add_paragraph()

    def _add_failed_section(self, doc):
        """Add validation failure section"""
        doc.add_heading('Validation Failures', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"Found {len(self.failed)} IP addresses where DNS validation failed. "
        )
        para.add_run(
            "This can indicate DNS timeouts, network connectivity issues, or DNS server problems."
        )

        doc.add_paragraph()

    def _add_recommendations(self, doc):
        """Add recommendations section"""
        doc.add_heading('Recommendations', 1)

        recommendations = []

        if self.stats['total_mismatches'] > 0:
            recommendations.append(
                "DNS Mismatches: Review and correct DNS records where forward and reverse "
                "DNS do not match. This can cause issues with hostname-based security policies "
                "and network monitoring tools."
            )

        if self.stats['total_valid_multiple'] > 0:
            recommendations.append(
                "Multiple IPs: Document VM and ESXi host relationships. Ensure network "
                "segmentation policies account for both IP addresses when defining security zones."
            )

        if self.stats['total_nxdomain'] > 0:
            recommendations.append(
                "NXDOMAIN: Investigate IP addresses without reverse DNS records. Consider adding "
                "DNS records for active systems to improve network visibility and troubleshooting."
            )

        if self.stats['total_failed'] > 0:
            recommendations.append(
                "Validation Failures: Review DNS timeout settings and network connectivity to DNS "
                "servers. Consider increasing timeout values for remote networks."
            )

        # Add general recommendation
        recommendations.append(
            "Regular Validation: Schedule periodic DNS validation checks to identify configuration "
            "drift and maintain accurate network documentation."
        )

        for idx, rec in enumerate(recommendations, start=1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(rec)

        doc.add_paragraph()

    def export_to_csv(self, output_path: Optional[str] = None) -> str:
        """
        Export all validation data to CSV

        Args:
            output_path: Optional custom output path

        Returns:
            Path to exported CSV
        """
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.output_dir / f'DNS_Validation_Data_{timestamp}.csv'

        logger.info(f"Exporting to CSV: {output_path}")

        # Convert to DataFrame
        df = pd.DataFrame(self.all_validations)

        # Reorder columns for readability
        column_order = [
            'app_id', 'ip', 'status', 'valid', 'reverse_hostname',
            'forward_ip', 'forward_ips', 'mismatch', 'timestamp'
        ]

        # Only include columns that exist
        columns = [col for col in column_order if col in df.columns]
        df = df[columns]

        # Sort by status (mismatches first)
        status_order = {'mismatch': 0, 'valid_multiple_ips': 1, 'nxdomain': 2, 'valid': 3}
        df['sort_order'] = df['status'].map(lambda x: status_order.get(x, 4))
        df = df.sort_values(['sort_order', 'app_id', 'ip'])
        df = df.drop('sort_order', axis=1)

        # Export
        df.to_csv(output_path, index=False)
        logger.info(f"  CSV exported: {output_path} ({len(df)} records)")

        return str(output_path)

    def export_to_json(self, output_path: Optional[str] = None) -> str:
        """
        Export summary data to JSON

        Args:
            output_path: Optional custom output path

        Returns:
            Path to exported JSON
        """
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.output_dir / f'DNS_Validation_Summary_{timestamp}.json'

        logger.info(f"Exporting to JSON: {output_path}")

        summary_data = {
            'metadata': {
                'generated': datetime.now().isoformat(),
                'applications_analyzed': self.stats['applications_analyzed'],
                'total_ips_validated': self.stats['total_ips_validated']
            },
            'statistics': self.stats,
            'mismatches': self.mismatches,
            'multiple_ips': self.multiple_ips,
            'nxdomain': self.nxdomain[:100],  # Limit NXDOMAIN to first 100
            'all_validations': self.all_validations
        }

        with open(output_path, 'w') as f:
            json.dump(summary_data, f, indent=2)

        logger.info(f"  JSON exported: {output_path}")

        return str(output_path)

    def generate_all_reports(self) -> Dict[str, str]:
        """
        Generate all report formats

        Returns:
            Dictionary of report type -> file path
        """
        logger.info("\n" + "="*80)
        logger.info("GENERATING DNS VALIDATION REPORTS")
        logger.info("="*80)

        reports = {}

        # Word report
        reports['word'] = self.generate_word_report()

        # CSV export
        reports['csv'] = self.export_to_csv()

        # JSON export
        reports['json'] = self.export_to_json()

        logger.info("\n" + "="*80)
        logger.info("DNS VALIDATION REPORTS COMPLETE")
        logger.info("="*80)
        logger.info(f"  Word Report: {reports['word']}")
        logger.info(f"  CSV Export:  {reports['csv']}")
        logger.info(f"  JSON Export: {reports['json']}")
        logger.info("="*80 + "\n")

        return reports


def collect_dns_validation_from_apps(topology_dir: str = './persistent_data/topology') -> DNSValidationReporter:
    """
    Collect DNS validation data from all application topology files

    Args:
        topology_dir: Directory containing application topology JSON files

    Returns:
        DNSValidationReporter with collected data
    """
    topology_path = Path(topology_dir)
    reporter = DNSValidationReporter()

    logger.info(f"Collecting DNS validation data from: {topology_path}")

    # Find all topology JSON files
    json_files = list(topology_path.glob('*.json'))
    logger.info(f"Found {len(json_files)} topology files")

    for json_file in json_files:
        try:
            with open(json_file, 'r') as f:
                topology_data = json.load(f)

            # Extract app_id from filename or data
            app_id = json_file.stem

            # Look for DNS validation data in topology
            dns_validation = topology_data.get('dns_validation')
            validation_metadata = topology_data.get('validation_metadata', {})

            if validation_metadata:
                reporter.add_validation_data(app_id, validation_metadata)
                logger.debug(f"  Loaded validation data for {app_id}")

        except Exception as e:
            logger.warning(f"  Failed to load {json_file.name}: {e}")

    logger.info(f"Collected validation data for {reporter.stats['applications_analyzed']} applications")

    return reporter

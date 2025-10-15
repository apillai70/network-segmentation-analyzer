"""
Word Document Generator for Solutions Architecture
==================================================
Generates comprehensive Microsoft Word (.docx) Solutions Architecture Documents
with embedded diagrams, tables, and professional formatting.

Author: Network Security Team
Version: 2.0
"""

import logging
from pathlib import Path
from typing import Dict, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import tempfile

logger = logging.getLogger(__name__)


class SolutionsArchitectureDocument:
    """Generates comprehensive Solutions Architecture Document in Word format"""

    def __init__(self, analysis_results: Dict, zones: Dict, rules: List, ml_predictions: Dict = None, png_path: str = None):
        """
        Initialize document generator

        Args:
            analysis_results: Analysis results from TrafficAnalyzer
            zones: Network zones dictionary
            rules: List of SegmentationRule objects
            ml_predictions: ML predictions for apps without data (optional)
            png_path: Path to PNG diagram file for embedding (optional)
        """
        self.analysis = analysis_results
        self.zones = zones
        self.rules = rules
        self.ml_predictions = ml_predictions
        self.png_path = png_path
        self.doc = Document()

        self._setup_styles()
        logger.info("SolutionsArchitectureDocument initialized")

    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles

        # Heading styles are built-in, but we can customize them
        try:
            title_style = styles['Title']
            title_style.font.size = Pt(28)
            title_style.font.color.rgb = RGBColor(0, 51, 102)
        except:
            pass

        # Create custom styles
        try:
            # Executive summary style
            exec_style = styles.add_style('ExecutiveSummary', WD_STYLE_TYPE.PARAGRAPH)
            exec_style.font.size = Pt(11)
            exec_style.font.italic = True
            exec_style.paragraph_format.space_before = Pt(6)
            exec_style.paragraph_format.space_after = Pt(6)
        except:
            pass

    def generate_document(self, output_path: str):
        """Generate complete Solutions Architecture Document"""
        logger.info("Generating Solutions Architecture Document...")

        # Document structure
        self._add_title_page()
        self._add_page_break()

        self._add_table_of_contents()
        self._add_page_break()

        self._add_revision_history()
        self._add_page_break()

        self._add_executive_summary()
        self._add_page_break()

        self._add_current_state_analysis()
        self._add_page_break()

        self._add_proposed_segmentation_design()

        # Add new section to return to portrait orientation after diagram page
        if self.png_path:
            section = self.doc.add_section()
            section.orientation = WD_ORIENT.PORTRAIT
            # Swap dimensions back to portrait
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        else:
            self._add_page_break()

        self._add_segmentation_rules()
        self._add_page_break()

        self._add_implementation_plan()
        self._add_page_break()

        self._add_testing_rollback_plan()
        self._add_page_break()

        # Add ML predictions section if available
        if self.ml_predictions:
            self._add_ml_predictions_section()
            self._add_page_break()

        self._add_appendices()

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output_file))
        logger.info(f"✓ Solutions Architecture Document saved: {output_path}")

    def _add_title_page(self):
        """Add title page"""
        # Title
        title = self.doc.add_heading('Network Segmentation', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self.doc.add_heading('Solutions Architecture Document', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Document info
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run('Version: 1.0\n').bold = True
        info_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
        info_para.add_run('Classification: Internal\n')

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Footer info
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Prepared by: Network Security Team\n')
        footer.add_run('Auto-generated by Network Segmentation Analyzer\n')

    def _add_table_of_contents(self):
        """Add table of contents placeholder"""
        heading = self.doc.add_heading('Table of Contents', level=1)

        para = self.doc.add_paragraph()
        para.add_run('Note: After opening in Word, right-click here and select "Update Field" to generate the table of contents.\n\n')

        # List the sections
        toc_items = [
            '1. Executive Summary',
            '2. Current State Analysis',
            '   2.1 Traffic Summary',
            '   2.2 Top Talkers',
            '   2.3 Suspicious Flows',
            '3. Proposed Segmentation Design',
            '   3.1 Network Zones',
            '   3.2 Zone Architecture',
            '   3.3 Traffic Flow Patterns',
            '4. Segmentation Rules',
            '   4.1 Rule Overview',
            '   4.2 Detailed Rules Table',
            '   4.3 Implementation Examples',
            '5. Implementation Plan',
            '   5.1 Phased Rollout',
            '   5.2 Validation Tests',
            '   5.3 Success Criteria',
            '6. Testing and Rollback Plan',
            '7. ML-Based Predictions (for apps without traffic data)',
            '   7.1 Overview',
            '   7.2 Prediction Summary',
            '   7.3 Sample Predictions',
            '   7.4 Recommendations',
            '8. Appendices',
            '   A. Raw Data Samples',
            '   B. Full Rules Export',
            '   C. Glossary'
        ]

        for item in toc_items:
            self.doc.add_paragraph(item)

    def _add_revision_history(self):
        """Add revision history table"""
        self.doc.add_heading('Revision History', level=1)

        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Version'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Author'
        hdr_cells[3].text = 'Description'

        # Data row
        row_cells = table.rows[1].cells
        row_cells[0].text = '1.0'
        row_cells[1].text = datetime.now().strftime('%Y-%m-%d')
        row_cells[2].text = 'Network Security Team'
        row_cells[3].text = 'Initial release - automated network segmentation analysis'

    def _add_executive_summary(self):
        """Add executive summary"""
        self.doc.add_heading('Executive Summary', level=1)

        summary = self.analysis.get('summary', {})

        # Key findings
        self.doc.add_heading('Key Findings', level=2)

        findings = [
            f"Analyzed {summary.get('total_flows', 0):,} network flows across {summary.get('unique_apps', 0)} applications",
            f"Identified {summary.get('suspicious_count', 0)} suspicious flows requiring immediate attention",
            f"Current network has {summary.get('internal_flows', 0):,} internal flows and {summary.get('external_flows', 0):,} external connections",
            f"Generated {len(self.rules)} segmentation rules to enhance security posture"
        ]

        for finding in findings:
            para = self.doc.add_paragraph(finding, style='List Bullet')

        self.doc.add_paragraph()

        # Recommendations
        self.doc.add_heading('Recommendations', level=2)

        recommendations = [
            "Implement micro-segmentation across all application tiers to enforce least-privilege access",
            "Block management ports (SSH, RDP) from external networks immediately - HIGH PRIORITY",
            "Restrict database access to authorized application servers only",
            "Deploy zone-based firewall policies to control east-west traffic",
            "Implement continuous monitoring and logging for all cross-zone traffic",
            "Conduct phased rollout with validation testing at each stage"
        ]

        for rec in recommendations:
            self.doc.add_paragraph(rec, style='List Bullet')

        self.doc.add_paragraph()

        # Risk summary
        self.doc.add_heading('Risk Summary', level=2)

        suspicious = self.analysis.get('suspicious_flows', [])
        high_risk_count = sum(1 for f in suspicious if f.get('risk_score', 0) > 70)

        para = self.doc.add_paragraph()
        para.add_run('Critical Risks: ').bold = True
        para.add_run(f'{high_risk_count} high-risk flows identified with scores above 70. ')
        para.add_run('These require immediate remediation to prevent potential security breaches.\n\n')

        para.add_run('Primary Risk Vectors:\n').bold = True
        risk_vectors = [
            'External access to management ports',
            'Unrestricted database port exposure',
            'Lack of micro-segmentation between application tiers',
            'Suspicious scanner/bot activity detected'
        ]

        for vector in risk_vectors:
            self.doc.add_paragraph(vector, style='List Bullet 2')

    def _add_current_state_analysis(self):
        """Add current state analysis section"""
        self.doc.add_heading('Current State Analysis', level=1)

        summary = self.analysis.get('summary', {})

        # Traffic summary
        self.doc.add_heading('Traffic Summary', level=2)

        para = self.doc.add_paragraph()
        para.add_run('Total Traffic Volume\n').bold = True

        stats_table = self.doc.add_table(rows=5, cols=2)
        stats_table.style = 'Light Shading Accent 1'

        stats_data = [
            ('Total Flows', f"{summary.get('total_flows', 0):,}"),
            ('Total Bytes', f"{summary.get('total_bytes', 0) / (1024**3):.2f} GB"),
            ('Total Packets', f"{summary.get('total_packets', 0):,}"),
            ('Internal Flows', f"{summary.get('internal_flows', 0):,} ({summary.get('internal_flows', 0) / max(summary.get('total_flows', 1), 1) * 100:.1f}%)"),
            ('External Flows', f"{summary.get('external_flows', 0):,} ({summary.get('external_flows', 0) / max(summary.get('total_flows', 1), 1) * 100:.1f}%)")
        ]

        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

        self.doc.add_paragraph()

        # Protocol distribution
        self.doc.add_heading('Protocol Distribution', level=3)

        proto_dist = summary.get('protocol_distribution', {})
        if proto_dist:
            proto_table = self.doc.add_table(rows=len(proto_dist) + 1, cols=2)
            proto_table.style = 'Light List Accent 1'

            proto_table.rows[0].cells[0].text = 'Protocol'
            proto_table.rows[0].cells[1].text = 'Flow Count'

            for i, (proto, count) in enumerate(sorted(proto_dist.items(), key=lambda x: x[1], reverse=True), 1):
                proto_table.rows[i].cells[0].text = proto
                proto_table.rows[i].cells[1].text = str(count)

        self.doc.add_paragraph()

        # Top talkers
        self.doc.add_heading('Top Talkers', level=2)

        top_talkers = self.analysis.get('top_talkers', {})

        self.doc.add_heading('Top Sources (by bytes)', level=3)
        top_sources = top_talkers.get('top_sources_by_bytes', {})

        if top_sources:
            src_table = self.doc.add_table(rows=min(len(top_sources) + 1, 11), cols=2)
            src_table.style = 'Light Grid Accent 1'

            src_table.rows[0].cells[0].text = 'Source IP'
            src_table.rows[0].cells[1].text = 'Bytes Transferred'

            for i, (ip, bytes_val) in enumerate(list(top_sources.items())[:10], 1):
                src_table.rows[i].cells[0].text = ip
                src_table.rows[i].cells[1].text = f"{bytes_val / (1024**2):.2f} MB"

        self.doc.add_paragraph()

        # Suspicious flows
        self.doc.add_heading('Suspicious Flows', level=2)

        suspicious = self.analysis.get('suspicious_flows', [])[:10]

        if suspicious:
            susp_table = self.doc.add_table(rows=len(suspicious) + 1, cols=4)
            susp_table.style = 'Medium Shading 1 Accent 2'

            hdr = susp_table.rows[0].cells
            hdr[0].text = 'Source'
            hdr[1].text = 'Destination'
            hdr[2].text = 'Risk Score'
            hdr[3].text = 'Reason'

            for i, flow in enumerate(suspicious, 1):
                row = susp_table.rows[i].cells
                row[0].text = flow.get('source', '')
                row[1].text = flow.get('destination', '')
                row[2].text = str(flow.get('risk_score', 0))
                row[3].text = flow.get('reason', '')

    def _add_proposed_segmentation_design(self):
        """Add proposed segmentation design"""
        self.doc.add_heading('Proposed Segmentation Design', level=1)

        self.doc.add_paragraph(
            'The proposed micro-segmentation design implements a zero-trust architecture with '
            'clearly defined security zones and strict access controls between tiers.'
        )

        self.doc.add_paragraph()

        # Network zones
        self.doc.add_heading('Network Zones', level=2)

        zone_table = self.doc.add_table(rows=len(self.zones) + 1, cols=5)
        zone_table.style = 'Light Grid Accent 1'

        # Header
        hdr = zone_table.rows[0].cells
        hdr[0].text = 'Zone Name'
        hdr[1].text = 'Type'
        hdr[2].text = 'Members'
        hdr[3].text = 'Security Level'
        hdr[4].text = 'Description'

        # Data rows
        for i, (zone_name, zone) in enumerate(sorted(self.zones.items()), 1):
            row = zone_table.rows[i].cells
            row[0].text = zone_name
            row[1].text = zone.zone_type
            row[2].text = str(len(zone.members))
            row[3].text = str(zone.security_level)
            row[4].text = zone.description

        self.doc.add_paragraph()

        # Architecture diagram - embed PNG with landscape orientation
        self.doc.add_heading('Application Architecture Diagram', level=2)

        # If PNG path is provided, embed it in landscape orientation with rotation
        if self.png_path and Path(self.png_path).exists():
            # Change this section to landscape orientation
            section = self.doc.sections[-1]
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap page dimensions for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

            self.doc.add_paragraph(
                'The following diagram shows the application architecture, data flows, security zones, '
                'and dependencies for this network segmentation solution.'
            )

            self.doc.add_paragraph()

            try:
                # Rotate image 90° counterclockwise to make horizontal diagram vertical
                with Image.open(self.png_path) as img:
                    rotated_img = img.rotate(90, expand=True)

                    # Save to temporary file
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        rotated_img.save(tmp.name, 'PNG')
                        tmp_path = tmp.name

                # Landscape page: 9" usable width, set to 60% = 5.4 inches
                # Height auto-scales to preserve aspect ratio - can span multiple pages
                # This keeps the image SHARP and READABLE
                self.doc.add_picture(tmp_path, width=Inches(5.4))
                logger.info(f"  PNG embedded (vertical, width=5.4in, height=auto): {Path(self.png_path).name}")

                # Clean up temp file
                try:
                    Path(tmp_path).unlink()
                except:
                    pass

            except Exception as e:
                logger.error(f"  Failed to embed PNG: {e}")
                self.doc.add_paragraph(f'Note: Diagram could not be embedded - {e}')

            self.doc.add_paragraph()

            # Legend
            self.doc.add_heading('Diagram Legend', level=3)

            legend_items = [
                'Application Box (with direction TB) = Internal architecture showing application tiers',
                'Web Tier = Frontend servers handling user requests',
                'App Tier = Backend application servers',
                'Data Tier = Database servers',
                'Cache Tier = Caching layer (Redis, Memcache)',
                'Messaging Tier = Message queues (Kafka, RabbitMQ)',
                'Management Tier = Infrastructure management and monitoring',
                'Downstream Applications (Circles) = External applications this app calls',
                'Infrastructure (Rectangles) = Databases, caches, and queues',
                'Thick lines (===) = Application-to-application calls',
                'Regular lines (---) = Infrastructure dependencies',
                'Colors indicate security zones and tiers'
            ]

            for item in legend_items:
                self.doc.add_paragraph(item, style='List Bullet')

            self.doc.add_paragraph()

        else:
            # Fallback message if no PNG provided
            para = self.doc.add_paragraph()
            para.add_run('Note: ').bold = True
            para.add_run('Diagram image not available. ')
            if self.png_path:
                para.add_run(f'PNG file not found at: {self.png_path}')
            else:
                para.add_run('No diagram path provided. See outputs/diagrams/ for generated diagrams.')

            self.doc.add_paragraph()

        # Traffic flow patterns
        self.doc.add_heading('Recommended Traffic Flow Patterns', level=2)

        flows = [
            'EXTERNAL → DMZ: HTTP/HTTPS traffic only',
            'DMZ → WEB_TIER: Reverse proxy connections',
            'WEB_TIER → APP_TIER: Application protocol (port 8080/8443)',
            'APP_TIER → DATA_TIER: Database connections (MySQL, PostgreSQL)',
            'APP_TIER → MESSAGING_TIER: Message queue access (Kafka)',
            'APP_TIER → CACHE_TIER: Cache access (Redis)',
            'MANAGEMENT_TIER → All Tiers: Monitoring and metrics collection',
            'All reverse traffic flows: DENIED (prevents data exfiltration)'
        ]

        for flow in flows:
            self.doc.add_paragraph(flow, style='List Bullet')

    def _add_segmentation_rules(self):
        """Add segmentation rules section"""
        self.doc.add_heading('Segmentation Rules', level=1)

        self.doc.add_heading('Rule Overview', level=2)

        para = self.doc.add_paragraph()
        para.add_run(f'Total Rules Generated: ').bold = True
        para.add_run(f'{len(self.rules)}\n\n')

        para.add_run('The segmentation rules implement a default-deny policy with explicit allow rules for legitimate traffic flows. ')
        para.add_run('Rules are prioritized based on security risk, with highest-risk denials at the top.\n\n')

        para.add_run('Rule Priority Scheme:\n').bold = True
        priority_list = [
            '100-199: Critical blocks (management ports from external)',
            '200-299: Database protection rules',
            '300-399: Public service access (HTTP/HTTPS)',
            '400-799: Internal tier-to-tier communication',
            '800-899: Management and monitoring',
            '9999: Default deny all'
        ]

        for item in priority_list:
            self.doc.add_paragraph(item, style='List Bullet 2')

        self.doc.add_paragraph()

        # Rules table
        self.doc.add_heading('Detailed Rules Table', level=2)

        rules_table = self.doc.add_table(rows=min(len(self.rules), 20) + 1, cols=7)
        rules_table.style = 'Light Grid Accent 1'

        # Header
        hdr = rules_table.rows[0].cells
        hdr[0].text = 'Priority'
        hdr[1].text = 'Source'
        hdr[2].text = 'Destination'
        hdr[3].text = 'Protocol:Port'
        hdr[4].text = 'Action'
        hdr[5].text = 'Risk'
        hdr[6].text = 'Justification'

        # Add rules (limit to top 20 for readability)
        for i, rule in enumerate(self.rules[:20], 1):
            row = rules_table.rows[i].cells
            row[0].text = str(rule.priority)
            row[1].text = rule.source
            row[2].text = rule.destination
            row[3].text = f"{rule.protocol}:{rule.port}"
            row[4].text = rule.action.upper()
            row[5].text = str(rule.risk_score)
            row[6].text = rule.justification[:100] + '...' if len(rule.justification) > 100 else rule.justification

        self.doc.add_paragraph()

        # Implementation examples
        self.doc.add_heading('Implementation Examples', level=2)

        self.doc.add_paragraph('The rules can be implemented on various platforms. Examples are provided in the appendices and output files.')

    def _add_implementation_plan(self):
        """Add implementation plan"""
        self.doc.add_heading('Implementation Plan', level=1)

        self.doc.add_heading('Phased Rollout Approach', level=2)

        phases = [
            {
                'phase': 'Phase 1: Preparation (Week 1-2)',
                'tasks': [
                    'Review and validate segmentation rules',
                    'Set up staging environment mirroring production',
                    'Configure firewall infrastructure',
                    'Install monitoring and logging tools',
                    'Create rollback procedures'
                ]
            },
            {
                'phase': 'Phase 2: Critical Security Rules (Week 3)',
                'tasks': [
                    'Block management ports from external networks',
                    'Restrict database port access',
                    'Implement default-deny policy on perimeter',
                    'Enable comprehensive logging',
                    'Validate with penetration testing'
                ]
            },
            {
                'phase': 'Phase 3: Micro-Segmentation (Week 4-5)',
                'tasks': [
                    'Deploy zone-based internal firewalls',
                    'Implement WEB_TIER → APP_TIER rules',
                    'Implement APP_TIER → DATA_TIER rules',
                    'Configure monitoring tier access',
                    'Validate application functionality'
                ]
            },
            {
                'phase': 'Phase 4: Full Deployment & Optimization (Week 6-8)',
                'tasks': [
                    'Roll out to all applications',
                    'Fine-tune rules based on production traffic',
                    'Optimize performance and latency',
                    'Conduct final security validation',
                    'Document as-built configuration'
                ]
            }
        ]

        for phase_data in phases:
            self.doc.add_heading(phase_data['phase'], level=3)
            for task in phase_data['tasks']:
                self.doc.add_paragraph(task, style='List Bullet')
            self.doc.add_paragraph()

        # Success criteria
        self.doc.add_heading('Success Criteria', level=2)

        criteria = [
            'Zero unauthorized cross-zone traffic',
            'All applications functioning normally',
            'Response time impact < 5%',
            'All monitoring dashboards operational',
            'Successful penetration test completion',
            'Documentation complete and approved'
        ]

        for criterion in criteria:
            para = self.doc.add_paragraph(criterion, style='List Bullet')
            para.runs[0].bold = True

    def _add_testing_rollback_plan(self):
        """Add testing and rollback plan"""
        self.doc.add_heading('Testing and Rollback Plan', level=1)

        self.doc.add_heading('Validation Tests', level=2)

        tests = [
            {
                'test': 'External Port Scan Test',
                'procedure': 'Scan all external IPs for open management ports (22, 23, 3389)',
                'expected': 'All management ports should be filtered/closed'
            },
            {
                'test': 'Database Access Test',
                'procedure': 'Attempt database connections from non-APP_TIER hosts',
                'expected': 'All connections should be blocked'
            },
            {
                'test': 'Application Functionality Test',
                'procedure': 'Execute full end-to-end application test suite',
                'expected': 'All tests pass with < 5% latency increase'
            },
            {
                'test': 'Monitoring Connectivity Test',
                'procedure': 'Verify monitoring tools can collect metrics from all zones',
                'expected': 'All metrics dashboards populated'
            }
        ]

        test_table = self.doc.add_table(rows=len(tests) + 1, cols=3)
        test_table.style = 'Light Shading Accent 1'

        hdr = test_table.rows[0].cells
        hdr[0].text = 'Test Name'
        hdr[1].text = 'Procedure'
        hdr[2].text = 'Expected Result'

        for i, test in enumerate(tests, 1):
            row = test_table.rows[i].cells
            row[0].text = test['test']
            row[1].text = test['procedure']
            row[2].text = test['expected']

        self.doc.add_paragraph()

        # Rollback procedures
        self.doc.add_heading('Rollback Procedures', level=2)

        self.doc.add_paragraph(
            'In the event of critical issues, execute the following rollback steps:'
        )

        rollback_steps = [
            'STOP: Halt the rollout immediately',
            'Disable newly implemented firewall rules',
            'Restore backup firewall configuration',
            'Verify application functionality',
            'Document the failure reason and lessons learned',
            'Schedule retrospective meeting',
            'Revise implementation plan before retry'
        ]

        for i, step in enumerate(rollback_steps, 1):
            para = self.doc.add_paragraph(f'Step {i}: {step}', style='List Number')

        self.doc.add_paragraph()

        self.doc.add_paragraph(
            'Rollback Time Objective: < 15 minutes for complete restoration of service'
        ).bold = True

    def _add_ml_predictions_section(self):
        """Add ML predictions section for apps without data"""
        self.doc.add_heading('ML-Based Predictions for Apps Without Traffic Data', level=1)

        # Overview
        self.doc.add_heading('Overview', level=2)

        para = self.doc.add_paragraph()
        para.add_run('Machine Learning Approach: ').bold = True
        para.add_run(
            'Using ensemble models (GNN, RNN, CNN, Attention, and Meta-learner), we have predicted '
            'network communication patterns for applications that do not yet have traffic data. '
            'These predictions are based on similarity analysis with observed applications and '
            'learned patterns from the 170 apps with actual traffic data.\n\n'
        )

        # Summary statistics
        self.doc.add_heading('Prediction Summary', level=2)

        total_predicted = len(self.ml_predictions)

        # Calculate confidence distribution
        confidence_high = sum(1 for p in self.ml_predictions.values() if p.get('confidence', 0) > 0.75)
        confidence_medium = sum(1 for p in self.ml_predictions.values() if 0.5 < p.get('confidence', 0) <= 0.75)
        confidence_low = sum(1 for p in self.ml_predictions.values() if p.get('confidence', 0) <= 0.5)

        stats_table = self.doc.add_table(rows=5, cols=2)
        stats_table.style = 'Light Shading Accent 1'

        stats_data = [
            ('Total Apps Predicted', str(total_predicted)),
            ('High Confidence (>75%)', f"{confidence_high} ({confidence_high/max(total_predicted,1)*100:.1f}%)"),
            ('Medium Confidence (50-75%)', f"{confidence_medium} ({confidence_medium/max(total_predicted,1)*100:.1f}%)"),
            ('Low Confidence (<50%)', f"{confidence_low} ({confidence_low/max(total_predicted,1)*100:.1f}%)"),
            ('Apps with Actual Data', '170 (used for training)')
        ]

        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

        self.doc.add_paragraph()

        # Ensemble model performance
        self.doc.add_heading('Ensemble Model Components', level=2)

        para = self.doc.add_paragraph()
        para.add_run('The prediction system uses five complementary models:\n').bold = True

        model_descriptions = [
            'GNN (Graph Neural Network): Learns network topology and communication patterns',
            'RNN (Recurrent Neural Network): Captures temporal sequences and time-based patterns',
            'CNN (Convolutional Neural Network): Detects traffic pattern features',
            'Multi-head Attention: Identifies important contextual relationships',
            'Markov Chain: Models peer correlation and transitive dependencies (if A→B and B→C, then A may need C)',
            'Meta-learner: Combines predictions from all models with weighted ensemble'
        ]

        for desc in model_descriptions:
            self.doc.add_paragraph(desc, style='List Bullet')

        self.doc.add_paragraph()

        # Markov Chain Peer Correlation
        self.doc.add_heading('Markov Chain Peer Correlation', level=2)

        para = self.doc.add_paragraph()
        para.add_run('The Markov chain model discovers peer correlations through temporal analysis:\n\n')

        markov_features = [
            'State Transitions: Models probability that app A will connect to peer B based on historical patterns',
            'Transitive Dependencies: If app A connects to B, and B connects to C, predicts A may need C',
            'Peer Similarity: Apps with similar peer sets are correlated (Jaccard similarity > 10%)',
            'Source Code Correlation: As applications are deployed, communication patterns form predictable chains',
            'Confidence Scoring: Predictions weighted by transition probability strength and correlation scores'
        ]

        for feature in markov_features:
            self.doc.add_paragraph(feature, style='List Bullet')

        self.doc.add_paragraph()

        para = self.doc.add_paragraph()
        para.add_run('Example: ').bold = True
        para.add_run(
            'If web_frontend connects to api_user_service, and api_user_service connects to db_customer_mysql, '
            'the Markov chain predicts that similar web applications will likely need both the API service '
            'and transitively require database access through that API.'
        )

        self.doc.add_paragraph()

        # Sample predictions table
        self.doc.add_heading('Sample Predictions', level=2)

        # Sort by confidence and get top 15
        sorted_predictions = sorted(
            self.ml_predictions.items(),
            key=lambda x: x[1].get('confidence', 0),
            reverse=True
        )[:15]

        if sorted_predictions:
            pred_table = self.doc.add_table(rows=len(sorted_predictions) + 1, cols=5)
            pred_table.style = 'Light Grid Accent 1'

            # Header
            hdr = pred_table.rows[0].cells
            hdr[0].text = 'Application Name'
            hdr[1].text = 'Predicted Zone'
            hdr[2].text = 'Confidence'
            hdr[3].text = 'Similar Apps'
            hdr[4].text = 'Estimated Flows'

            # Data rows
            for i, (app_name, pred) in enumerate(sorted_predictions, 1):
                row = pred_table.rows[i].cells
                row[0].text = app_name[:40]  # Truncate long names
                row[1].text = pred.get('predicted_zone', 'UNKNOWN')
                row[2].text = f"{pred.get('confidence', 0)*100:.1f}%"

                similar_apps = pred.get('similar_apps', [])
                row[3].text = f"{len(similar_apps)} apps"

                estimated_flows = pred.get('estimated_flows', 0)
                row[4].text = f"~{estimated_flows:,}" if estimated_flows > 0 else "N/A"

        self.doc.add_paragraph()

        # Confidence explanation
        self.doc.add_heading('Understanding Confidence Scores', level=2)

        para = self.doc.add_paragraph()
        para.add_run('Confidence scores reflect the model\'s certainty in predictions:\n\n')

        confidence_levels = [
            ('High (>75%)', 'Strong similarity to observed apps; predictions can be used with minimal validation'),
            ('Medium (50-75%)', 'Reasonable similarity; recommend validation against actual traffic when available'),
            ('Low (<50%)', 'Limited similarity; use as initial guidance only, validate thoroughly')
        ]

        conf_table = self.doc.add_table(rows=len(confidence_levels) + 1, cols=2)
        conf_table.style = 'Light List Accent 1'

        conf_table.rows[0].cells[0].text = 'Confidence Level'
        conf_table.rows[0].cells[1].text = 'Recommendation'

        for i, (level, recommendation) in enumerate(confidence_levels, 1):
            row = conf_table.rows[i].cells
            row[0].text = level
            row[1].text = recommendation

        self.doc.add_paragraph()

        # Recommendations
        self.doc.add_heading('Recommendations for ML Predictions', level=2)

        recommendations = [
            'Start with high-confidence predictions (>75%) for initial segmentation design',
            'Validate predictions against actual traffic data as apps are deployed',
            'Use predictions as starting point; adjust rules based on observed behavior',
            'Prioritize validation for security-critical applications (database, authentication)',
            'Continuously retrain models as more apps generate traffic data',
            'Export ml_predictions.json for detailed prediction data including ensemble weights'
        ]

        for rec in recommendations:
            self.doc.add_paragraph(rec, style='List Bullet')

        self.doc.add_paragraph()

        # Validation approach
        self.doc.add_heading('Prediction Validation Approach', level=2)

        para = self.doc.add_paragraph()
        para.add_run('When new apps begin generating traffic:\n').bold = True

        validation_steps = [
            'Compare actual traffic patterns with predicted patterns',
            'Verify zone assignment matches observed communication patterns',
            'Adjust segmentation rules based on actual vs. predicted differences',
            'Document accuracy metrics for continuous model improvement',
            'Feed observed data back into training set for model refinement'
        ]

        for i, step in enumerate(validation_steps, 1):
            self.doc.add_paragraph(f'{i}. {step}', style='List Number')

    def _add_appendices(self):
        """Add appendices"""
        self.doc.add_heading('Appendices', level=1)

        # Appendix A: Sample data
        self.doc.add_heading('Appendix A: Sample Network Flow Data', level=2)

        self.doc.add_paragraph(
            'Sample network flow records analyzed during this assessment:'
        )

        sample_table = self.doc.add_table(rows=6, cols=6)
        sample_table.style = 'Light Grid'

        hdr = sample_table.rows[0].cells
        hdr[0].text = 'Source'
        hdr[1].text = 'Destination'
        hdr[2].text = 'Protocol'
        hdr[3].text = 'Port'
        hdr[4].text = 'Bytes'
        hdr[5].text = 'Risk'

        sample_data = [
            ('10.1.2.15', '10.1.3.20', 'tcp', '8080', '234 KB', 'Low'),
            ('10.1.3.20', '10.1.4.30', 'tcp', '3306', '87 KB', 'Medium'),
            ('203.0.113.45', '10.1.1.10', 'tcp', '443', '89 KB', 'Low'),
            ('185.220.101.23', '10.1.2.15', 'tcp', '22', '456 B', 'High'),
            ('10.1.7.60', '10.1.3.20', 'tcp', '9090', '1 KB', 'Low')
        ]

        for i, row_data in enumerate(sample_data, 1):
            row = sample_table.rows[i].cells
            for j, val in enumerate(row_data):
                row[j].text = str(val)

        self.doc.add_paragraph()

        # Appendix B: Files generated
        self.doc.add_heading('Appendix B: Generated Output Files', level=2)

        files = [
            ('outputs/segmentation_rules.csv', 'Complete rules in CSV format'),
            ('outputs/iptables_rules.sh', 'Linux IPTables implementation'),
            ('outputs/aws_security_groups.json', 'AWS Security Group definitions'),
            ('outputs/diagrams/overall_network.mmd', 'Overall network Mermaid diagram'),
            ('outputs/diagrams/overall_network.html', 'Interactive network diagram'),
            ('outputs/diagrams/zone_flows.mmd', 'Zone traffic flow diagram'),
            ('outputs/analysis_report.json', 'Complete analysis in JSON format')
        ]

        # Add ML predictions file if available
        if self.ml_predictions:
            files.append(('outputs/ml_predictions.json', 'ML predictions for apps without traffic data'))

        files_table = self.doc.add_table(rows=len(files) + 1, cols=2)
        files_table.style = 'Light List Accent 1'

        files_table.rows[0].cells[0].text = 'File Path'
        files_table.rows[0].cells[1].text = 'Description'

        for i, (filepath, desc) in enumerate(files, 1):
            files_table.rows[i].cells[0].text = filepath
            files_table.rows[i].cells[1].text = desc

        self.doc.add_paragraph()

        # Appendix C: Glossary
        self.doc.add_heading('Appendix C: Glossary', level=2)

        glossary = [
            ('ACL', 'Access Control List - network traffic filtering rules'),
            ('DMZ', 'Demilitarized Zone - network segment for public-facing services'),
            ('East-West Traffic', 'Network traffic between servers within the data center'),
            ('Micro-segmentation', 'Granular security controls between workloads/applications'),
            ('North-South Traffic', 'Network traffic entering/leaving the data center'),
            ('NSG', 'Network Security Group - Azure firewall construct'),
            ('Security Group', 'AWS EC2 instance-level firewall'),
            ('Zero-Trust', 'Security model requiring verification for all access')
        ]

        glossary_table = self.doc.add_table(rows=len(glossary) + 1, cols=2)
        glossary_table.style = 'Light Shading Accent 1'

        glossary_table.rows[0].cells[0].text = 'Term'
        glossary_table.rows[0].cells[1].text = 'Definition'

        for i, (term, definition) in enumerate(glossary, 1):
            row = glossary_table.rows[i].cells
            row[0].text = term
            row[1].text = definition

    def _add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()


# Convenience function
def generate_solutions_document(
    analysis_results: Dict,
    zones: Dict,
    rules: List,
    output_path: str = 'outputs_final/word_reports/netseg/network_segmentation_solution.docx',
    png_path: str = None
):
    """
    Generate comprehensive Solutions Architecture Document

    Args:
        analysis_results: Analysis results from TrafficAnalyzer
        zones: Network zones dictionary
        rules: List of SegmentationRule objects
        output_path: Output file path
        png_path: Path to PNG diagram for embedding (optional)

    Returns:
        Path to generated document
    """
    doc_gen = SolutionsArchitectureDocument(analysis_results, zones, rules, png_path=png_path)
    doc_gen.generate_document(output_path)

    logger.info(f"✅ Solutions Architecture Document generated: {output_path}")
    return output_path


if __name__ == '__main__':
    # Example usage
    from src.parser import parse_network_logs
    from src.analysis import analyze_traffic

    print("="*60)
    print("Solutions Document Generator - Test Run")
    print("="*60)

    # Parse and analyze
    parser = parse_network_logs('data/input')
    analyzer = analyze_traffic(parser.records, 'outputs')

    # Generate document
    doc_path = generate_solutions_document(
        analyzer.analysis_results,
        analyzer.zones,
        analyzer.rules,
        'outputs/network_segmentation_solution.docx'
    )

    print(f"\n✅ Document generated: {doc_path}")

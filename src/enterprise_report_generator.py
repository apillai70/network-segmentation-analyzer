#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Enterprise Network Analysis Report Generator
==============================================
Generates comprehensive enterprise-wide network analysis reports

Consolidates data from all applications to provide:
- Enterprise-wide statistics
- Cross-application dependencies
- Security zone analysis
- DNS validation summaries
- Network segmentation recommendations

Author: Enterprise Security Team
Version: 1.0
"""

import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Set, Tuple
from collections import defaultdict, Counter
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class EnterpriseNetworkReportGenerator:
    """
    Generate enterprise-wide network analysis reports

    Consolidates topology data from all applications to provide
    comprehensive network security and segmentation analysis
    """

    def __init__(self, topology_dir: str = './persistent_data/topology',
                 output_dir: str = './results'):
        """
        Initialize enterprise report generator

        Args:
            topology_dir: Directory containing application topology JSON files
            output_dir: Output directory for reports
        """
        self.topology_dir = Path(topology_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Consolidated data
        self.applications: List[Dict] = []
        self.all_dependencies: List[Dict] = []
        self.security_zones: Dict[str, List[str]] = defaultdict(list)
        self.cross_zone_flows: List[Tuple[str, str, str]] = []  # (app, src_zone, dst_zone)

        # Statistics
        self.stats = {
            'total_applications': 0,
            'total_dependencies': 0,
            'total_unique_ips': 0,
            'apps_by_zone': {},
            'cross_zone_connections': 0,
            'dns_validation_summary': {},
            'timestamp': datetime.now().isoformat()
        }

        logger.info("Enterprise Network Report Generator initialized")
        logger.info(f"  Topology directory: {self.topology_dir}")
        logger.info(f"  Output directory: {self.output_dir}")

    def load_all_topology_data(self):
        """Load topology data from all application JSON files"""
        logger.info(f"\nLoading topology data from: {self.topology_dir}")

        # Find all topology JSON files
        json_files = list(self.topology_dir.glob('*.json'))
        logger.info(f"Found {len(json_files)} application topology files")

        if not json_files:
            logger.warning("No topology files found!")
            return

        # Load each application
        for json_file in json_files:
            try:
                with open(json_file, 'r') as f:
                    topology_data = json.load(f)

                app_id = topology_data.get('app_id', json_file.stem)
                security_zone = topology_data.get('security_zone', 'UNKNOWN')
                dependencies = topology_data.get('dependencies', [])

                # Store application data
                self.applications.append({
                    'app_id': app_id,
                    'security_zone': security_zone,
                    'confidence': topology_data.get('confidence', 0.0),
                    'num_dependencies': len(dependencies),
                    'dns_validation': topology_data.get('dns_validation', {}),
                    'validation_metadata': topology_data.get('validation_metadata', {}),
                    'characteristics': topology_data.get('characteristics', []),
                    'timestamp': topology_data.get('timestamp')
                })

                # Track security zones
                self.security_zones[security_zone].append(app_id)

                # Store dependencies
                for dep in dependencies:
                    dep['source_app'] = app_id
                    dep['source_zone'] = security_zone
                    self.all_dependencies.append(dep)

                logger.debug(f"  Loaded: {app_id} ({security_zone})")

            except Exception as e:
                logger.warning(f"  Failed to load {json_file.name}: {e}")

        logger.info(f"✓ Loaded {len(self.applications)} applications")

    def analyze_network_topology(self):
        """Analyze enterprise network topology"""
        logger.info("\nAnalyzing enterprise network topology...")

        # Basic statistics
        self.stats['total_applications'] = len(self.applications)
        self.stats['total_dependencies'] = len(self.all_dependencies)

        # Count apps by security zone
        for zone, apps in self.security_zones.items():
            self.stats['apps_by_zone'][zone] = len(apps)

        # Analyze cross-zone connections
        cross_zone_count = 0
        for dep in self.all_dependencies:
            source_zone = dep.get('source_zone', 'UNKNOWN')
            # Dependencies might not have explicit zones, infer from component types
            if source_zone != 'UNKNOWN':
                cross_zone_count += 1
                # Track cross-zone flows
                self.cross_zone_flows.append((
                    dep.get('source_app'),
                    source_zone,
                    dep.get('component_type', 'UNKNOWN')
                ))

        self.stats['cross_zone_connections'] = cross_zone_count

        # Count unique IPs from all dependencies
        unique_ips = set()
        for dep in self.all_dependencies:
            components = dep.get('components', [])
            for comp in components:
                if 'ip' in comp:
                    unique_ips.add(comp['ip'])
                elif 'host' in comp:
                    unique_ips.add(comp['host'])

        self.stats['total_unique_ips'] = len(unique_ips)

        # Consolidate DNS validation statistics
        dns_stats = {
            'total_validated': 0,
            'total_valid': 0,
            'total_valid_multiple': 0,
            'total_mismatches': 0,
            'total_nxdomain': 0,
            'total_failed': 0,
            'apps_with_validation': 0
        }

        for app in self.applications:
            dns_validation = app.get('dns_validation', {})
            if dns_validation and dns_validation.get('total_validated', 0) > 0:
                dns_stats['apps_with_validation'] += 1
                dns_stats['total_validated'] += dns_validation.get('total_validated', 0)
                dns_stats['total_valid'] += dns_validation.get('valid', 0)
                dns_stats['total_valid_multiple'] += dns_validation.get('valid_multiple_ips', 0)
                dns_stats['total_mismatches'] += dns_validation.get('mismatch', 0)
                dns_stats['total_nxdomain'] += dns_validation.get('nxdomain', 0)
                dns_stats['total_failed'] += dns_validation.get('failed', 0)

        self.stats['dns_validation_summary'] = dns_stats

        logger.info(f"✓ Analysis complete")
        logger.info(f"  Applications: {self.stats['total_applications']}")
        logger.info(f"  Dependencies: {self.stats['total_dependencies']}")
        logger.info(f"  Unique IPs: {self.stats['total_unique_ips']}")
        logger.info(f"  Security Zones: {len(self.security_zones)}")

    def generate_json_report(self, output_path: str = None) -> str:
        """
        Generate comprehensive JSON report

        Args:
            output_path: Optional custom output path

        Returns:
            Path to generated report
        """
        if not output_path:
            output_path = self.output_dir / 'enterprise_network_analysis_report.json'

        logger.info(f"\nGenerating JSON report: {output_path}")

        # Build comprehensive report
        report = {
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'total_applications': self.stats['total_applications'],
                'total_dependencies': self.stats['total_dependencies'],
                'report_version': '1.0'
            },
            'statistics': self.stats,
            'security_zones': {
                zone: {
                    'count': len(apps),
                    'applications': apps
                }
                for zone, apps in self.security_zones.items()
            },
            'applications': self.applications,
            'dependencies_summary': {
                'total': len(self.all_dependencies),
                'by_type': self._summarize_dependencies_by_type(),
                'cross_zone_flows': len(self.cross_zone_flows)
            },
            'dns_validation': self.stats['dns_validation_summary'],
            'top_applications': self._get_top_applications(),
            'recommendations': self._generate_recommendations()
        }

        # Save report
        with open(output_path, 'w') as f:
            json.dump(report, f, indent=2, default=str)

        logger.info(f"✓ JSON report saved: {output_path}")
        return str(output_path)

    def _summarize_dependencies_by_type(self) -> Dict:
        """Summarize dependencies by component type"""
        type_counts = Counter()

        for dep in self.all_dependencies:
            comp_type = dep.get('component_type', 'unknown')
            type_counts[comp_type] += 1

        return dict(type_counts)

    def _get_top_applications(self, limit: int = 10) -> List[Dict]:
        """Get top applications by number of dependencies"""
        # Sort by number of dependencies
        sorted_apps = sorted(
            self.applications,
            key=lambda x: x['num_dependencies'],
            reverse=True
        )

        return [
            {
                'app_id': app['app_id'],
                'security_zone': app['security_zone'],
                'num_dependencies': app['num_dependencies'],
                'confidence': app['confidence']
            }
            for app in sorted_apps[:limit]
        ]

    def _generate_recommendations(self) -> List[Dict]:
        """Generate security recommendations based on analysis"""
        recommendations = []

        # Zone distribution recommendation
        zone_counts = self.stats['apps_by_zone']
        if zone_counts:
            largest_zone = max(zone_counts, key=zone_counts.get)
            if zone_counts[largest_zone] > len(self.applications) * 0.5:
                recommendations.append({
                    'priority': 'HIGH',
                    'category': 'Zone Distribution',
                    'issue': f'{zone_counts[largest_zone]} applications ({zone_counts[largest_zone]/len(self.applications)*100:.1f}%) are in {largest_zone}',
                    'recommendation': 'Consider further segmentation to reduce blast radius and improve security isolation'
                })

        # DNS validation recommendation
        dns_stats = self.stats['dns_validation_summary']
        if dns_stats.get('total_mismatches', 0) > 0:
            recommendations.append({
                'priority': 'MEDIUM',
                'category': 'DNS Configuration',
                'issue': f'{dns_stats["total_mismatches"]} DNS mismatches detected across enterprise',
                'recommendation': 'Review and correct DNS records where forward and reverse DNS do not match'
            })

        # NXDOMAIN recommendation
        if dns_stats.get('total_nxdomain', 0) > 10:
            recommendations.append({
                'priority': 'LOW',
                'category': 'DNS Coverage',
                'issue': f'{dns_stats["total_nxdomain"]} IP addresses have no reverse DNS records',
                'recommendation': 'Add DNS records for active systems to improve network visibility'
            })

        # Cross-zone flows recommendation
        if self.stats['cross_zone_connections'] > 100:
            recommendations.append({
                'priority': 'HIGH',
                'category': 'Network Segmentation',
                'issue': f'{self.stats["cross_zone_connections"]} cross-zone connections detected',
                'recommendation': 'Review cross-zone flows and implement strict firewall policies between security zones'
            })

        return recommendations

    def generate_word_report(self, output_path: str = None) -> str:
        """
        Generate comprehensive Word document report

        Args:
            output_path: Optional custom output path

        Returns:
            Path to generated report
        """
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.output_dir / f'Enterprise_Network_Analysis_{timestamp}.docx'

        logger.info(f"\nGenerating Word report: {output_path}")

        doc = Document()

        # Title
        title = doc.add_heading('Enterprise Network Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(
            f'Comprehensive Network Segmentation Analysis\n'
            f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Executive Summary
        self._add_word_executive_summary(doc)

        # Network Statistics
        self._add_word_statistics(doc)

        # Security Zones
        self._add_word_security_zones(doc)

        # DNS Validation Summary
        self._add_word_dns_summary(doc)

        # Top Applications
        self._add_word_top_applications(doc)

        # Recommendations
        self._add_word_recommendations(doc)

        # Save document
        doc.save(output_path)
        logger.info(f"✓ Word report saved: {output_path}")

        return str(output_path)

    def _add_word_executive_summary(self, doc):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"This report provides a comprehensive analysis of the enterprise network infrastructure "
            f"covering {self.stats['total_applications']} applications. "
        )
        para.add_run(
            f"The analysis includes network segmentation assessment, dependency mapping, "
            f"DNS validation findings, and security recommendations."
        )

        doc.add_paragraph()

        # Key findings
        findings_para = doc.add_paragraph()
        findings_para.add_run("Key Findings:\n").bold = True

        findings_para.add_run(
            f"• {self.stats['total_applications']} applications analyzed across "
            f"{len(self.security_zones)} security zones\n"
        )

        findings_para.add_run(
            f"• {self.stats['total_dependencies']} dependencies identified across "
            f"{self.stats['total_unique_ips']} unique IP addresses\n"
        )

        dns_stats = self.stats['dns_validation_summary']
        if dns_stats.get('apps_with_validation', 0) > 0:
            findings_para.add_run(
                f"• DNS validation performed on {dns_stats['apps_with_validation']} applications, "
                f"validating {dns_stats['total_validated']} IP addresses\n"
            )

        if self.stats['cross_zone_connections'] > 0:
            findings_para.add_run(
                f"• {self.stats['cross_zone_connections']} cross-zone network connections require policy review\n"
            )

        doc.add_paragraph()

    def _add_word_statistics(self, doc):
        """Add statistics table"""
        doc.add_heading('Network Statistics', 1)

        # Create table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        rows_data = [
            ('Total Applications', self.stats['total_applications']),
            ('Total Dependencies', self.stats['total_dependencies']),
            ('Unique IP Addresses', self.stats['total_unique_ips']),
            ('Security Zones', len(self.security_zones)),
            ('Cross-Zone Connections', self.stats['cross_zone_connections']),
        ]

        for idx, (label, value) in enumerate(rows_data, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = label
            row_cells[1].text = str(value)

        doc.add_paragraph()

    def _add_word_security_zones(self, doc):
        """Add security zones section"""
        doc.add_heading('Security Zones', 1)

        para = doc.add_paragraph()
        para.add_run(
            f"Applications are distributed across {len(self.security_zones)} security zones:"
        )

        doc.add_paragraph()

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Security Zone'
        header_cells[1].text = 'Applications'
        header_cells[2].text = 'Percentage'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        total_apps = self.stats['total_applications']
        for zone in sorted(self.security_zones.keys()):
            count = len(self.security_zones[zone])
            percentage = (count / total_apps * 100) if total_apps > 0 else 0

            row_cells = table.add_row().cells
            row_cells[0].text = zone
            row_cells[1].text = str(count)
            row_cells[2].text = f"{percentage:.1f}%"

        doc.add_paragraph()

    def _add_word_dns_summary(self, doc):
        """Add DNS validation summary"""
        doc.add_heading('DNS Validation Summary', 1)

        dns_stats = self.stats['dns_validation_summary']

        if dns_stats.get('apps_with_validation', 0) == 0:
            para = doc.add_paragraph()
            para.add_run("DNS validation data not available. ")
            para.add_run("Run batch processing with DNS validation enabled to collect data.").italic = True
            doc.add_paragraph()
            return

        para = doc.add_paragraph()
        para.add_run(
            f"DNS validation was performed on {dns_stats['apps_with_validation']} applications, "
            f"validating {dns_stats['total_validated']} unique IP addresses."
        )

        doc.add_paragraph()

        # Create table
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Validation Status'
        header_cells[1].text = 'Count'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        rows_data = [
            ('Applications with Validation', dns_stats.get('apps_with_validation', 0)),
            ('Total IPs Validated', dns_stats.get('total_validated', 0)),
            ('Valid DNS (Perfect Match)', dns_stats.get('total_valid', 0)),
            ('Valid DNS (Multiple IPs)', dns_stats.get('total_valid_multiple', 0)),
            ('DNS Mismatches', dns_stats.get('total_mismatches', 0)),
            ('NXDOMAIN', dns_stats.get('total_nxdomain', 0)),
        ]

        for idx, (label, value) in enumerate(rows_data, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = label
            row_cells[1].text = str(value)

        doc.add_paragraph()

    def _add_word_top_applications(self, doc):
        """Add top applications section"""
        doc.add_heading('Top 10 Applications by Dependencies', 1)

        top_apps = self._get_top_applications(limit=10)

        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Rank'
        header_cells[1].text = 'Application ID'
        header_cells[2].text = 'Security Zone'
        header_cells[3].text = 'Dependencies'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for idx, app in enumerate(top_apps, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = app['app_id']
            row_cells[2].text = app['security_zone']
            row_cells[3].text = str(app['num_dependencies'])

        doc.add_paragraph()

    def _add_word_recommendations(self, doc):
        """Add recommendations section"""
        doc.add_heading('Recommendations', 1)

        recommendations = self._generate_recommendations()

        if not recommendations:
            para = doc.add_paragraph()
            para.add_run("No critical recommendations at this time. Network segmentation appears healthy.")
            doc.add_paragraph()
            return

        for idx, rec in enumerate(recommendations, start=1):
            # Recommendation heading
            rec_heading = doc.add_paragraph(style='List Number')
            rec_heading.add_run(f"{rec['category']} - ").bold = True
            rec_heading.add_run(f"[{rec['priority']}]").font.color.rgb = RGBColor(255, 0, 0) if rec['priority'] == 'HIGH' else RGBColor(255, 165, 0)

            # Issue
            issue_para = doc.add_paragraph(style='List Bullet 2')
            issue_para.add_run("Issue: ").bold = True
            issue_para.add_run(rec['issue'])

            # Recommendation
            rec_para = doc.add_paragraph(style='List Bullet 2')
            rec_para.add_run("Recommendation: ").bold = True
            rec_para.add_run(rec['recommendation'])

        doc.add_paragraph()

    def generate_all_reports(self) -> Dict[str, str]:
        """
        Generate all report formats

        Returns:
            Dictionary of report type -> file path
        """
        logger.info("\n" + "="*80)
        logger.info("GENERATING ENTERPRISE NETWORK ANALYSIS REPORTS")
        logger.info("="*80)

        reports = {}

        # JSON report
        reports['json'] = self.generate_json_report()

        # Word report
        reports['word'] = self.generate_word_report()

        logger.info("\n" + "="*80)
        logger.info("ENTERPRISE REPORTS COMPLETE")
        logger.info("="*80)
        logger.info(f"  JSON Report: {reports['json']}")
        logger.info(f"  Word Report: {reports['word']}")
        logger.info("="*80 + "\n")

        return reports

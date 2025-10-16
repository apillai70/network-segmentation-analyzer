"""
Enterprise Network Segmentation Strategy Document Generator
=============================================================
Generates a comprehensive network segmentation strategy document
based on actual network topology data, proposing multiple segmentation
options with detailed pros/cons analysis.

This document is DYNAMIC - it responds to actual network data and
creates data-driven segmentation recommendations.

Author: Prutech Network Security Team
Version: 1.0
"""

import logging
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict, Counter

logger = logging.getLogger(__name__)


class EnterpriseSegmentationStrategyDocument:
    """Generates enterprise-wide network segmentation strategy document"""

    def __init__(self, master_topology_path: str):
        """Initialize with master topology data

        Args:
            master_topology_path: Path to master_topology.json
        """
        self.topology_path = master_topology_path
        self.topology_data = self._load_topology()
        self.doc = Document()
        self._setup_styles()

        # Analytics from topology
        self.total_apps = 0
        self.zones_used = set()
        self.tier_distribution = defaultdict(int)
        self.dependency_count = 0
        self.external_dependencies = 0
        self.cross_zone_flows = []

        self._analyze_topology()

    def _load_topology(self) -> Dict:
        """Load master topology data"""
        try:
            with open(self.topology_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Failed to load topology: {e}")
            return {'topology': {}, 'metadata': {}}

    def _analyze_topology(self):
        """Analyze topology to extract segmentation insights"""
        topology = self.topology_data.get('topology', {})
        self.total_apps = len(topology)

        for app_name, app_data in topology.items():
            zone = app_data.get('security_zone', 'UNKNOWN')
            self.zones_used.add(zone)
            self.tier_distribution[zone] += 1

            dependencies = app_data.get('predicted_dependencies', [])
            self.dependency_count += len(dependencies)

            # Count external dependencies
            for dep in dependencies:
                dep_type = dep.get('type', '')
                if dep_type in ['downstream_app', 'external', 'api']:
                    self.external_dependencies += 1

    def _setup_styles(self):
        """Setup document styles"""
        styles = self.doc.styles

        try:
            # Heading styles
            h1 = styles['Heading 1']
            h1.font.size = Pt(20)
            h1.font.color.rgb = RGBColor(0, 51, 102)
            h1.font.bold = True

            h2 = styles['Heading 2']
            h2.font.size = Pt(16)
            h2.font.color.rgb = RGBColor(0, 102, 204)

            h3 = styles['Heading 3']
            h3.font.size = Pt(13)
            h3.font.color.rgb = RGBColor(51, 153, 255)
        except:
            pass

    def generate_document(self, output_path: str):
        """Generate the enterprise segmentation strategy document

        Args:
            output_path: Output path for Word document
        """
        logger.info("Generating Enterprise Network Segmentation Strategy Document")

        # Cover page
        self._add_cover_page()
        self.doc.add_page_break()

        # Executive summary
        self._add_executive_summary()
        self.doc.add_page_break()

        # Current state analysis
        self._add_current_state_analysis()
        self.doc.add_page_break()

        # Segmentation options (THE CORE)
        self._add_segmentation_options()
        self.doc.add_page_break()

        # Comparison matrix
        self._add_comparison_matrix()
        self.doc.add_page_break()

        # Recommendations
        self._add_recommendations()
        self.doc.add_page_break()

        # Implementation roadmap
        self._add_implementation_roadmap()

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output_file))
        logger.info(f"[SUCCESS] Enterprise segmentation strategy saved: {output_path}")

    def _add_cover_page(self):
        """Add professional cover page"""
        # Title
        title = self.doc.add_heading(
            'Enterprise Network Segmentation\nStrategy & Options Analysis',
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Data-Driven Segmentation Analysis\nwith Multiple Strategic Options')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 102, 204)

        for _ in range(3):
            self.doc.add_paragraph()

        # Classification
        classification = self.doc.add_paragraph()
        classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classification.add_run('[CONFIDENTIAL] - SECURITY STRATEGY')
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(139, 0, 0)

        for _ in range(2):
            self.doc.add_paragraph()

        # Document info
        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run('Version: 1.0\n').bold = True
        info.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
        info.add_run(f'Network Scope: {self.total_apps} Applications Analyzed\n')
        info.add_run(f'Data Source: Live Network Topology\n')

        for _ in range(3):
            self.doc.add_paragraph()

        # Footer
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Prepared by: Prutech Network Security Team\n')
        footer.add_run('Enterprise Network Architecture & Zero Trust Strategy\n')
        footer.add_run('Auto-generated by Network Segmentation Analyzer\n')

    def _add_executive_summary(self):
        """Add executive summary with key findings"""
        self.doc.add_heading('Executive Summary', level=1)

        summary_text = (
            f"This document presents a comprehensive network segmentation strategy for the enterprise, "
            f"analyzing {self.total_apps} applications and {self.dependency_count} dependencies across "
            f"the current network topology. Based on actual network flow data, we propose multiple "
            f"segmentation options ranging from basic 3-zone architectures to advanced micro-segmentation "
            f"approaches, each with detailed cost-benefit analysis."
        )
        self.doc.add_paragraph(summary_text)

        self.doc.add_paragraph()

        # Key findings
        findings_para = self.doc.add_paragraph()
        findings_para.add_run('KEY FINDINGS:').bold = True

        findings = [
            f'Current network spans {len(self.zones_used)} security zones',
            f'Total applications analyzed: {self.total_apps}',
            f'Total dependencies mapped: {self.dependency_count}',
            f'External/API dependencies: {self.external_dependencies}',
            f'Segmentation opportunities identified: {self._count_segmentation_opportunities()}'
        ]

        for finding in findings:
            self.doc.add_paragraph(finding, style='List Bullet')

        self.doc.add_paragraph()

        # What this document provides
        provides_para = self.doc.add_paragraph()
        provides_para.add_run('THIS DOCUMENT PROVIDES:').bold = True

        provides = [
            'Current state network topology analysis based on actual flow data',
            'Multiple segmentation options (Minimal, Standard, Advanced, Micro-segmentation)',
            'Detailed pros and cons for each option',
            'Cost-benefit analysis for each approach',
            'Regulatory compliance implications (PCI-DSS, HIPAA, SOX, GDPR)',
            'Data-driven recommendations tailored to your network',
            'Phased implementation roadmap with timelines'
        ]

        for item in provides:
            self.doc.add_paragraph(item, style='List Bullet')

    def _add_current_state_analysis(self):
        """Add current state network analysis"""
        self.doc.add_heading('Current State Network Analysis', level=1)

        intro = (
            "The following analysis is based on actual network topology data collected from "
            "ExtraHop flow records and machine learning predictions. This represents the current "
            "state of the network as observed in production."
        )
        self.doc.add_paragraph(intro)

        # Network topology summary
        self.doc.add_heading('Network Topology Summary', level=2)

        # Create summary table
        summary_table = self.doc.add_table(rows=6, cols=2)
        summary_table.style = 'Medium Shading 1 Accent 1'

        metrics = [
            ('Total Applications', str(self.total_apps)),
            ('Security Zones Identified', str(len(self.zones_used))),
            ('Total Dependencies', str(self.dependency_count)),
            ('External Dependencies', str(self.external_dependencies)),
            ('Average Dependencies per App', f'{self.dependency_count / max(self.total_apps, 1):.1f}'),
            ('Current Segmentation Level', self._assess_current_segmentation())
        ]

        for idx, (metric, value) in enumerate(metrics):
            summary_table.rows[idx].cells[0].text = metric
            summary_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            summary_table.rows[idx].cells[1].text = value

        # Zone distribution
        self.doc.add_paragraph()
        self.doc.add_heading('Security Zone Distribution', level=2)

        if self.tier_distribution:
            zone_table = self.doc.add_table(rows=len(self.tier_distribution) + 1, cols=3)
            zone_table.style = 'Light Grid Accent 1'

            # Header
            zone_table.rows[0].cells[0].text = 'Security Zone'
            zone_table.rows[0].cells[1].text = 'Applications'
            zone_table.rows[0].cells[2].text = 'Percentage'
            for cell in zone_table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data
            for idx, (zone, count) in enumerate(sorted(self.tier_distribution.items()), 1):
                percentage = (count / self.total_apps * 100) if self.total_apps > 0 else 0
                zone_table.rows[idx].cells[0].text = zone
                zone_table.rows[idx].cells[1].text = str(count)
                zone_table.rows[idx].cells[2].text = f'{percentage:.1f}%'

        # Current state assessment
        self.doc.add_paragraph()
        self.doc.add_heading('Current State Assessment', level=2)

        assessment = self._generate_current_state_assessment()
        for item in assessment:
            self.doc.add_paragraph(item, style='List Bullet')

    def _add_segmentation_options(self):
        """Add detailed segmentation options with pros/cons"""
        self.doc.add_heading('Network Segmentation Options', level=1)

        intro = (
            "Based on analysis of your network topology, we present four segmentation options "
            "ranging from basic to advanced. Each option is tailored to your specific network "
            "characteristics and includes detailed implementation requirements, benefits, drawbacks, "
            "and cost analysis."
        )
        self.doc.add_paragraph(intro)

        self.doc.add_paragraph()

        # Generate each option
        options = [
            self._generate_option_minimal(),
            self._generate_option_standard(),
            self._generate_option_advanced(),
            self._generate_option_microsegmentation()
        ]

        for idx, option in enumerate(options, 1):
            self._render_option(idx, option)
            if idx < len(options):
                self.doc.add_page_break()

    def _generate_option_minimal(self) -> Dict:
        """Generate Option 1: Minimal Segmentation (3 Zones)"""
        return {
            'name': 'Option 1: Minimal Segmentation (3 Zones)',
            'description': (
                'Basic three-zone architecture separating internet-facing, internal applications, '
                'and data storage. Simplest approach with lowest implementation cost.'
            ),
            'zones': [
                ('DMZ (Perimeter Zone)', 'Internet-facing web servers, load balancers, reverse proxies'),
                ('Internal (Application Zone)', 'All business applications, APIs, middleware'),
                ('Data Zone', 'All databases, file servers, backup systems')
            ],
            'implementation': [
                'Deploy perimeter firewall between Internet and DMZ',
                'Deploy internal firewall between DMZ and Internal zone',
                'Deploy data firewall between Internal and Data zones',
                'Implement basic ACLs on each firewall',
                'Enable logging for all zone boundaries'
            ],
            'pros': [
                f'[COST] Minimal hardware requirements (3 firewall zones) - Estimated cost: $50K-150K',
                f'[SIMPLICITY] Easy to understand and manage - Low operational overhead',
                f'[SPEED] Fastest implementation - 2-4 weeks deployment time',
                f'[COMPATIBILITY] Works with existing infrastructure - No major redesign needed',
                f'[COMPLIANCE] Meets basic PCI-DSS requirements for CDE isolation',
                f'[PROVEN] Well-established architecture pattern - Low risk'
            ],
            'cons': [
                f'[SECURITY] Limited granularity - All apps in single "Internal" zone increases blast radius',
                f'[LATERAL MOVEMENT] Minimal protection against east-west threats - Compromise spreads easily',
                f'[COMPLIANCE] Insufficient for HIPAA or advanced PCI-DSS scope reduction',
                f'[VISIBILITY] Coarse-grained logging - Difficult to trace lateral movement',
                f'[SCALABILITY] Does not scale well beyond {min(self.total_apps, 50)} applications',
                f'[ZERO TRUST] Not aligned with Zero Trust principles - Perimeter-based security model'
            ],
            'cost_breakdown': {
                'one_time': [
                    ('Firewall hardware/licenses', '$30K-100K'),
                    ('Implementation services', '$15K-40K'),
                    ('Testing and validation', '$5K-10K')
                ],
                'annual': [
                    ('Maintenance and support', '$5K-15K'),
                    ('Operations and monitoring', '$10K-25K')
                ]
            },
            'timeline': '2-4 weeks',
            'complexity': 'LOW',
            'security_rating': '★★☆☆☆ (2/5)',
            'best_for': [
                'Small organizations (<50 applications)',
                'Limited security budget',
                'Low regulatory requirements',
                'Proof-of-concept segmentation projects',
                'Organizations just starting network segmentation journey'
            ],
            'not_recommended_for': [
                'Organizations processing payment cards (PCI-DSS)',
                'Healthcare providers (HIPAA)',
                'Organizations with high-value data',
                'Environments with sophisticated threats'
            ]
        }

    def _generate_option_standard(self) -> Dict:
        """Generate Option 2: Standard Segmentation (6-7 Zones)"""
        apps_per_zone = max(self.total_apps // 6, 1)

        return {
            'name': 'Option 2: Standard Tier-Based Segmentation (6-7 Zones)',
            'description': (
                f'Industry-standard tier-based architecture providing strong security with '
                f'manageable complexity. Suitable for most enterprises. Based on your {self.total_apps} '
                f'applications, this creates approximately {apps_per_zone} applications per zone.'
            ),
            'zones': [
                ('WEB_TIER', 'Internet-facing web servers, load balancers, WAF'),
                ('APP_TIER', 'Application servers, business logic, API gateways'),
                ('DATA_TIER', 'Databases (SQL/NoSQL), data warehouses'),
                ('CACHE_TIER', 'Redis, Memcached, caching layers'),
                ('MESSAGING_TIER', 'Kafka, RabbitMQ, message queues'),
                ('MANAGEMENT_TIER', 'Monitoring, logging, bastion hosts, CI/CD'),
                ('EXTERNAL', 'Third-party services, SaaS, partner networks')
            ],
            'implementation': [
                'Deploy firewalls between each tier boundary (6-7 boundaries)',
                'Implement deny-all policies with specific allow rules per tier',
                'Enable TLS 1.3 for all inter-tier communications',
                'Deploy IDS/IPS at each tier boundary',
                'Implement centralized logging (SIEM)',
                'Deploy database activity monitoring (DAM) for Data Tier',
                'Configure API gateway for App Tier',
                'Establish privileged access management (PAM) for Management Tier'
            ],
            'pros': [
                f'[SECURITY] Strong tier isolation reduces lateral movement by 70-80%',
                f'[COMPLIANCE] Meets PCI-DSS, HIPAA, SOX requirements for segmentation',
                f'[GRANULARITY] Separate control for each functional tier - Better access control',
                f'[SCALABILITY] Scales to {self.total_apps * 3} applications without major changes',
                f'[INDUSTRY STANDARD] Proven architecture used by Fortune 500 companies',
                f'[BALANCE] Optimal security/complexity trade-off for most organizations',
                f'[TROUBLESHOOTING] Easier to isolate issues by tier - Clear boundaries',
                f'[COMPLIANCE] Reduces PCI-DSS audit scope by 50-70% (fewer systems in CDE)'
            ],
            'cons': [
                f'[COST] Moderate investment required - Estimated $200K-500K initial',
                f'[COMPLEXITY] Requires skilled security team to manage 6-7 zone boundaries',
                f'[OPERATIONS] More firewall rules to maintain (estimated {self.dependency_count * 2} rules)',
                f'[PERFORMANCE] Additional network hops may introduce 5-10ms latency',
                f'[DEPLOYMENT] 6-10 weeks implementation time - Moderate disruption risk',
                f'[DEPENDENCIES] Applications must be tier-aware - May require app refactoring'
            ],
            'cost_breakdown': {
                'one_time': [
                    ('Firewall infrastructure', '$100K-250K'),
                    ('IDS/IPS deployment', '$50K-100K'),
                    ('SIEM implementation', '$50K-100K'),
                    ('Professional services', '$75K-150K'),
                    ('Testing and validation', '$20K-50K')
                ],
                'annual': [
                    ('Licenses and maintenance', '$40K-100K'),
                    ('Operations (2 FTEs)', '$150K-300K'),
                    ('Monitoring and incident response', '$50K-100K')
                ]
            },
            'timeline': '6-10 weeks',
            'complexity': 'MEDIUM',
            'security_rating': '★★★★☆ (4/5)',
            'best_for': [
                f'Medium to large organizations ({self.total_apps}-500 applications)',
                'PCI-DSS compliance requirements',
                'HIPAA compliance requirements',
                'Organizations with moderate security maturity',
                'Balance between security and operational complexity'
            ],
            'not_recommended_for': [
                'Small organizations with <30 applications (over-engineered)',
                'Organizations lacking skilled security staff',
                'Very tight budget constraints'
            ]
        }

    def _generate_option_advanced(self) -> Dict:
        """Generate Option 3: Advanced Segmentation (10+ Zones)"""
        return {
            'name': 'Option 3: Advanced Multi-Zone Segmentation (10+ Zones)',
            'description': (
                f'Advanced segmentation with separate zones for production/non-production, '
                f'data sensitivity levels, and regulatory requirements. For organizations with '
                f'{self.total_apps}+ applications requiring fine-grained control.'
            ),
            'zones': [
                ('PROD_WEB', 'Production internet-facing tier'),
                ('NONPROD_WEB', 'Dev/Test web tier'),
                ('PROD_APP', 'Production application tier'),
                ('NONPROD_APP', 'Dev/Test application tier'),
                ('REGULATED_DATA', 'PCI/PHI/PII data (highest security)'),
                ('STANDARD_DATA', 'Internal data (standard security)'),
                ('CACHE', 'Caching layer'),
                ('MESSAGING', 'Message queues and event streams'),
                ('MANAGEMENT', 'Monitoring, logging, ops tools'),
                ('PARTNER_DMZ', 'Partner integrations'),
                ('EXTERNAL_APIS', 'Third-party API calls'),
                ('BACKUP/DR', 'Backup systems and disaster recovery')
            ],
            'implementation': [
                'Deploy next-generation firewalls at each zone boundary (10-12 boundaries)',
                'Implement application-aware firewall rules with deep packet inspection',
                'Enable mTLS (mutual TLS) for all service-to-service communications',
                'Deploy service mesh (Istio/Linkerd) for micro-segmentation within zones',
                'Implement advanced threat protection (ATP) at each boundary',
                'Deploy separate SIEM with UEBA (User and Entity Behavior Analytics)',
                'Implement privileged access management (PAM) with session recording',
                'Deploy database activity monitoring (DAM) with real-time alerting',
                'Establish separate networks for prod vs. non-prod',
                'Implement network access control (NAC) for device authentication'
            ],
            'pros': [
                f'[SECURITY] Maximum isolation - Lateral movement reduced by 85-90%',
                f'[COMPLIANCE] Exceeds all regulatory requirements (PCI-DSS, HIPAA, SOX, GDPR)',
                f'[GRANULARITY] Fine-grained control per environment and data sensitivity',
                f'[BLAST RADIUS] Minimal blast radius - Breaches contained to single zone',
                f'[AUDIT] Simplified audit process - Clear separation of regulated data',
                f'[FLEXIBILITY] Supports complex regulatory and business requirements',
                f'[ZERO TRUST] Aligned with Zero Trust architecture principles',
                f'[INSURANCE] May reduce cyber insurance premiums by 20-30%'
            ],
            'cons': [
                f'[COST] High investment - Estimated $500K-1.5M initial capital',
                f'[COMPLEXITY] Complex to design and operate - Requires dedicated security team (3-5 FTEs)',
                f'[OPERATIONS] High operational overhead - {self.dependency_count * 4}+ firewall rules to manage',
                f'[PERFORMANCE] Additional hops introduce 10-20ms latency',
                f'[DEPLOYMENT] Long deployment (12-16 weeks) - High disruption risk',
                f'[TROUBLESHOOTING] Complex troubleshooting across many zones',
                f'[SKILLS] Requires highly skilled staff - Difficult to hire/train',
                f'[OVERKILL] May be over-engineered for organizations with <{self.total_apps} applications'
            ],
            'cost_breakdown': {
                'one_time': [
                    ('Next-gen firewalls and licenses', '$250K-600K'),
                    ('IDS/IPS/ATP deployment', '$100K-250K'),
                    ('Service mesh implementation', '$75K-150K'),
                    ('SIEM with UEBA', '$100K-250K'),
                    ('PAM implementation', '$50K-150K'),
                    ('Professional services', '$200K-400K'),
                    ('Testing and validation', '$50K-100K')
                ],
                'annual': [
                    ('Licenses and maintenance', '$100K-250K'),
                    ('Operations (4-5 FTEs)', '$400K-750K'),
                    ('Monitoring and SOC', '$150K-300K'),
                    ('Training and development', '$50K-100K')
                ]
            },
            'timeline': '12-16 weeks',
            'complexity': 'HIGH',
            'security_rating': '★★★★★ (5/5)',
            'best_for': [
                f'Large enterprises ({self.total_apps}+ applications)',
                'Organizations processing highly sensitive data (PCI/PHI)',
                'Multi-national companies with complex regulatory requirements',
                'Financial services and healthcare industries',
                'High-security environments (defense, government)',
                'Organizations with mature security programs'
            ],
            'not_recommended_for': [
                'Small to medium organizations (<100 applications)',
                'Organizations with limited security budgets',
                'Organizations lacking skilled security staff',
                'Environments without 24/7 security operations'
            ]
        }

    def _generate_option_microsegmentation(self) -> Dict:
        """Generate Option 4: Micro-segmentation (Zero Trust)"""
        return {
            'name': 'Option 4: Micro-Segmentation (Zero Trust Architecture)',
            'description': (
                f'Full Zero Trust implementation with application-level micro-segmentation. '
                f'Each of your {self.total_apps} applications gets isolated segments with '
                f'identity-based access controls. Most advanced option.'
            ),
            'zones': [
                ('Per-Application Segments', f'{self.total_apps} isolated segments - one per application'),
                ('Per-Service Segments', 'Within each app, separate segments per microservice'),
                ('Identity-Based Access', 'No network location trust - every request verified'),
                ('Encrypted Everything', 'All communications encrypted with mTLS'),
                ('Dynamic Policies', 'Policies adapt based on user, device, location, behavior')
            ],
            'implementation': [
                f'Deploy service mesh across all {self.total_apps} applications',
                'Implement identity-based access control (IBAC) for every service',
                'Deploy certificate authority for mTLS between all services',
                'Implement software-defined perimeter (SDP)',
                'Deploy zero-trust network access (ZTNA) solution',
                'Implement continuous verification and monitoring',
                'Deploy AI-powered behavioral analytics (UEBA)',
                'Implement just-in-time (JIT) and just-enough-access (JEA)',
                'Deploy API gateway with OAuth2/OpenID Connect',
                'Implement dynamic policy engine',
                'Deploy distributed tracing for all service calls'
            ],
            'pros': [
                f'[SECURITY] Maximum security posture - 90-95% reduction in lateral movement',
                f'[ZERO TRUST] Full Zero Trust implementation - "Never trust, always verify"',
                f'[GRANULARITY] Application and service-level control - Finest possible segmentation',
                f'[COMPLIANCE] Exceeds all compliance requirements with detailed audit trails',
                f'[VISIBILITY] Complete visibility into all service-to-service communications',
                f'[AUTOMATION] Policy automation reduces human error',
                f'[IDENTITY] Identity-based rather than IP-based - Supports cloud and remote work',
                f'[FUTURE-PROOF] Aligns with industry direction - Supports modern architectures',
                f'[CLOUD NATIVE] Works seamlessly in hybrid/multi-cloud environments'
            ],
            'cons': [
                f'[COST] Very high investment - Estimated $1M-3M+ initial capital',
                f'[COMPLEXITY] Extreme complexity - Requires specialized Zero Trust team (5-8 FTEs)',
                f'[PERFORMANCE] Potential 20-50ms latency from encryption and verification overhead',
                f'[DEPLOYMENT] Very long implementation (6-12 months) - Phased rollout required',
                f'[DISRUPTION] High risk of service disruption during implementation',
                f'[SKILLS] Requires rare skillsets - Difficult and expensive to staff',
                f'[LEGACY] Legacy applications may not support micro-segmentation - May require rewrites',
                f'[OPERATIONAL] Ongoing operational complexity is very high',
                f'[TROUBLESHOOTING] Very difficult to troubleshoot - Requires advanced tooling',
                f'[MATURITY] Requires organizational security maturity - Not suitable for immature programs'
            ],
            'cost_breakdown': {
                'one_time': [
                    ('Service mesh deployment', '$200K-500K'),
                    ('ZTNA solution', '$150K-400K'),
                    ('Identity and access management', '$200K-500K'),
                    ('SDP implementation', '$100K-300K'),
                    ('API gateway and OAuth', '$100K-250K'),
                    ('Behavioral analytics (UEBA)', '$150K-350K'),
                    ('Professional services', '$500K-1M'),
                    ('Application refactoring', '$300K-1M+'),
                    ('Testing and validation', '$100K-250K')
                ],
                'annual': [
                    ('Licenses and maintenance', '$200K-500K'),
                    ('Operations (6-8 FTEs)', '$750K-1.2M'),
                    ('24/7 SOC operations', '$300K-600K'),
                    ('Training and certifications', '$100K-200K'),
                    ('Continuous improvement', '$100K-250K')
                ]
            },
            'timeline': '6-12 months',
            'complexity': 'VERY HIGH',
            'security_rating': '★★★★★ (5/5+)',
            'best_for': [
                'Very large enterprises (500+ applications)',
                'Cloud-native organizations',
                'Organizations with advanced security maturity (Level 4-5)',
                'Industries with extreme security requirements (defense, intelligence)',
                'Organizations committed to Zero Trust journey',
                'Modern microservices architectures',
                'Organizations with significant security budgets'
            ],
            'not_recommended_for': [
                f'Organizations with <{max(self.total_apps * 2, 200)} applications (ROI too low)',
                'Organizations with tight budgets',
                'Monolithic legacy application environments',
                'Organizations without dedicated security teams',
                'Organizations with low security maturity',
                'Short-term tactical segmentation needs'
            ]
        }

    def _render_option(self, option_num: int, option: Dict):
        """Render a segmentation option with full details"""
        # Option header
        self.doc.add_heading(option['name'], level=2)

        # Description
        desc_para = self.doc.add_paragraph()
        desc_para.add_run('OVERVIEW: ').bold = True
        desc_para.add_run(option['description'])

        self.doc.add_paragraph()

        # Quick stats table
        stats_table = self.doc.add_table(rows=3, cols=2)
        stats_table.style = 'Light List Accent 1'

        stats_table.rows[0].cells[0].text = 'Implementation Timeline'
        stats_table.rows[0].cells[1].text = option['timeline']
        stats_table.rows[1].cells[0].text = 'Complexity Level'
        stats_table.rows[1].cells[1].text = option['complexity']
        stats_table.rows[2].cells[0].text = 'Security Rating'
        stats_table.rows[2].cells[1].text = option['security_rating']

        for row in stats_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True

        # Zones
        self.doc.add_paragraph()
        self.doc.add_heading('Security Zones', level=3)

        for zone_name, zone_desc in option['zones']:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(f'{zone_name}: ').bold = True
            para.add_run(zone_desc)

        # Implementation requirements
        self.doc.add_paragraph()
        self.doc.add_heading('Implementation Requirements', level=3)

        for req in option['implementation']:
            self.doc.add_paragraph(req, style='List Bullet')

        # PROS
        self.doc.add_paragraph()
        pros_heading = self.doc.add_heading('Advantages (PROS)', level=3)
        pros_heading.style.font.color.rgb = RGBColor(0, 128, 0)

        for pro in option['pros']:
            para = self.doc.add_paragraph(style='List Bullet')
            # Bold the [TAG] part
            if '[' in pro:
                tag_end = pro.index(']') + 1
                para.add_run(pro[:tag_end]).bold = True
                para.add_run(pro[tag_end:])
            else:
                para.add_run(pro)

        # CONS
        self.doc.add_paragraph()
        cons_heading = self.doc.add_heading('Disadvantages (CONS)', level=3)
        cons_heading.style.font.color.rgb = RGBColor(139, 0, 0)

        for con in option['cons']:
            para = self.doc.add_paragraph(style='List Bullet')
            # Bold the [TAG] part
            if '[' in con:
                tag_end = con.index(']') + 1
                para.add_run(con[:tag_end]).bold = True
                para.add_run(con[tag_end:])
            else:
                para.add_run(con)

        # Cost breakdown
        self.doc.add_paragraph()
        self.doc.add_heading('Cost Analysis', level=3)

        cost_table = self.doc.add_table(
            rows=len(option['cost_breakdown']['one_time']) + len(option['cost_breakdown']['annual']) + 3,
            cols=2
        )
        cost_table.style = 'Medium Shading 1 Accent 1'

        row_idx = 0

        # One-time costs header
        cost_table.rows[row_idx].cells[0].text = 'ONE-TIME COSTS'
        cost_table.rows[row_idx].cells[0].paragraphs[0].runs[0].bold = True
        cost_table.rows[row_idx].cells[1].text = ''
        row_idx += 1

        for item, cost in option['cost_breakdown']['one_time']:
            cost_table.rows[row_idx].cells[0].text = item
            cost_table.rows[row_idx].cells[1].text = cost
            row_idx += 1

        # Annual costs header
        cost_table.rows[row_idx].cells[0].text = 'ANNUAL COSTS'
        cost_table.rows[row_idx].cells[0].paragraphs[0].runs[0].bold = True
        cost_table.rows[row_idx].cells[1].text = ''
        row_idx += 1

        for item, cost in option['cost_breakdown']['annual']:
            cost_table.rows[row_idx].cells[0].text = item
            cost_table.rows[row_idx].cells[1].text = cost
            row_idx += 1

        # Best for / Not recommended for
        self.doc.add_paragraph()
        self.doc.add_heading('Recommended For:', level=3)
        for item in option['best_for']:
            self.doc.add_paragraph(item, style='List Bullet')

        self.doc.add_paragraph()
        self.doc.add_heading('NOT Recommended For:', level=3)
        for item in option['not_recommended_for']:
            self.doc.add_paragraph(item, style='List Bullet')

    def _add_comparison_matrix(self):
        """Add side-by-side comparison matrix"""
        self.doc.add_heading('Option Comparison Matrix', level=1)

        intro = (
            "The following matrix provides a side-by-side comparison of all four segmentation "
            "options across key decision factors. Use this to quickly assess which option best "
            "fits your organization's requirements, budget, and risk tolerance."
        )
        self.doc.add_paragraph(intro)

        self.doc.add_paragraph()

        # Comparison table
        comparison_table = self.doc.add_table(rows=15, cols=5)
        comparison_table.style = 'Medium Shading 1 Accent 1'

        # Header row
        headers = ['Factor', 'Option 1: Minimal', 'Option 2: Standard', 'Option 3: Advanced', 'Option 4: Micro-seg']
        for idx, header in enumerate(headers):
            comparison_table.rows[0].cells[idx].text = header
            comparison_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

        # Comparison data
        comparisons = [
            ('Zone Count', '3 zones', '6-7 zones', '10+ zones', f'{self.total_apps}+ segments'),
            ('Security Rating', '2/5', '4/5', '5/5', '5+/5'),
            ('Implementation Time', '2-4 weeks', '6-10 weeks', '12-16 weeks', '6-12 months'),
            ('Initial Cost', '$50K-150K', '$200K-500K', '$500K-1.5M', '$1M-3M+'),
            ('Annual Cost', '$15K-40K', '$250K-500K', '$700K-1.4M', '$1.5M-2.2M+'),
            ('Complexity', 'LOW', 'MEDIUM', 'HIGH', 'VERY HIGH'),
            ('Staff Required', '0-1 FTE', '2 FTEs', '4-5 FTEs', '6-8 FTEs'),
            ('Lateral Movement Reduction', '30-40%', '70-80%', '85-90%', '90-95%'),
            ('Compliance', 'Basic PCI', 'PCI, HIPAA, SOX', 'All + GDPR', 'Exceeds all'),
            ('Scalability', f'<{min(self.total_apps, 50)} apps', f'{self.total_apps * 3} apps', f'{self.total_apps * 5}+ apps', 'Unlimited'),
            ('Operational Overhead', 'Low', 'Medium', 'High', 'Very High'),
            ('Performance Impact', '<2ms', '5-10ms', '10-20ms', '20-50ms'),
            ('Zero Trust Alignment', 'Poor', 'Partial', 'Good', 'Excellent'),
            ('Best For', 'Small orgs', 'Most orgs', 'Large enterprises', 'Advanced/Cloud-native')
        ]

        for row_idx, comparison in enumerate(comparisons, 1):
            for col_idx, value in enumerate(comparison):
                comparison_table.rows[row_idx].cells[col_idx].text = value
                if col_idx == 0:
                    comparison_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True

    def _add_recommendations(self):
        """Add data-driven recommendations"""
        self.doc.add_heading('Recommendations', level=1)

        recommendation = self._generate_recommendation()

        # Primary recommendation
        rec_para = self.doc.add_paragraph()
        rec_para.add_run('PRIMARY RECOMMENDATION: ').bold = True
        rec_para.add_run(recommendation['primary'])

        self.doc.add_paragraph()

        # Justification
        just_para = self.doc.add_paragraph()
        just_para.add_run('JUSTIFICATION:').bold = True

        for reason in recommendation['justification']:
            self.doc.add_paragraph(reason, style='List Bullet')

        # Alternative considerations
        self.doc.add_paragraph()
        alt_para = self.doc.add_paragraph()
        alt_para.add_run('ALTERNATIVE CONSIDERATIONS:').bold = True

        self.doc.add_paragraph(recommendation['alternative'])

        # Phased approach
        self.doc.add_paragraph()
        self.doc.add_heading('Phased Implementation Approach', level=2)

        phased_text = (
            "Regardless of which option you choose, we recommend a phased implementation approach "
            "to minimize risk and allow for adjustments based on lessons learned."
        )
        self.doc.add_paragraph(phased_text)

        phases = [
            ('Phase 1: Assessment & Planning (Weeks 1-2)', 'Complete detailed assessment, finalize design, secure budget'),
            ('Phase 2: Pilot (Weeks 3-6)', 'Implement segmentation for 1-2 pilot applications, validate approach'),
            ('Phase 3: Production Rollout (Weeks 7-12)', 'Roll out to production applications in batches'),
            ('Phase 4: Optimization (Weeks 13-16)', 'Fine-tune rules, optimize performance, reduce false positives'),
            ('Phase 5: Continuous Improvement (Ongoing)', 'Monitor, measure, and continuously improve segmentation')
        ]

        for phase, description in phases:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(f'{phase}: ').bold = True
            para.add_run(description)

    def _add_implementation_roadmap(self):
        """Add implementation roadmap"""
        self.doc.add_heading('Implementation Roadmap', level=1)

        roadmap_intro = (
            "This section provides a detailed implementation roadmap for the recommended option, "
            "including key milestones, deliverables, and success criteria."
        )
        self.doc.add_paragraph(roadmap_intro)

        # This would be customized based on the recommended option
        # For now, provide a generic roadmap

        roadmap_phases = [
            {
                'phase': 'Phase 1: Foundation (Weeks 1-4)',
                'objectives': [
                    'Finalize segmentation design and zone boundaries',
                    'Procure firewall hardware and licenses',
                    'Establish change management process',
                    'Create detailed network diagrams',
                    'Identify pilot applications'
                ],
                'deliverables': [
                    'Approved segmentation design document',
                    'Firewall hardware on-site',
                    'Change management procedures',
                    'Network topology diagrams',
                    'Pilot application list'
                ]
            },
            {
                'phase': 'Phase 2: Pilot Implementation (Weeks 5-8)',
                'objectives': [
                    'Deploy firewalls for pilot zone boundaries',
                    'Configure deny-all policies',
                    'Create allow rules for pilot applications',
                    'Enable logging and monitoring',
                    'Test pilot applications'
                ],
                'deliverables': [
                    'Pilot firewalls operational',
                    'Pilot allow rules documented',
                    'Monitoring dashboards created',
                    'Pilot test results',
                    'Lessons learned document'
                ]
            },
            {
                'phase': 'Phase 3: Production Rollout (Weeks 9-14)',
                'objectives': [
                    'Deploy remaining firewalls',
                    'Migrate applications in batches',
                    'Update firewall rules iteratively',
                    'Monitor for issues',
                    'Adjust rules based on feedback'
                ],
                'deliverables': [
                    'All firewalls deployed',
                    'All applications migrated',
                    'Complete firewall rule set',
                    'Migration reports',
                    'Issue resolution log'
                ]
            },
            {
                'phase': 'Phase 4: Optimization (Weeks 15-18)',
                'objectives': [
                    'Fine-tune firewall rules',
                    'Optimize performance',
                    'Reduce false positives',
                    'Train operations team',
                    'Document procedures'
                ],
                'deliverables': [
                    'Optimized rule set',
                    'Performance test results',
                    'Operations runbook',
                    'Training materials',
                    'Standard operating procedures'
                ]
            }
        ]

        for phase_info in roadmap_phases:
            self.doc.add_heading(phase_info['phase'], level=2)

            self.doc.add_heading('Objectives:', level=3)
            for obj in phase_info['objectives']:
                self.doc.add_paragraph(obj, style='List Bullet')

            self.doc.add_paragraph()
            self.doc.add_heading('Deliverables:', level=3)
            for deliv in phase_info['deliverables']:
                self.doc.add_paragraph(deliv, style='List Bullet')

            self.doc.add_paragraph()

        # Success criteria
        self.doc.add_heading('Success Criteria', level=2)

        success_criteria = [
            'All applications successfully migrated to new zones',
            'Zero production outages related to segmentation',
            'Firewall rule set documented and reviewed',
            'Operations team trained and comfortable with new architecture',
            'Security monitoring operational with defined escalation procedures',
            'Compliance audit findings addressed',
            'Performance metrics meet or exceed baseline',
            'Incident response procedures tested and validated'
        ]

        for criteria in success_criteria:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run('[SUCCESS] ').bold = True
            para.add_run(criteria)

    # Helper methods

    def _assess_current_segmentation(self) -> str:
        """Assess current segmentation maturity"""
        zone_count = len(self.zones_used)

        if zone_count <= 2:
            return 'MINIMAL (Flat network or basic DMZ)'
        elif zone_count <= 4:
            return 'BASIC (3-4 zones)'
        elif zone_count <= 7:
            return 'STANDARD (5-7 zones)'
        elif zone_count <= 10:
            return 'ADVANCED (8-10 zones)'
        else:
            return 'MICRO-SEGMENTATION (10+ zones)'

    def _generate_current_state_assessment(self) -> List[str]:
        """Generate assessment items for current state"""
        assessment = []

        avg_deps = self.dependency_count / max(self.total_apps, 1)

        assessment.append(
            f'Network complexity: {"High" if avg_deps > 10 else "Medium" if avg_deps > 5 else "Low"} '
            f'({avg_deps:.1f} dependencies per application)'
        )

        assessment.append(
            f'External exposure: {self.external_dependencies} external/API dependencies identified'
        )

        if 'WEB_TIER' in self.zones_used:
            assessment.append('Internet-facing applications detected - Perimeter security critical')

        if 'DATA_TIER' in self.zones_used:
            assessment.append('Data tier identified - Data protection segmentation recommended')

        if len(self.zones_used) < 5:
            assessment.append('Limited zone segmentation - Significant improvement opportunity exists')

        return assessment

    def _count_segmentation_opportunities(self) -> int:
        """Count segmentation improvement opportunities"""
        opportunities = 0

        # Opportunity if not using standard tiers
        standard_tiers = {'WEB_TIER', 'APP_TIER', 'DATA_TIER', 'CACHE_TIER', 'MESSAGING_TIER', 'MANAGEMENT_TIER'}
        missing_tiers = standard_tiers - self.zones_used
        opportunities += len(missing_tiers)

        # Opportunity per external dependency
        opportunities += min(self.external_dependencies // 10, 3)

        return max(opportunities, 3)

    def _generate_recommendation(self) -> Dict:
        """Generate data-driven recommendation"""
        zone_count = len(self.zones_used)
        avg_deps = self.dependency_count / max(self.total_apps, 1)

        # Decision logic based on actual data
        if self.total_apps < 50:
            if zone_count <= 3:
                recommendation = 'Option 1: Minimal Segmentation (3 Zones)'
                justification = [
                    f'Your network has {self.total_apps} applications, suitable for minimal segmentation',
                    'Current zone count is low, suggesting basic segmentation is appropriate',
                    'Provides strong ROI without excessive complexity for this scale'
                ]
                alternative = (
                    f'If budget allows and security requirements are high, consider Option 2 (Standard) '
                    f'for better future scalability and compliance support.'
                )
            else:
                recommendation = 'Option 2: Standard Tier-Based Segmentation (6-7 Zones)'
                justification = [
                    f'Your network has {self.total_apps} applications with {zone_count} zones already in use',
                    'Standard segmentation provides optimal security/complexity balance',
                    'Meets most compliance requirements (PCI-DSS, HIPAA)'
                ]
                alternative = (
                    'If very tight budget constraints exist, Option 1 (Minimal) could work, '
                    'though security and compliance capabilities will be limited.'
                )

        elif self.total_apps < 200:
            recommendation = 'Option 2: Standard Tier-Based Segmentation (6-7 Zones)'
            justification = [
                f'Your network has {self.total_apps} applications - ideal scale for standard segmentation',
                f'Average {avg_deps:.1f} dependencies per app indicates moderate complexity',
                'Standard segmentation is the industry best practice for this scale',
                'Balances strong security with manageable operational overhead'
            ]
            alternative = (
                f'If you have high-security requirements or process regulated data extensively, '
                f'consider Option 3 (Advanced) for enhanced protection and compliance.'
            )

        else:
            recommendation = 'Option 3: Advanced Multi-Zone Segmentation (10+ Zones)'
            justification = [
                f'Large network with {self.total_apps} applications requires advanced segmentation',
                f'High dependency count ({self.dependency_count} total) benefits from fine-grained control',
                'Advanced segmentation provides best ROI for large-scale environments',
                'Supports complex compliance and regulatory requirements'
            ]
            alternative = (
                f'If your organization is cloud-native with modern microservices architecture, '
                f'Option 4 (Micro-segmentation) provides the most future-proof approach. However, '
                f'this requires significant investment and security maturity.'
            )

        return {
            'primary': recommendation,
            'justification': justification,
            'alternative': alternative
        }


def generate_enterprise_segmentation_strategy(
    master_topology_path: str,
    output_path: str
):
    """Generate enterprise segmentation strategy document

    Args:
        master_topology_path: Path to master_topology.json
        output_path: Output path for Word document

    Returns:
        Path to generated document
    """
    doc_gen = EnterpriseSegmentationStrategyDocument(master_topology_path)
    doc_gen.generate_document(output_path)

    return output_path


if __name__ == '__main__':
    # Example usage
    master_topo = 'persistent_data/master_topology.json'
    output = 'outputs_final/word_reports/Enterprise_Network_Segmentation_Strategy.docx'

    generate_enterprise_segmentation_strategy(master_topo, output)

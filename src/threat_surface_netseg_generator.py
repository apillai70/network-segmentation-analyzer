"""
Threat Surface Analysis & Network Segmentation Document Generator
===================================================================
Generates security-focused documents analyzing threat surfaces and
providing network segmentation best practices recommendations.

Focus Areas:
- Attack surface mapping (external & internal)
- Threat vector identification
- Zero Trust segmentation strategy
- Actionable security recommendations

Author: Network Security Team
Version: 1.0
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

logger = logging.getLogger(__name__)


class ThreatSurfaceNetSegDocument:
    """Generates threat surface analysis and network segmentation documents"""

    # Risk scoring weights
    TIER_RISK_SCORES = {
        'WEB_TIER': 10,          # Highest risk - internet-facing
        'APP_TIER': 7,           # High risk - business logic
        'DATA_TIER': 9,          # Critical - sensitive data
        'CACHE_TIER': 5,         # Medium risk
        'MESSAGING_TIER': 6,     # Medium-high risk
        'MANAGEMENT_TIER': 8,    # High risk - administrative access
        'EXTERNAL': 10,          # External dependencies
        'UNKNOWN': 7             # Unknown = assume high risk
    }

    def __init__(self, app_name: str, app_data: Optional[Dict] = None):
        """Initialize document generator

        Args:
            app_name: Application name
            app_data: Application topology and analysis data
        """
        self.app_name = app_name
        self.app_data = app_data or {}
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles"""
        styles = self.doc.styles

        try:
            # Heading styles
            h1_style = styles['Heading 1']
            h1_style.font.size = Pt(18)
            h1_style.font.color.rgb = RGBColor(139, 0, 0)  # Dark red for security docs
            h1_style.font.bold = True

            h2_style = styles['Heading 2']
            h2_style.font.size = Pt(14)
            h2_style.font.color.rgb = RGBColor(178, 34, 34)  # Fire brick

            h3_style = styles['Heading 3']
            h3_style.font.size = Pt(12)
            h3_style.font.color.rgb = RGBColor(205, 92, 92)  # Indian red
        except:
            pass

    def generate_document(self, output_path: str):
        """Generate threat surface and network segmentation document

        Args:
            output_path: Output path for Word document
        """
        logger.info(f"Generating threat surface document for: {self.app_name}")

        # Cover page
        self._add_cover_page()
        self.doc.add_page_break()

        # Executive summary
        self._add_executive_summary()
        self.doc.add_page_break()

        # Threat surface analysis
        self._add_threat_surface_analysis()
        self.doc.add_page_break()

        # DNS Configuration Security
        self._add_dns_security_analysis()
        self.doc.add_page_break()

        # Attack vector analysis
        self._add_attack_vector_analysis()
        self.doc.add_page_break()

        # Zero Trust segmentation strategy
        self._add_zero_trust_strategy()
        self.doc.add_page_break()

        # Risk-Based Segmentation Decision Framework
        self._add_risk_based_segmentation_framework()
        self.doc.add_page_break()

        # Network segmentation recommendations
        self._add_segmentation_recommendations()
        self.doc.add_page_break()

        # Firewall rules and ACLs
        self._add_firewall_rules()
        self.doc.add_page_break()

        # Monitoring and detection
        self._add_monitoring_detection()
        self.doc.add_page_break()

        # Implementation roadmap
        self._add_implementation_roadmap()

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output_file))
        logger.info(f"✓ Threat surface document saved: {output_path}")

    def _add_cover_page(self):
        """Add professional cover page"""
        # Title
        title = self.doc.add_heading(
            f'Threat Surface Analysis &\nNetwork Segmentation Best Practices\n\n{self.app_name}',
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Security classification
        classification = self.doc.add_paragraph()
        classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classification.add_run('🔒 CONFIDENTIAL - SECURITY ANALYSIS 🔒')
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(139, 0, 0)

        for _ in range(2):
            self.doc.add_paragraph()

        # Document info
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run('Version: 1.0\n').bold = True
        info_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
        info_para.add_run('Classification: Confidential\n')
        info_para.add_run(f'Application: {self.app_name}\n')

        zone = self.app_data.get('security_zone', 'Unknown')
        risk = self.app_data.get('risk_level', 'MEDIUM')
        info_para.add_run(f'Security Zone: {zone}\n')
        info_para.add_run(f'Risk Level: {risk}\n')

        for _ in range(3):
            self.doc.add_paragraph()

        # Footer
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Prepared by: Prutech Network Security Team\n')
        footer.add_run('Enterprise Security & Zero Trust Architecture\n')
        footer.add_run('Auto-generated by Network Segmentation Analyzer v3.0\n')

    def _add_executive_summary(self):
        """Add executive summary with key findings"""
        self.doc.add_heading('Executive Summary', level=1)

        # Calculate threat score
        threat_score = self._calculate_threat_score()
        risk_rating = self._get_risk_rating(threat_score)

        # Break into shorter sentences for better Word rendering
        para1 = self.doc.add_paragraph()
        para1.add_run(f"This document provides a comprehensive security architecture assessment and network segmentation strategy for the {self.app_name} application. ")
        para1.add_run("The analysis evaluates the current security posture and provides actionable recommendations to enhance protection through modern Zero Trust security principles.")

        self.doc.add_paragraph()
        threat_para = self.doc.add_paragraph()
        threat_para.add_run('SECURITY POSTURE ASSESSMENT:').bold = True
        self.doc.add_paragraph(f"• Security Maturity Score: {threat_score}/100")
        self.doc.add_paragraph(f"• Enhancement Priority: {risk_rating}")
        self.doc.add_paragraph(f"• Security Zone: {self.app_data.get('security_zone', 'Unknown')}")
        self.doc.add_paragraph(f"• Focus Areas: {self._get_primary_concerns()}")

        self.doc.add_paragraph()
        key_findings_para = self.doc.add_paragraph()
        key_findings_para.add_run('KEY RECOMMENDATIONS:').bold = True

        # Break into shorter sentences for better Word rendering
        para2 = self.doc.add_paragraph()
        para2.add_run("This application has opportunities to strengthen its security architecture through strategic network segmentation and enhanced security controls. ")
        para2.add_run("This document outlines specific recommendations to implement industry-leading defense-in-depth security practices aligned with modern Zero Trust principles.")

        # Quick stats table
        self.doc.add_paragraph()
        self.doc.add_heading('Security Metrics', level=2)

        metrics_table = self.doc.add_table(rows=6, cols=2)
        metrics_table.style = 'Medium Shading 1 Accent 2'

        dependencies = self.app_data.get('predicted_dependencies', [])
        external_deps = [d for d in dependencies if d.get('type') not in ['database', 'cache']]

        metrics = [
            ('Threat Score', f'{threat_score}/100'),
            ('Risk Rating', risk_rating),
            ('External Dependencies', str(len(external_deps))),
            ('Infrastructure Dependencies', str(len(dependencies) - len(external_deps))),
            ('Recommended Segmentation Rules', self._count_required_rules()),
            ('Critical Vulnerabilities', self._count_critical_issues())
        ]

        for idx, (metric, value) in enumerate(metrics):
            metrics_table.rows[idx].cells[0].text = metric
            metrics_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            metrics_table.rows[idx].cells[1].text = value

    def _add_threat_surface_analysis(self):
        """Add comprehensive threat surface analysis"""
        self.doc.add_heading('Threat Surface Analysis', level=1)

        # External attack surface
        self.doc.add_heading('External Attack Surface', level=2)

        external_text = "The external attack surface represents entry points accessible from outside the organization's network perimeter. These are the primary targets for external threat actors."
        self.doc.add_paragraph(external_text)

        # Analyze external exposure
        zone = self.app_data.get('security_zone', 'UNKNOWN')
        is_web_facing = zone == 'WEB_TIER'

        self.doc.add_heading('Internet-Facing Components', level=3)

        if is_web_facing:
            external_components = [
                'Web Tier: Direct internet exposure through load balancers',
                'API Endpoints: RESTful APIs accessible from public internet',
                'Authentication Services: Login and identity verification endpoints',
                'Static Content: CDN and web assets served to end users',
                'Health Check Endpoints: Status monitoring URLs'
            ]

            for comp in external_components:
                self.doc.add_paragraph(comp, style='List Bullet')

            self.doc.add_paragraph()
            self.doc.add_paragraph(
                '⚠️ CRITICAL: All internet-facing components must have WAF, DDoS protection, '
                'rate limiting, and comprehensive logging enabled.',
                style='Intense Quote'
            )
        else:
            self.doc.add_paragraph(
                f'✓ Application is deployed in {zone} - No direct internet exposure detected. '
                'This significantly reduces external attack surface.'
            )

        # Internal attack surface
        self.doc.add_paragraph()
        self.doc.add_heading('Internal Attack Surface', level=2)

        internal_text = "The internal attack surface represents lateral movement opportunities for threat actors who have gained initial access to the internal network. Proper segmentation is critical to prevent lateral movement and contain breaches."
        self.doc.add_paragraph(internal_text)

        dependencies = self.app_data.get('predicted_dependencies', [])

        self.doc.add_heading('Internal Communication Paths', level=3)

        # Group dependencies by type
        dep_by_type = defaultdict(list)
        unknown_count = 0
        for dep in dependencies:
            dep_type = dep.get('type', 'unknown')
            if dep_type == 'unknown':
                unknown_count += 1
            dep_by_type[dep_type].append(dep.get('name', 'Unknown'))

        for dep_type, dep_names in dep_by_type.items():
            if dep_type == 'unknown':
                self.doc.add_paragraph(f"{dep_type.title()}*: {len(dep_names)} connection(s)", style='List Bullet')
            else:
                self.doc.add_paragraph(f"{dep_type.title()}: {len(dep_names)} connection(s)", style='List Bullet')

        # Add asterisk explanation for unknown connections
        if unknown_count > 0:
            self.doc.add_paragraph()
            note_para = self.doc.add_paragraph()
            note_para.add_run('* Unknown connections: ').bold = True
            explanation_run = note_para.add_run(
                f'These {unknown_count} connection(s) could not be definitively classified based on available '
                'ExtraHop network flow data. This may occur when: (1) destination endpoints do not have clear '
                'service type indicators in their network signatures, (2) flow data lacks sufficient context '
                'to determine the application protocol, or (3) connections involve custom or proprietary services '
                'without standard port/protocol patterns. Manual investigation and correlation with application '
                'configuration is recommended to properly classify these dependencies.'
            )
            explanation_run.italic = True

        # Trust boundaries
        self.doc.add_paragraph()
        self.doc.add_heading('Trust Boundaries', level=3)

        boundaries = [
            'Internet → Web Tier: PUBLIC to INTERNAL boundary (highest risk)',
            'Web Tier → App Tier: DMZ to INTERNAL boundary',
            'App Tier → Data Tier: APPLICATION to DATA boundary (critical)',
            'App Tier → External Services: INTERNAL to EXTERNAL boundary',
            'Management Tier → All Tiers: ADMIN to PRODUCTION boundary'
        ]

        for boundary in boundaries:
            parts = boundary.split(':')
            if len(parts) == 2:
                para = self.doc.add_paragraph(style='List Bullet')
                para.add_run(parts[0] + ': ').bold = True
                para.add_run(parts[1])

    def _add_dns_security_analysis(self):
        """Add DNS configuration security analysis"""
        self.doc.add_heading('DNS Configuration Security', level=1)

        # Get DNS validation data
        dns_validation = self.app_data.get('dns_validation', {})
        validation_metadata = self.app_data.get('validation_metadata', {})

        intro_text = "DNS (Domain Name System) configuration is a critical component of network security. " \
                     "Misconfigured DNS can lead to man-in-the-middle attacks, service disruptions, " \
                     "and reconnaissance opportunities for attackers. This section analyzes DNS " \
                     "configuration for this application."
        self.doc.add_paragraph(intro_text)

        # DNS Validation Summary
        self.doc.add_heading('DNS Validation Results', level=2)

        if not dns_validation or dns_validation.get('total_validated', 0) == 0:
            para = self.doc.add_paragraph()
            para.add_run("⚠️ DNS validation data not available. ").bold = True
            para.add_run(
                "DNS validation was not performed during analysis. This gap prevents "
                "identification of potential DNS misconfigurations, split-brain DNS scenarios, "
                "and stale DNS records. "
            )
            para.add_run("\n\nRecommendation: ").bold = True
            para.add_run(
                "Enable DNS validation in batch processing to identify configuration issues "
                "that could impact security and availability."
            )
            self.doc.add_paragraph()
            return

        # DNS validation statistics table
        validation_table = self.doc.add_table(rows=6, cols=2)
        validation_table.style = 'Medium Shading 1 Accent 2'

        validation_stats = [
            ('IPs Validated', dns_validation.get('total_validated', 0)),
            ('Valid DNS (Perfect Match)', dns_validation.get('valid', 0)),
            ('Multiple IPs Detected', dns_validation.get('valid_multiple_ips', 0)),
            ('DNS Mismatches', dns_validation.get('mismatch', 0)),
            ('Non-Existent Records (NXDOMAIN)', dns_validation.get('nxdomain', 0)),
            ('Validation Failures', dns_validation.get('failed', 0))
        ]

        for idx, (metric, value) in enumerate(validation_stats):
            validation_table.rows[idx].cells[0].text = metric
            validation_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            validation_table.rows[idx].cells[1].text = str(value)

        self.doc.add_paragraph()

        # Security Assessment
        total_issues = (dns_validation.get('mismatch', 0) +
                       dns_validation.get('nxdomain', 0) +
                       dns_validation.get('failed', 0))

        assessment_para = self.doc.add_paragraph()
        if total_issues == 0:
            assessment_para.add_run("✓ DNS Security Status: HEALTHY").bold = True
            assessment_para.add_run("\n\nNo DNS configuration issues detected. All forward and reverse DNS lookups match correctly.")
            assessment_para.font.color.rgb = RGBColor(0, 128, 0)
        elif total_issues <= 5:
            assessment_para.add_run("⚠️ DNS Security Status: MINOR ISSUES").bold = True
            assessment_para.font.color.rgb = RGBColor(255, 165, 0)
            assessment_para.add_run(f"\n\n{total_issues} DNS configuration issue(s) detected. These should be addressed to improve network reliability and security posture.")
        else:
            assessment_para.add_run("🔴 DNS Security Status: ATTENTION REQUIRED").bold = True
            assessment_para.font.color.rgb = RGBColor(255, 0, 0)
            assessment_para.add_run(f"\n\n{total_issues} DNS configuration issues detected. Immediate remediation recommended.")

        self.doc.add_paragraph()

        # DNS Mismatches (Security Risk)
        if dns_validation.get('mismatch', 0) > 0:
            self.doc.add_heading('DNS Mismatches - Security Risk', level=2)

            mismatch_text = "DNS mismatches occur when forward DNS (hostname → IP) does not match reverse DNS (IP → hostname). " \
                           "This can indicate:"
            self.doc.add_paragraph(mismatch_text)

            risks = [
                'Stale DNS records that could be exploited for attacks',
                'Split-brain DNS configurations leading to inconsistent routing',
                'Potential DNS cache poisoning or man-in-the-middle scenarios',
                'Configuration drift between DNS servers',
                'Hostname spoofing opportunities'
            ]

            for risk in risks:
                para = self.doc.add_paragraph(style='List Bullet')
                para.add_run('⚠️ ').font.color.rgb = RGBColor(255, 165, 0)
                para.add_run(risk)

            self.doc.add_paragraph()

            # Extract actual mismatches from validation metadata
            mismatches = []
            for ip, metadata in validation_metadata.items():
                if metadata.get('status') == 'mismatch':
                    mismatches.append({
                        'ip': ip,
                        'reverse': metadata.get('reverse_hostname', 'N/A'),
                        'forward_ip': metadata.get('forward_ip', 'N/A'),
                        'mismatch': metadata.get('mismatch', 'Unknown')
                    })

            if mismatches[:5]:  # Show up to 5 examples
                self.doc.add_heading('Example Mismatches:', level=3)
                mismatch_table = self.doc.add_table(rows=len(mismatches[:5])+1, cols=3)
                mismatch_table.style = 'Light Grid Accent 2'

                # Header
                mismatch_table.rows[0].cells[0].text = 'IP Address'
                mismatch_table.rows[0].cells[1].text = 'Reverse DNS'
                mismatch_table.rows[0].cells[2].text = 'Forward IP'
                for cell in mismatch_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].bold = True

                # Data
                for idx, mm in enumerate(mismatches[:5], 1):
                    mismatch_table.rows[idx].cells[0].text = mm['ip']
                    mismatch_table.rows[idx].cells[1].text = mm['reverse']
                    mismatch_table.rows[idx].cells[2].text = mm['forward_ip']

                self.doc.add_paragraph()

            # Remediation
            remediation_para = self.doc.add_paragraph()
            remediation_para.add_run('Remediation Steps:').bold = True

            remediation_steps = [
                'Audit all DNS records (forward and reverse) for accuracy',
                'Ensure PTR records match A records for all IPs',
                'Synchronize DNS changes across all authoritative servers',
                'Implement DNS monitoring to detect future mismatches',
                'Review and update any hardcoded IP addresses in configurations'
            ]

            for step in remediation_steps:
                self.doc.add_paragraph(step, style='List Bullet')

            self.doc.add_paragraph()

        # Multiple IPs (VM + ESXi Scenarios)
        if dns_validation.get('valid_multiple_ips', 0) > 0:
            self.doc.add_heading('Multiple IP Addresses - Virtualization Awareness', level=2)

            multi_ip_text = f"{dns_validation.get('valid_multiple_ips', 0)} hostname(s) were detected with multiple IP addresses. " \
                           "This is common in virtualized environments where both the VM and ESXi host share the same hostname."
            self.doc.add_paragraph(multi_ip_text)

            self.doc.add_paragraph()

            security_considerations = [
                {
                    'title': 'Network Segmentation Impact',
                    'description': 'Firewall rules must account for BOTH IP addresses (VM and ESXi host). '
                                 'Missing either IP could create security gaps or cause legitimate traffic blocking.'
                },
                {
                    'title': 'Monitoring and Alerting',
                    'description': 'Security monitoring tools should track both IPs under the same hostname. '
                                 'Treat traffic from either IP as belonging to the same logical entity.'
                },
                {
                    'title': 'Attack Surface',
                    'description': 'ESXi management interfaces represent a critical attack vector. Ensure ESXi '
                                 'management networks are properly isolated from production VM networks.'
                },
                {
                    'title': 'Incident Response',
                    'description': 'During security incidents, investigate BOTH the VM and hypervisor layer. '
                                 'Compromise of the hypervisor affects all VMs on that host.'
                }
            ]

            for consideration in security_considerations:
                self.doc.add_heading(consideration['title'], level=3)
                self.doc.add_paragraph(consideration['description'])

            self.doc.add_paragraph()

            # Best practices
            self.doc.add_heading('Best Practices for Virtualized Environments:', level=3)

            best_practices = [
                'Maintain separate DNS zones for VM management and ESXi management',
                'Use distinct hostname patterns (e.g., vm-web01 vs esxi-host01)',
                'Implement micro-segmentation between VM and management networks',
                'Apply security hardening guides for both guest OS and hypervisor',
                'Monitor both VM and ESXi networks for anomalous behavior'
            ]

            for practice in best_practices:
                self.doc.add_paragraph(practice, style='List Bullet')

            self.doc.add_paragraph()

        # NXDOMAIN Issues
        if dns_validation.get('nxdomain', 0) > 0:
            self.doc.add_heading('Non-Existent DNS Records (NXDOMAIN)', level=2)

            nxdomain_text = f"{dns_validation.get('nxdomain', 0)} IP address(es) have no reverse DNS records. " \
                           "While not always a security risk, this can impact:"
            self.doc.add_paragraph(nxdomain_text)

            impacts = [
                'Network troubleshooting and forensics (harder to identify systems)',
                'Log analysis and correlation (IPs instead of meaningful hostnames)',
                'Compliance requirements (some standards require reverse DNS)',
                'Third-party integrations (may reject connections without reverse DNS)',
                'Professional appearance (lack of reverse DNS appears unprofessional)'
            ]

            for impact in impacts:
                self.doc.add_paragraph(impact, style='List Bullet')

            self.doc.add_paragraph()

            recommendation_para = self.doc.add_paragraph()
            recommendation_para.add_run('Recommendation: ').bold = True
            recommendation_para.add_run(
                'Add PTR records for all active production systems to improve network '
                'visibility and operational efficiency. Exclude only temporary or '
                'ephemeral addresses (DHCP pools, container networks).'
            )

            self.doc.add_paragraph()

        # DNS Security Best Practices
        self.doc.add_heading('DNS Security Best Practices', level=2)

        best_practice_categories = [
            {
                'category': 'DNSSEC (DNS Security Extensions)',
                'practices': [
                    'Enable DNSSEC for all authoritative zones to prevent DNS spoofing',
                    'Validate DNSSEC signatures on DNS responses',
                    'Monitor DNSSEC validation failures as potential attack indicators'
                ]
            },
            {
                'category': 'Access Controls',
                'practices': [
                    'Restrict DNS zone transfers to authorized secondary servers only',
                    'Implement DNS query rate limiting to prevent DDoS amplification',
                    'Use split-horizon DNS to separate internal and external views'
                ]
            },
            {
                'category': 'Monitoring and Logging',
                'practices': [
                    'Log all DNS queries and responses for security analysis',
                    'Alert on DNS queries to known malicious domains',
                    'Monitor for DNS tunneling and exfiltration attempts',
                    'Track DNS query volumes for DDoS detection'
                ]
            },
            {
                'category': 'Configuration Management',
                'practices': [
                    'Maintain DNS records in version control with approval workflows',
                    'Implement automated DNS consistency checking',
                    'Regular audits of DNS records for accuracy and necessity',
                    'Document all DNS changes with business justification'
                ]
            }
        ]

        for category_info in best_practice_categories:
            self.doc.add_heading(category_info['category'], level=3)
            for practice in category_info['practices']:
                self.doc.add_paragraph(practice, style='List Bullet')
            self.doc.add_paragraph()

    def _add_attack_vector_analysis(self):
        """Add attack vector analysis"""
        self.doc.add_heading('Attack Vector Analysis', level=1)

        attack_vectors = [
            {
                'name': 'External Web Application Attacks',
                'likelihood': 'HIGH' if self.app_data.get('security_zone') == 'WEB_TIER' else 'LOW',
                'impact': 'HIGH',
                'vectors': [
                    'SQL Injection through web forms and API parameters',
                    'Cross-Site Scripting (XSS) in user input fields',
                    'Authentication bypass and credential stuffing',
                    'API abuse and parameter tampering',
                    'DDoS attacks targeting availability'
                ],
                'mitigations': [
                    'Deploy Web Application Firewall (WAF) with OWASP rules',
                    'Implement input validation and output encoding',
                    'Use parameterized queries and ORM frameworks',
                    'Enable rate limiting and IP reputation filtering',
                    'Implement multi-factor authentication (MFA)'
                ]
            },
            {
                'name': 'Lateral Movement (East-West)',
                'likelihood': 'MEDIUM',
                'impact': 'HIGH',
                'vectors': [
                    'Compromised service account credentials',
                    'Unencrypted internal communications',
                    'Overly permissive network ACLs',
                    'Shared service accounts across tiers',
                    'Unrestricted database access'
                ],
                'mitigations': [
                    'Implement Zero Trust micro-segmentation',
                    'Enforce TLS for all internal communications',
                    'Apply principle of least privilege for service accounts',
                    'Use unique credentials per service/tier',
                    'Deploy network intrusion detection (IDS/IPS)'
                ]
            },
            {
                'name': 'Data Exfiltration',
                'likelihood': 'MEDIUM',
                'impact': 'CRITICAL',
                'vectors': [
                    'Direct database access from compromised app tier',
                    'API abuse to extract sensitive data',
                    'Backup file theft',
                    'Log file containing sensitive data',
                    'Unencrypted data in transit'
                ],
                'mitigations': [
                    'Implement database activity monitoring (DAM)',
                    'Apply data loss prevention (DLP) controls',
                    'Encrypt all data at rest and in transit',
                    'Implement egress filtering and monitoring',
                    'Sanitize logs to remove sensitive data'
                ]
            },
            {
                'name': 'Supply Chain Attacks',
                'likelihood': 'MEDIUM',
                'impact': 'HIGH',
                'vectors': [
                    'Compromised third-party dependencies',
                    'Malicious container images',
                    'Vulnerable open-source libraries',
                    'Compromised CI/CD pipeline',
                    'Malicious infrastructure-as-code templates'
                ],
                'mitigations': [
                    'Scan all dependencies for known vulnerabilities',
                    'Use trusted container registries with image signing',
                    'Implement software composition analysis (SCA)',
                    'Secure CI/CD pipeline with access controls',
                    'Verify cryptographic signatures of all artifacts'
                ]
            },
            {
                'name': 'Privilege Escalation',
                'likelihood': 'LOW',
                'impact': 'CRITICAL',
                'vectors': [
                    'Misconfigured RBAC permissions',
                    'Kernel exploits in container runtime',
                    'Service account token theft',
                    'Sudo/admin credential compromise',
                    'Unpatched OS vulnerabilities'
                ],
                'mitigations': [
                    'Implement strict RBAC with regular audits',
                    'Use security contexts and pod security policies',
                    'Rotate service account tokens regularly',
                    'Disable or restrict sudo/admin access',
                    'Maintain aggressive patch management program'
                ]
            }
        ]

        for idx, vector in enumerate(attack_vectors, 1):
            self.doc.add_heading(f"{idx}. {vector['name']}", level=2)

            # Risk assessment
            risk_para = self.doc.add_paragraph()
            risk_para.add_run('Likelihood: ').bold = True
            risk_para.add_run(f"{vector['likelihood']}  |  ")
            risk_para.add_run('Impact: ').bold = True
            risk_para.add_run(vector['impact'])

            self.doc.add_paragraph()

            # Attack vectors
            self.doc.add_heading('Attack Vectors:', level=3)
            for v in vector['vectors']:
                self.doc.add_paragraph(v, style='List Bullet')

            self.doc.add_paragraph()

            # Mitigations
            self.doc.add_heading('Recommended Mitigations:', level=3)
            for m in vector['mitigations']:
                self.doc.add_paragraph(f'✓ {m}', style='List Bullet')

            if idx < len(attack_vectors):
                self.doc.add_paragraph()

    def _add_zero_trust_strategy(self):
        """Add Zero Trust segmentation strategy"""
        self.doc.add_heading('Zero Trust Network Segmentation Strategy', level=1)

        intro_text = 'Zero Trust is a security model that eliminates implicit trust and requires continuous verification of all connections. The core principle is "never trust, always verify" regardless of the network location.'
        self.doc.add_paragraph(intro_text)

        # Core principles
        self.doc.add_heading('Core Zero Trust Principles', level=2)

        principles = [
            ('Verify Explicitly', 'Always authenticate and authorize based on all available data points including user identity, location, device health, service, workload, data classification, and anomalies.'),
            ('Least Privilege Access', 'Limit user and service access with Just-In-Time and Just-Enough-Access (JIT/JEA), risk-based adaptive policies, and data protection.'),
            ('Assume Breach', 'Minimize blast radius and segment access. Verify end-to-end encryption and use analytics to get visibility, drive threat detection, and improve defenses.')
        ]

        for principle, description in principles:
            self.doc.add_heading(principle, level=3)
            self.doc.add_paragraph(description)

        # Segmentation model
        self.doc.add_paragraph()
        self.doc.add_heading('Micro-Segmentation Model', level=2)

        segment_text = f"For {self.app_name}, implement the following micro-segmentation strategy:"
        self.doc.add_paragraph(segment_text)

        segments = [
            {
                'name': 'Web Tier Segment',
                'description': 'Isolate public-facing web tier',
                'ingress': 'Internet (HTTPS/443 only)',
                'egress': 'App Tier (authenticated API calls only)',
                'controls': 'WAF, DDoS protection, rate limiting, geo-blocking'
            },
            {
                'name': 'Application Tier Segment',
                'description': 'Business logic isolation',
                'ingress': 'Web Tier (authenticated requests)',
                'egress': 'Data Tier, Cache Tier, External APIs',
                'controls': 'Service mesh, mTLS, API gateway, RBAC'
            },
            {
                'name': 'Data Tier Segment',
                'description': 'Database and storage isolation',
                'ingress': 'App Tier (authenticated database queries)',
                'egress': 'None (data should not initiate external connections)',
                'controls': 'Database firewall, encryption, audit logging, query monitoring'
            },
            {
                'name': 'Management Tier Segment',
                'description': 'Administrative access isolation',
                'ingress': 'Admin networks only (MFA required)',
                'egress': 'All tiers (monitoring and management)',
                'controls': 'Bastion hosts, VPN, privileged access management (PAM)'
            }
        ]

        for segment in segments:
            self.doc.add_heading(segment['name'], level=3)
            self.doc.add_paragraph(segment['description'])

            seg_table = self.doc.add_table(rows=3, cols=2)
            seg_table.style = 'Light Grid Accent 2'

            seg_table.rows[0].cells[0].text = 'Allowed Ingress'
            seg_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            seg_table.rows[0].cells[1].text = segment['ingress']

            seg_table.rows[1].cells[0].text = 'Allowed Egress'
            seg_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
            seg_table.rows[1].cells[1].text = segment['egress']

            seg_table.rows[2].cells[0].text = 'Security Controls'
            seg_table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
            seg_table.rows[2].cells[1].text = segment['controls']

            self.doc.add_paragraph()

    def _add_risk_based_segmentation_framework(self):
        """Add risk-based segmentation decision framework with regulatory compliance"""
        self.doc.add_heading('Risk-Based Segmentation Decision Framework', level=1)

        intro_text = ("This section provides a structured framework for making segmentation decisions based on "
                     "risk-reward analysis, regulatory mandates, and operational requirements. The goal is to "
                     "balance security benefits against implementation costs and complexity.")
        self.doc.add_paragraph(intro_text)

        # Risk Assessment Methodology
        self.doc.add_heading('Risk Assessment Methodology', level=2)

        risk_intro = ("Segmentation decisions should be driven by quantitative risk analysis. Use the following "
                     "formula to calculate the segmentation priority for each application or data classification:")
        self.doc.add_paragraph(risk_intro)

        self.doc.add_paragraph()
        formula_para = self.doc.add_paragraph()
        formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formula_run = formula_para.add_run('Risk Score = (Data Sensitivity × Exposure Level × Business Impact) / Mitigation Cost')
        formula_run.bold = True
        formula_run.font.size = Pt(12)
        self.doc.add_paragraph()

        # Risk scoring table
        risk_table = self.doc.add_table(rows=5, cols=5)
        risk_table.style = 'Medium Shading 1 Accent 2'

        # Header row
        risk_table.rows[0].cells[0].text = 'Factor'
        risk_table.rows[0].cells[1].text = 'Low (1-3)'
        risk_table.rows[0].cells[2].text = 'Medium (4-6)'
        risk_table.rows[0].cells[3].text = 'High (7-9)'
        risk_table.rows[0].cells[4].text = 'Critical (10)'
        for cell in risk_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        risk_factors = [
            ('Data Sensitivity', 'Public data', 'Internal use only', 'Confidential/PII', 'PCI/PHI/Regulated'),
            ('Exposure Level', 'Internal only', 'Partner networks', 'Controlled internet', 'Public internet'),
            ('Business Impact', 'Dev/Test', 'Non-critical prod', 'Business-critical', 'Revenue-generating'),
            ('Mitigation Cost', 'Policy-based ($)', 'Firewall rules ($$)', 'New infrastructure ($$$)', 'Complete redesign ($$$$)')
        ]

        for idx, (factor, low, med, high, crit) in enumerate(risk_factors, 1):
            risk_table.rows[idx].cells[0].text = factor
            risk_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            risk_table.rows[idx].cells[1].text = low
            risk_table.rows[idx].cells[2].text = med
            risk_table.rows[idx].cells[3].text = high
            risk_table.rows[idx].cells[4].text = crit

        self.doc.add_paragraph()

        # Risk score interpretation
        self.doc.add_heading('Risk Score Interpretation:', level=3)
        risk_scores = [
            ('100-300: IMMEDIATE ACTION REQUIRED', 'Segment immediately. Critical security gap that poses existential risk.'),
            ('50-99: HIGH PRIORITY', 'Segment within 30 days. Significant risk with clear regulatory or business drivers.'),
            ('20-49: MEDIUM PRIORITY', 'Segment within 90 days as part of planned architecture improvements.'),
            ('10-19: LOW PRIORITY', 'Consider in next major infrastructure refresh or when budget allows.'),
            ('<10: DEFER', 'Low risk/reward ratio. Focus resources on higher-priority segmentation efforts.')
        ]

        for score_range, interpretation in risk_scores:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(score_range + ': ').bold = True
            para.add_run(interpretation)

        # Regulatory Compliance Mapping
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self.doc.add_heading('Regulatory Compliance Mapping', level=2)

        reg_intro = ("Different regulatory frameworks mandate specific network segmentation requirements. "
                    "Understanding which regulations apply to your data helps determine segmentation scope and rigor.")
        self.doc.add_paragraph(reg_intro)

        # PCI-DSS
        self.doc.add_heading('PCI-DSS (Payment Card Industry)', level=3)
        pci_req = ("PCI-DSS Requirement 1.2.1: Explicitly REQUIRES network segmentation to reduce CDE scope. "
                  "Segmentation must isolate systems that store, process, or transmit cardholder data.")
        self.doc.add_paragraph(pci_req)

        pci_requirements = [
            ('Requirement 1.1', 'Firewall configuration standards with documented business justification for all allowed services'),
            ('Requirement 1.2', 'Build firewall configurations that restrict connections between untrusted and trusted networks'),
            ('Requirement 1.3', 'Prohibit direct public access between internet and CDE - require DMZ architecture'),
            ('Requirement 11.3.4', 'Penetration testing must verify segmentation controls effectively isolate CDE')
        ]

        pci_table = self.doc.add_table(rows=len(pci_requirements)+1, cols=2)
        pci_table.style = 'Light Grid Accent 2'
        pci_table.rows[0].cells[0].text = 'PCI-DSS Requirement'
        pci_table.rows[0].cells[1].text = 'Segmentation Implication'
        for cell in pci_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (req, impl) in enumerate(pci_requirements, 1):
            pci_table.rows[idx].cells[0].text = req
            pci_table.rows[idx].cells[1].text = impl

        self.doc.add_paragraph()
        pci_verdict = self.doc.add_paragraph()
        pci_verdict.add_run('PCI-DSS Verdict: ').bold = True
        pci_verdict.add_run('SEGMENTATION MANDATORY. Non-compliance = Loss of card processing ability.')

        # HIPAA
        self.doc.add_paragraph()
        self.doc.add_heading('HIPAA (Healthcare)', level=3)
        hipaa_req = ("HIPAA Security Rule 164.312(a)(1) requires 'access controls' to limit access to ePHI "
                    "to authorized users. Network segmentation is considered an industry best practice to achieve compliance.")
        self.doc.add_paragraph(hipaa_req)

        hipaa_requirements = [
            ('164.308(a)(3)', 'Workforce access - Implement policies limiting access to ePHI to authorized personnel'),
            ('164.308(a)(4)', 'Security incident procedures - Segmentation limits blast radius of breaches (reduces penalties)'),
            ('164.312(a)(1)', 'Access controls - Technical safeguards to allow only authorized access to ePHI systems'),
            ('164.312(e)(1)', 'Transmission security - Protect ePHI during transmission over networks')
        ]

        hipaa_table = self.doc.add_table(rows=len(hipaa_requirements)+1, cols=2)
        hipaa_table.style = 'Light Grid Accent 2'
        hipaa_table.rows[0].cells[0].text = 'HIPAA Requirement'
        hipaa_table.rows[0].cells[1].text = 'Segmentation Benefit'
        for cell in hipaa_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (req, benefit) in enumerate(hipaa_requirements, 1):
            hipaa_table.rows[idx].cells[0].text = req
            hipaa_table.rows[idx].cells[1].text = benefit

        self.doc.add_paragraph()
        hipaa_verdict = self.doc.add_paragraph()
        hipaa_verdict.add_run('HIPAA Verdict: ').bold = True
        hipaa_verdict.add_run('STRONGLY RECOMMENDED. Required by most cyber insurance policies for covered entities.')

        # SOX
        self.doc.add_paragraph()
        self.doc.add_heading('SOX (Sarbanes-Oxley - Financial Reporting)', level=3)
        sox_req = ("SOX Section 404 requires management to assess internal controls over financial reporting. "
                  "Segmentation supports segregation of duties and audit trail requirements.")
        self.doc.add_paragraph(sox_req)

        sox_points = [
            'Segregation of Duties: Separate dev, test, and production environments to prevent unauthorized financial data modification',
            'Access Controls: Limit access to financial reporting systems to authorized personnel only',
            'Change Management: Network segmentation enables strict change control for production financial systems',
            'Audit Trails: Segmentation facilitates comprehensive logging of access to financial data systems'
        ]

        for point in sox_points:
            self.doc.add_paragraph(point, style='List Bullet')

        sox_verdict = self.doc.add_paragraph()
        sox_verdict.add_run('SOX Verdict: ').bold = True
        sox_verdict.add_run('RECOMMENDED for internal controls. Demonstrates mature risk management to auditors.')

        # GDPR
        self.doc.add_paragraph()
        self.doc.add_heading('GDPR (EU Data Protection)', level=3)
        gdpr_req = ("GDPR Article 32 requires 'appropriate technical and organizational measures' including "
                   "network segmentation to protect personal data. Segmentation reduces breach notification scope.")
        self.doc.add_paragraph(gdpr_req)

        gdpr_benefits = [
            ('Data Minimization', 'Segment systems processing EU personal data to limit exposure and access'),
            ('Breach Notification', 'Segmentation can limit breach scope, potentially avoiding 72-hour notification if contained'),
            ('Privacy by Design', 'Network segmentation demonstrates proactive technical measures under Article 25'),
            ('Data Subject Rights', 'Easier to locate and manage personal data when systems are logically segmented')
        ]

        for benefit, description in gdpr_benefits:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(benefit + ': ').bold = True
            para.add_run(description)

        gdpr_verdict = self.doc.add_paragraph()
        gdpr_verdict.add_run('GDPR Verdict: ').bold = True
        gdpr_verdict.add_run('REQUIRED for high-risk processing. Reduces potential fines (up to 4% global revenue).')

        # Zone Sizing Decision Framework
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self.doc.add_heading('Zone Sizing Decision Framework: How Many Zones?', level=2)

        zone_intro = ("The optimal number of security zones depends on your organization's size, complexity, "
                     "risk tolerance, and regulatory requirements. Use this decision tree:")
        self.doc.add_paragraph(zone_intro)

        # Decision matrix table
        decision_table = self.doc.add_table(rows=4, cols=5)
        decision_table.style = 'Medium Shading 1 Accent 2'

        decision_table.rows[0].cells[0].text = 'Organization Profile'
        decision_table.rows[0].cells[1].text = 'Zone Count'
        decision_table.rows[0].cells[2].text = 'Use Case'
        decision_table.rows[0].cells[3].text = 'Regulatory Drivers'
        decision_table.rows[0].cells[4].text = 'Complexity'
        for cell in decision_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        profiles = [
            ('MINIMAL (3 zones)', '3-5', 'Small organizations, <50 apps, low regulatory requirements', 'None or minimal', 'Low - DMZ, Internal, Management'),
            ('STANDARD (5-7 zones)', '5-7', 'Mid-size enterprises, 50-200 apps, moderate compliance', 'PCI-DSS, SOX, or HIPAA', 'Medium - Web, App, Data, Cache, Messaging, Mgmt, External'),
            ('ADVANCED (10+ zones)', '10+', 'Large enterprises, 200+ apps, heavy regulation, Zero Trust', 'Multiple: PCI + HIPAA + GDPR', 'High - Micro-segmentation per app tier + environment + sensitivity')
        ]

        for idx, (profile, count, use_case, reg, complexity) in enumerate(profiles, 1):
            decision_table.rows[idx].cells[0].text = profile
            decision_table.rows[idx].cells[1].text = count
            decision_table.rows[idx].cells[2].text = use_case
            decision_table.rows[idx].cells[3].text = reg
            decision_table.rows[idx].cells[4].text = complexity

        self.doc.add_paragraph()

        # Recommended zone breakdown for this application
        self.doc.add_heading(f'Recommended Segmentation for {self.app_name}', level=3)

        dependencies = self.app_data.get('predicted_dependencies', [])
        zone = self.app_data.get('security_zone', 'APP_TIER')
        app_count = 1  # Simplified - would need context for real count

        if len(dependencies) > 15 or zone == 'DATA_TIER':
            recommendation = 'ADVANCED (10+ zones)'
            justification = (f"This application has {len(dependencies)} dependencies and operates in {zone}. "
                           f"Recommend advanced micro-segmentation to limit blast radius and support compliance requirements.")
        elif len(dependencies) > 5 or zone == 'WEB_TIER':
            recommendation = 'STANDARD (5-7 zones)'
            justification = (f"This application has {len(dependencies)} dependencies. "
                           f"Standard segmentation provides adequate security with manageable operational overhead.")
        else:
            recommendation = 'STANDARD (5-7 zones)'
            justification = ("Even simple applications benefit from standard tier-based segmentation for defense-in-depth.")

        rec_para = self.doc.add_paragraph()
        rec_para.add_run('Recommendation: ').bold = True
        rec_para.add_run(recommendation)
        self.doc.add_paragraph(justification)

        # Classification Criteria - What Goes Where?
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self.doc.add_heading('Classification Criteria: What Goes Where?', level=2)

        class_intro = ("Use the following decision matrix to classify applications and assign them to security zones. "
                      "Classification should be based on multiple factors, not just function.")
        self.doc.add_paragraph(class_intro)

        # Classification matrix
        self.doc.add_heading('Zone Assignment Decision Matrix', level=3)

        class_matrix_table = self.doc.add_table(rows=8, cols=5)
        class_matrix_table.style = 'Light Grid Accent 2'

        # Header
        class_matrix_table.rows[0].cells[0].text = 'Security Zone'
        class_matrix_table.rows[0].cells[1].text = 'Function'
        class_matrix_table.rows[0].cells[2].text = 'Data Classification'
        class_matrix_table.rows[0].cells[3].text = 'Exposure'
        class_matrix_table.rows[0].cells[4].text = 'Example Applications'
        for cell in class_matrix_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        zone_classifications = [
            ('WEB_TIER', 'User-facing services', 'Public, Internal', 'Internet (HTTPS/443)', 'Load balancers, Web servers, CDN, API gateways'),
            ('APP_TIER', 'Business logic', 'Confidential, Internal', 'Internal only (from Web tier)', 'Application servers, Microservices, API backends'),
            ('DATA_TIER', 'Data storage & processing', 'Highly Confidential, Regulated', 'Internal only (from App tier)', 'Databases (SQL/NoSQL), Data warehouses, File servers'),
            ('CACHE_TIER', 'Performance optimization', 'Confidential (cached data)', 'Internal only', 'Redis, Memcached, Varnish'),
            ('MESSAGING_TIER', 'Async communication', 'Confidential', 'Internal only', 'Kafka, RabbitMQ, Message queues'),
            ('MANAGEMENT_TIER', 'Admin & operations', 'Highly Confidential (admin access)', 'Admin networks only + MFA', 'Monitoring, Logging, Bastion hosts, CI/CD'),
            ('EXTERNAL', 'Third-party services', 'Variable', 'Controlled outbound', 'SaaS APIs, Payment processors, Cloud services')
        ]

        for idx, (zone_name, function, data_class, exposure, examples) in enumerate(zone_classifications, 1):
            class_matrix_table.rows[idx].cells[0].text = zone_name
            class_matrix_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            class_matrix_table.rows[idx].cells[1].text = function
            class_matrix_table.rows[idx].cells[2].text = data_class
            class_matrix_table.rows[idx].cells[3].text = exposure
            class_matrix_table.rows[idx].cells[4].text = examples

        self.doc.add_paragraph()

        # Decision flowchart in text
        self.doc.add_heading('Zone Assignment Decision Flowchart', level=3)
        flowchart_steps = [
            ('1. Does it accept inbound connections from the internet?', 'YES → WEB_TIER | NO → Continue'),
            ('2. Does it store or process regulated data (PCI/PHI/PII)?', 'YES → DATA_TIER | NO → Continue'),
            ('3. Does it provide administrative/management functions?', 'YES → MANAGEMENT_TIER | NO → Continue'),
            ('4. Does it cache frequently accessed data for performance?', 'YES → CACHE_TIER | NO → Continue'),
            ('5. Does it handle asynchronous messaging or event streams?', 'YES → MESSAGING_TIER | NO → Continue'),
            ('6. Is it a third-party service outside your control?', 'YES → EXTERNAL | NO → APP_TIER (default)')
        ]

        for step, decision in flowchart_steps:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(step).bold = True
            para.add_run(f'\n   {decision}')

        # Firewall Policy Templates
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self.doc.add_heading('Firewall Policy Templates by Zone Boundary', level=2)

        policy_intro = ("Each zone boundary requires specific firewall policies. Use these templates as a starting point "
                       "and customize based on your application's actual requirements.")
        self.doc.add_paragraph(policy_intro)

        # Internet → WEB_TIER template
        self.doc.add_heading('Template 1: Internet → WEB_TIER', level=3)

        template1_table = self.doc.add_table(rows=6, cols=5)
        template1_table.style = 'Light Grid Accent 2'
        template1_table.rows[0].cells[0].text = 'Rule #'
        template1_table.rows[0].cells[1].text = 'Source'
        template1_table.rows[0].cells[2].text = 'Destination'
        template1_table.rows[0].cells[3].text = 'Port/Protocol'
        template1_table.rows[0].cells[4].text = 'Purpose'
        for cell in template1_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        template1_rules = [
            ('100', 'Internet (any)', 'WEB_TIER (LB VIP)', 'TCP/443', 'HTTPS traffic'),
            ('110', 'Internet (any)', 'WEB_TIER (LB VIP)', 'TCP/80', 'HTTP→HTTPS redirect'),
            ('900', 'Internet (any)', 'WEB_TIER (any)', 'ANY', 'Default DENY all other'),
            ('910', 'WEB_TIER (any)', 'Internet (any)', 'TCP/443', 'Outbound HTTPS (CDN, APIs)'),
            ('999', 'WEB_TIER (any)', 'Internet (any)', 'ANY', 'Default DENY all other')
        ]

        for idx, (rule_num, source, dest, port, purpose) in enumerate(template1_rules, 1):
            template1_table.rows[idx].cells[0].text = rule_num
            template1_table.rows[idx].cells[1].text = source
            template1_table.rows[idx].cells[2].text = dest
            template1_table.rows[idx].cells[3].text = port
            template1_table.rows[idx].cells[4].text = purpose

        self.doc.add_paragraph()
        template1_note = self.doc.add_paragraph()
        template1_note.add_run('Additional Controls: ').bold = True
        template1_note.add_run('WAF with OWASP Core Rule Set, DDoS protection (rate limiting: 1000 req/sec per IP), Geo-blocking (allow only operational countries), Bot detection')

        # WEB_TIER → APP_TIER template
        self.doc.add_paragraph()
        self.doc.add_heading('Template 2: WEB_TIER → APP_TIER', level=3)

        template2_table = self.doc.add_table(rows=4, cols=5)
        template2_table.style = 'Light Grid Accent 2'
        template2_table.rows[0].cells[0].text = 'Rule #'
        template2_table.rows[0].cells[1].text = 'Source'
        template2_table.rows[0].cells[2].text = 'Destination'
        template2_table.rows[0].cells[3].text = 'Port/Protocol'
        template2_table.rows[0].cells[4].text = 'Purpose'
        for cell in template2_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        template2_rules = [
            ('200', 'WEB_TIER (web servers)', 'APP_TIER (app servers)', 'TCP/8080', 'API calls (authenticated)'),
            ('210', 'APP_TIER (app servers)', 'WEB_TIER (web servers)', 'TCP/8080', 'API responses'),
            ('999', 'WEB_TIER (any)', 'APP_TIER (any)', 'ANY', 'Default DENY all other')
        ]

        for idx, (rule_num, source, dest, port, purpose) in enumerate(template2_rules, 1):
            template2_table.rows[idx].cells[0].text = rule_num
            template2_table.rows[idx].cells[1].text = source
            template2_table.rows[idx].cells[2].text = dest
            template2_table.rows[idx].cells[3].text = port
            template2_table.rows[idx].cells[4].text = purpose

        self.doc.add_paragraph()
        template2_note = self.doc.add_paragraph()
        template2_note.add_run('Additional Controls: ').bold = True
        template2_note.add_run('TLS 1.3 required, mTLS for service-to-service auth, API gateway with JWT validation, Connection rate limiting: 100 conn/sec per web server')

        # APP_TIER → DATA_TIER template
        self.doc.add_paragraph()
        self.doc.add_heading('Template 3: APP_TIER → DATA_TIER', level=3)

        template3_table = self.doc.add_table(rows=5, cols=5)
        template3_table.style = 'Light Grid Accent 2'
        template3_table.rows[0].cells[0].text = 'Rule #'
        template3_table.rows[0].cells[1].text = 'Source'
        template3_table.rows[0].cells[2].text = 'Destination'
        template3_table.rows[0].cells[3].text = 'Port/Protocol'
        template3_table.rows[0].cells[4].text = 'Purpose'
        for cell in template3_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        template3_rules = [
            ('300', 'APP_TIER (app servers)', 'DATA_TIER (PostgreSQL)', 'TCP/5432', 'Database queries'),
            ('310', 'APP_TIER (app servers)', 'DATA_TIER (MySQL)', 'TCP/3306', 'Database queries'),
            ('320', 'APP_TIER (app servers)', 'CACHE_TIER (Redis)', 'TCP/6379', 'Cache operations'),
            ('999', 'APP_TIER (any)', 'DATA_TIER (any)', 'ANY', 'Default DENY all other')
        ]

        for idx, (rule_num, source, dest, port, purpose) in enumerate(template3_rules, 1):
            template3_table.rows[idx].cells[0].text = rule_num
            template3_table.rows[idx].cells[1].text = source
            template3_table.rows[idx].cells[2].text = dest
            template3_table.rows[idx].cells[3].text = port
            template3_table.rows[idx].cells[4].text = purpose

        self.doc.add_paragraph()
        template3_note = self.doc.add_paragraph()
        template3_note.add_run('Additional Controls: ').bold = True
        template3_note.add_run('Database firewall with query inspection, TLS encryption for all database connections, Database Activity Monitoring (DAM), Query anomaly detection, Connection pooling limits')

        # MANAGEMENT_TIER template
        self.doc.add_paragraph()
        self.doc.add_heading('Template 4: MANAGEMENT_TIER → All Tiers', level=3)

        template4_table = self.doc.add_table(rows=5, cols=5)
        template4_table.style = 'Light Grid Accent 2'
        template4_table.rows[0].cells[0].text = 'Rule #'
        template4_table.rows[0].cells[1].text = 'Source'
        template4_table.rows[0].cells[2].text = 'Destination'
        template4_table.rows[0].cells[3].text = 'Port/Protocol'
        template4_table.rows[0].cells[4].text = 'Purpose'
        for cell in template4_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        template4_rules = [
            ('400', 'Admin Networks (VPN)', 'MANAGEMENT_TIER (Bastion)', 'TCP/22', 'SSH to bastion (MFA required)'),
            ('410', 'MANAGEMENT_TIER (Bastion)', 'All Tiers (servers)', 'TCP/22', 'SSH to servers'),
            ('420', 'All Tiers (servers)', 'MANAGEMENT_TIER (SIEM)', 'TCP/514', 'Syslog forwarding'),
            ('999', 'MANAGEMENT_TIER (any)', 'All Tiers (any)', 'ANY', 'Default DENY all other')
        ]

        for idx, (rule_num, source, dest, port, purpose) in enumerate(template4_rules, 1):
            template4_table.rows[idx].cells[0].text = rule_num
            template4_table.rows[idx].cells[1].text = source
            template4_table.rows[idx].cells[2].text = dest
            template4_table.rows[idx].cells[3].text = port
            template4_table.rows[idx].cells[4].text = purpose

        self.doc.add_paragraph()
        template4_note = self.doc.add_paragraph()
        template4_note.add_run('Additional Controls: ').bold = True
        template4_note.add_run('Privileged Access Management (PAM) with session recording, MFA required for all admin access, Time-based access (deny outside business hours without approval), Just-In-Time (JIT) access provisioning')

        # Cost-Benefit Analysis
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self.doc.add_heading('Cost-Benefit Analysis', level=2)

        costbenefit_intro = ("Understanding the costs and benefits of segmentation helps justify investments and prioritize efforts.")
        self.doc.add_paragraph(costbenefit_intro)

        # Benefits table
        self.doc.add_heading('Quantifiable Benefits', level=3)

        benefits_table = self.doc.add_table(rows=6, cols=3)
        benefits_table.style = 'Medium Shading 1 Accent 2'
        benefits_table.rows[0].cells[0].text = 'Benefit Category'
        benefits_table.rows[0].cells[1].text = 'Impact'
        benefits_table.rows[0].cells[2].text = 'Estimated Value'
        for cell in benefits_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        benefits_data = [
            ('Reduced Breach Scope', '70-90% reduction in lateral movement', '$2-5M avg. savings per breach (IBM Cost of Data Breach 2024)'),
            ('Compliance Fines Avoidance', 'Meet PCI-DSS, HIPAA, GDPR requirements', '$50K-50M+ in avoided fines'),
            ('Cyber Insurance Premium Reduction', '10-30% discount with mature segmentation', '$50K-500K annually for enterprises'),
            ('Reduced PCI Audit Scope', '50-80% fewer systems in scope', '$100K-1M in annual compliance costs'),
            ('Faster Incident Response', 'Mean Time to Contain (MTTC) reduced 60%', '3-5 fewer days of business disruption per incident')
        ]

        for idx, (category, impact, value) in enumerate(benefits_data, 1):
            benefits_table.rows[idx].cells[0].text = category
            benefits_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            benefits_table.rows[idx].cells[1].text = impact
            benefits_table.rows[idx].cells[2].text = value

        self.doc.add_paragraph()

        # Costs table
        self.doc.add_heading('Implementation Costs', level=3)

        costs_table = self.doc.add_table(rows=6, cols=3)
        costs_table.style = 'Medium Shading 1 Accent 2'
        costs_table.rows[0].cells[0].text = 'Cost Category'
        costs_table.rows[0].cells[1].text = 'One-Time'
        costs_table.rows[0].cells[2].text = 'Ongoing (Annual)'
        for cell in costs_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        costs_data = [
            ('Firewall Infrastructure', '$50K-500K (hardware/cloud)', '$10K-100K (maintenance, licenses)'),
            ('Network Redesign & Implementation', '$100K-1M (consulting, engineering)', '$20K-200K (operations, changes)'),
            ('Security Tools (IDS/IPS, SIEM, PAM)', '$200K-2M (licenses, deployment)', '$50K-500K (licenses, support)'),
            ('Staff Training & Documentation', '$50K-200K (training, process development)', '$20K-100K (ongoing training)'),
            ('Operational Overhead', '$100K-500K (initial process changes)', '$50K-300K (change management, monitoring)')
        ]

        for idx, (category, onetime, ongoing) in enumerate(costs_data, 1):
            costs_table.rows[idx].cells[0].text = category
            costs_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            costs_table.rows[idx].cells[1].text = onetime
            costs_table.rows[idx].cells[2].text = ongoing

        self.doc.add_paragraph()

        # ROI summary
        roi_para = self.doc.add_paragraph()
        roi_para.add_run('Return on Investment (ROI): ').bold = True
        roi_para.add_run('For most enterprises, network segmentation pays for itself after preventing a SINGLE significant breach. '
                        'Typical ROI positive within 12-18 months when factoring in compliance cost reduction and insurance savings.')

        # Final recommendations summary
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self.doc.add_heading('Implementation Recommendations Summary', level=2)

        final_recommendations = [
            ('Start with Regulatory Requirements', 'If you process payment cards (PCI), healthcare data (HIPAA), or EU personal data (GDPR), segmentation is effectively mandatory. Start there.'),
            ('Use Risk Scoring', 'Calculate risk scores (Data Sensitivity × Exposure × Business Impact / Mitigation Cost) to prioritize which applications to segment first.'),
            ('Adopt Standard Zone Model', 'For most organizations, the 5-7 zone model (Web, App, Data, Cache, Messaging, Management, External) provides optimal security/complexity balance.'),
            ('Implement Incrementally', 'Start with the highest-risk zone boundary (typically App → Data tier for PCI/HIPAA compliance), then expand.'),
            ('Use Policy Templates', 'Customize the firewall policy templates provided in this document rather than starting from scratch.'),
            ('Measure and Improve', 'Track metrics like MTTC (Mean Time to Contain), segmentation coverage %, and compliance audit findings to demonstrate value.')
        ]

        for recommendation, description in final_recommendations:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(recommendation + ': ').bold = True
            para.add_run(description)

    def _add_segmentation_recommendations(self):
        """Add specific segmentation recommendations"""
        self.doc.add_heading('Network Segmentation Recommendations', level=1)

        # Immediate actions
        self.doc.add_heading('Immediate Actions (Priority 1)', level=2)

        p1_actions = [
            'Deploy network firewalls between all tier boundaries',
            'Implement explicit deny-all policies with allow-listing for required traffic',
            'Enable TLS 1.3 for all inter-tier communications',
            'Deploy database activity monitoring (DAM) for Data Tier',
            'Implement privileged access management (PAM) for Management Tier',
            'Enable comprehensive logging for all cross-tier communications'
        ]

        for action in p1_actions:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run('🔴 ').font.color.rgb = RGBColor(255, 0, 0)
            para.add_run(action)

        # Short-term actions
        self.doc.add_paragraph()
        self.doc.add_heading('Short-Term Actions (30-60 days)', level=2)

        p2_actions = [
            'Deploy service mesh (Istio/Linkerd) for application tier',
            'Implement API gateway with rate limiting and throttling',
            'Enable container network policies in Kubernetes',
            'Deploy intrusion detection/prevention systems (IDS/IPS)',
            'Implement secrets management solution (Vault/AWS Secrets Manager)',
            'Enable runtime application self-protection (RASP)',
            'Deploy security information and event management (SIEM)'
        ]

        for action in p2_actions:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run('🟡 ').font.color.rgb = RGBColor(255, 165, 0)
            para.add_run(action)

        # Long-term actions
        self.doc.add_paragraph()
        self.doc.add_heading('Long-Term Actions (90+ days)', level=2)

        p3_actions = [
            'Implement full Zero Trust architecture with continuous verification',
            'Deploy advanced threat protection (ATP) solutions',
            'Implement user and entity behavior analytics (UEBA)',
            'Enable automated threat response and remediation',
            'Conduct regular penetration testing and red team exercises',
            'Implement chaos engineering for security resilience testing'
        ]

        for action in p3_actions:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run('🟢 ').font.color.rgb = RGBColor(0, 128, 0)
            para.add_run(action)

    def _add_firewall_rules(self):
        """Add firewall rules and ACLs"""
        self.doc.add_heading('Firewall Rules and Network ACLs', level=1)

        rules_intro = "The following firewall rules implement the Zero Trust micro-segmentation strategy. All rules follow the principle of explicit deny with specific allow rules for required traffic."
        self.doc.add_paragraph(rules_intro)

        # Default policies
        self.doc.add_heading('Default Policies', level=2)

        default_policies = [
            ('Ingress Default', 'DENY ALL', 'Block all inbound traffic by default'),
            ('Egress Default', 'DENY ALL', 'Block all outbound traffic by default'),
            ('Inter-Tier Default', 'DENY ALL', 'Block all cross-tier traffic by default'),
            ('Logging', 'LOG ALL', 'Log all allowed and denied connections')
        ]

        policy_table = self.doc.add_table(rows=len(default_policies)+1, cols=3)
        policy_table.style = 'Medium Shading 1 Accent 2'

        # Header
        policy_table.rows[0].cells[0].text = 'Policy Type'
        policy_table.rows[0].cells[1].text = 'Action'
        policy_table.rows[0].cells[2].text = 'Description'
        for cell in policy_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (policy, action, desc) in enumerate(default_policies, 1):
            policy_table.rows[idx].cells[0].text = policy
            policy_table.rows[idx].cells[1].text = action
            policy_table.rows[idx].cells[2].text = desc

        # Allow rules
        self.doc.add_paragraph()
        self.doc.add_heading('Required Allow Rules', level=2)

        zone = self.app_data.get('security_zone', 'APP_TIER')
        dependencies = self.app_data.get('predicted_dependencies', [])

        allow_rules = []

        # Web tier rules
        if zone == 'WEB_TIER':
            allow_rules.extend([
                ('Internet → Web Tier', 'TCP/443', 'HTTPS traffic from internet', 'HIGH'),
                ('Internet → Web Tier', 'TCP/80', 'HTTP redirect to HTTPS', 'MEDIUM'),
                ('Web Tier → App Tier', 'TCP/8080', 'API calls to backend', 'HIGH'),
                ('App Tier → Web Tier', 'TCP/8080', 'API responses', 'HIGH')
            ])

        # App tier rules (always needed)
        allow_rules.extend([
            ('App Tier → Data Tier', 'TCP/5432', 'PostgreSQL database access', 'HIGH'),
            ('App Tier → Data Tier', 'TCP/3306', 'MySQL database access', 'HIGH'),
            ('App Tier → Cache Tier', 'TCP/6379', 'Redis cache access', 'MEDIUM'),
            ('App Tier → Messaging Tier', 'TCP/9092', 'Kafka messaging', 'MEDIUM')
        ])

        # Add dependency-specific rules
        for dep in dependencies[:5]:  # Limit to first 5 for brevity
            dep_type = dep.get('type', 'service')
            dep_name = dep.get('name', 'Unknown')
            port = self._guess_port(dep_type)
            allow_rules.append((
                f'App Tier → {dep_name}',
                port,
                f'Access to {dep_type}',
                'MEDIUM'
            ))

        # Management rules
        allow_rules.extend([
            ('Management → All Tiers', 'TCP/22', 'SSH administrative access', 'CRITICAL'),
            ('All Tiers → Management', 'TCP/514', 'Syslog to SIEM', 'HIGH'),
            ('All Tiers → Management', 'TCP/9090', 'Metrics to monitoring', 'HIGH')
        ])

        # Create rules table
        rules_table = self.doc.add_table(rows=len(allow_rules)+1, cols=4)
        rules_table.style = 'Light Grid Accent 2'

        # Header
        headers = ['Source → Destination', 'Port/Protocol', 'Purpose', 'Priority']
        for idx, header in enumerate(headers):
            rules_table.rows[0].cells[idx].text = header
            rules_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

        for idx, (flow, port, purpose, priority) in enumerate(allow_rules, 1):
            rules_table.rows[idx].cells[0].text = flow
            rules_table.rows[idx].cells[1].text = port
            rules_table.rows[idx].cells[2].text = purpose
            rules_table.rows[idx].cells[3].text = priority

        # Additional security controls
        self.doc.add_paragraph()
        self.doc.add_heading('Additional Security Controls', level=2)

        additional_controls = [
            'Enable stateful inspection on all firewalls',
            'Implement connection rate limiting (max connections per second)',
            'Enable geo-blocking for countries outside operational regions',
            'Implement application-layer gateway (ALG) inspection',
            'Enable threat intelligence feed integration',
            'Configure automatic IP blacklisting for repeated violations'
        ]

        for control in additional_controls:
            self.doc.add_paragraph(control, style='List Bullet')

    def _add_monitoring_detection(self):
        """Add monitoring and threat detection requirements"""
        self.doc.add_heading('Monitoring and Threat Detection', level=1)

        monitor_intro = "Comprehensive monitoring and threat detection are essential components of defense-in-depth security. The following capabilities should be implemented to detect and respond to security incidents."
        self.doc.add_paragraph(monitor_intro)

        # Logging requirements
        self.doc.add_heading('Logging Requirements', level=2)

        log_categories = [
            {
                'category': 'Authentication and Access',
                'events': [
                    'All authentication attempts (success and failure)',
                    'Authorization decisions and access denials',
                    'Privileged access and administrative actions',
                    'Service account usage and API key access',
                    'Session creation, modification, and termination'
                ]
            },
            {
                'category': 'Network Activity',
                'events': [
                    'All firewall allow and deny decisions',
                    'Cross-tier communications and data flows',
                    'External API calls and responses',
                    'DNS queries and resolutions',
                    'Connection attempts to unauthorized ports/services'
                ]
            },
            {
                'category': 'Data Access',
                'events': [
                    'Database queries and transactions',
                    'Sensitive data access and modifications',
                    'Large data exports or unusual queries',
                    'Schema changes and permission modifications',
                    'Backup and restore operations'
                ]
            },
            {
                'category': 'System Events',
                'events': [
                    'Configuration changes',
                    'Software installations and updates',
                    'Service starts, stops, and crashes',
                    'Resource exhaustion events',
                    'Security control bypass attempts'
                ]
            }
        ]

        for category_info in log_categories:
            self.doc.add_heading(category_info['category'], level=3)
            for event in category_info['events']:
                self.doc.add_paragraph(event, style='List Bullet')
            self.doc.add_paragraph()

        # Detection rules
        self.doc.add_heading('Threat Detection Rules', level=2)

        detection_rules = [
            ('Brute Force Detection', 'Alert on >10 failed authentication attempts within 5 minutes from same source'),
            ('Lateral Movement', 'Alert on unusual cross-tier communications or first-time connections'),
            ('Data Exfiltration', 'Alert on large data transfers (>100MB) or unusual egress patterns'),
            ('Privilege Escalation', 'Alert on privilege changes or attempts to access restricted resources'),
            ('Anomalous Behavior', 'Alert on deviations from baseline behavior (ML-based detection)'),
            ('Known Threats', 'Alert on matches with threat intelligence feeds (IPs, domains, hashes)'),
            ('Protocol Anomalies', 'Alert on malformed packets or protocol violations'),
            ('Time-Based Anomalies', 'Alert on access during non-business hours or from unusual locations')
        ]

        detection_table = self.doc.add_table(rows=len(detection_rules)+1, cols=2)
        detection_table.style = 'Light List Accent 2'

        detection_table.rows[0].cells[0].text = 'Detection Rule'
        detection_table.rows[0].cells[1].text = 'Trigger Condition'
        for cell in detection_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (rule, condition) in enumerate(detection_rules, 1):
            detection_table.rows[idx].cells[0].text = rule
            detection_table.rows[idx].cells[1].text = condition

        # Response procedures
        self.doc.add_paragraph()
        self.doc.add_heading('Incident Response Procedures', level=2)

        response_text = "Upon detection of security incidents, the following response procedures should be followed:"
        self.doc.add_paragraph(response_text)

        response_steps = [
            ('DETECT', 'Automated detection through SIEM, IDS/IPS, or manual discovery'),
            ('ASSESS', 'Determine severity (Critical/High/Medium/Low) and scope of incident'),
            ('CONTAIN', 'Isolate affected systems, block malicious IPs, disable compromised accounts'),
            ('INVESTIGATE', 'Collect forensic evidence, analyze logs, determine root cause'),
            ('ERADICATE', 'Remove malicious artifacts, patch vulnerabilities, strengthen controls'),
            ('RECOVER', 'Restore systems from clean backups, verify integrity, resume operations'),
            ('LESSONS LEARNED', 'Post-incident review, update procedures, improve detection')
        ]

        for step, description in response_steps:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(f'{step}: ').bold = True
            para.add_run(description)

    def _add_implementation_roadmap(self):
        """Add implementation roadmap"""
        self.doc.add_heading('Implementation Roadmap', level=1)

        roadmap_intro = "The following phased approach ensures systematic implementation of network segmentation and threat surface reduction while minimizing operational disruption."
        self.doc.add_paragraph(roadmap_intro)

        phases = [
            {
                'phase': 'Phase 1: Assessment and Planning (Week 1-2)',
                'objectives': [
                    'Complete threat surface assessment',
                    'Document current network architecture',
                    'Identify critical assets and data flows',
                    'Define security zones and trust boundaries',
                    'Develop detailed implementation plan'
                ],
                'deliverables': [
                    'Threat surface analysis report (this document)',
                    'Current state network diagram',
                    'Target state security architecture',
                    'Project plan with timeline and resources'
                ]
            },
            {
                'phase': 'Phase 2: Foundation (Week 3-6)',
                'objectives': [
                    'Deploy network firewalls between tiers',
                    'Implement deny-all default policies',
                    'Enable comprehensive logging to SIEM',
                    'Deploy privileged access management (PAM)',
                    'Establish change management process'
                ],
                'deliverables': [
                    'Firewall rules documentation',
                    'SIEM integration and dashboards',
                    'PAM system operational',
                    'Change control procedures'
                ]
            },
            {
                'phase': 'Phase 3: Hardening (Week 7-10)',
                'objectives': [
                    'Implement TLS for all inter-tier communications',
                    'Deploy service mesh with mTLS',
                    'Configure database activity monitoring',
                    'Deploy IDS/IPS systems',
                    'Implement API gateway with security controls'
                ],
                'deliverables': [
                    'TLS certificate management system',
                    'Service mesh operational',
                    'Database monitoring active',
                    'IDS/IPS signatures deployed'
                ]
            },
            {
                'phase': 'Phase 4: Advanced Security (Week 11-14)',
                'objectives': [
                    'Deploy threat intelligence integration',
                    'Implement behavioral analytics (UEBA)',
                    'Configure automated incident response',
                    'Deploy runtime application protection (RASP)',
                    'Establish security metrics and KPIs'
                ],
                'deliverables': [
                    'Threat intelligence feeds operational',
                    'UEBA baselines established',
                    'Automated response playbooks',
                    'Security dashboard and reports'
                ]
            },
            {
                'phase': 'Phase 5: Optimization (Week 15+)',
                'objectives': [
                    'Fine-tune detection rules and reduce false positives',
                    'Conduct penetration testing',
                    'Optimize performance of security controls',
                    'Establish continuous improvement process',
                    'Train security operations team'
                ],
                'deliverables': [
                    'Penetration test report',
                    'Optimized security controls',
                    'Training materials and documentation',
                    'Continuous monitoring procedures'
                ]
            }
        ]

        for phase_info in phases:
            self.doc.add_heading(phase_info['phase'], level=2)

            self.doc.add_heading('Objectives:', level=3)
            for objective in phase_info['objectives']:
                self.doc.add_paragraph(objective, style='List Bullet')

            self.doc.add_paragraph()
            self.doc.add_heading('Deliverables:', level=3)
            for deliverable in phase_info['deliverables']:
                self.doc.add_paragraph(deliverable, style='List Bullet')

            self.doc.add_paragraph()

        # Success criteria
        self.doc.add_heading('Success Criteria', level=2)

        success_criteria = [
            'All tier-to-tier communications encrypted with TLS 1.3',
            'Zero trust policies enforced with explicit allow rules only',
            'Mean time to detect (MTTD) security incidents < 15 minutes',
            'Mean time to respond (MTTR) to incidents < 4 hours',
            '100% of security events logged and retained for 90 days',
            'All critical and high vulnerabilities remediated within SLA',
            'Regular penetration testing with no critical findings',
            'Security team trained and incident response tested quarterly'
        ]

        for criteria in success_criteria:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run('✓ ').bold = True
            para.add_run(criteria)

    # Helper methods
    def _calculate_threat_score(self) -> int:
        """Calculate overall threat score (0-100)"""
        score = 0

        # Zone risk
        zone = self.app_data.get('security_zone', 'UNKNOWN')
        score += self.TIER_RISK_SCORES.get(zone, 50)

        # Dependency risk
        dependencies = self.app_data.get('predicted_dependencies', [])
        external_deps = [d for d in dependencies if d.get('type') == 'downstream_app']
        score += min(len(external_deps) * 5, 30)  # Cap at 30

        # Data sensitivity (if available)
        if self.app_data.get('has_sensitive_data', False):
            score += 20

        # Complexity
        if len(dependencies) > 10:
            score += 15

        return min(score, 100)

    def _get_risk_rating(self, score: int) -> str:
        """Get risk rating from threat score"""
        if score >= 80:
            return 'CRITICAL'
        elif score >= 60:
            return 'HIGH'
        elif score >= 40:
            return 'MEDIUM'
        else:
            return 'LOW'

    def _get_primary_concerns(self) -> str:
        """Get primary security focus areas (positive framing)"""
        focus_areas = []

        zone = self.app_data.get('security_zone', 'UNKNOWN')
        if zone == 'WEB_TIER':
            focus_areas.append('Public-facing tier hardening')

        dependencies = self.app_data.get('predicted_dependencies', [])
        if len(dependencies) > 10:
            focus_areas.append('Dependency chain optimization')

        external_deps = [d for d in dependencies if d.get('type') == 'downstream_app']
        if len(external_deps) > 5:
            focus_areas.append('External integration security')

        if not focus_areas:
            focus_areas.append('Standard security best practices')

        return ', '.join(focus_areas)

    def _count_required_rules(self) -> str:
        """Estimate number of required segmentation rules"""
        dependencies = self.app_data.get('predicted_dependencies', [])
        base_rules = 15  # Minimum rules
        dep_rules = len(dependencies) * 2  # 2 rules per dependency (ingress/egress)
        return f'{base_rules + dep_rules}+'

    def _count_critical_issues(self) -> str:
        """Count critical security issues"""
        issues = 0

        zone = self.app_data.get('security_zone', 'UNKNOWN')
        if zone == 'WEB_TIER':
            issues += 1  # Internet exposure

        if zone == 'UNKNOWN':
            issues += 1  # Unclassified zone

        dependencies = self.app_data.get('predicted_dependencies', [])
        if len(dependencies) > 15:
            issues += 1  # Complex dependencies

        return str(issues)

    def _guess_port(self, dep_type: str) -> str:
        """Guess port based on dependency type"""
        port_map = {
            'database': 'TCP/5432',
            'cache': 'TCP/6379',
            'queue': 'TCP/9092',
            'service': 'TCP/8080',
            'api': 'TCP/443'
        }
        return port_map.get(dep_type, 'TCP/443')


def generate_threat_surface_document(
    app_name: str,
    app_data: Dict,
    output_path: str
):
    """Generate threat surface and network segmentation document

    Args:
        app_name: Application name
        app_data: Application topology data
        output_path: Output path for Word document

    Returns:
        Path to generated document
    """
    doc_gen = ThreatSurfaceNetSegDocument(app_name, app_data)
    doc_gen.generate_document(output_path)

    return output_path

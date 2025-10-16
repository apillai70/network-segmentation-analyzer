"""
Topology and Network Analysis Document Generator
================================================
Generates comprehensive network topology analysis documents with
banking-specific segmentation guidance and compliance frameworks.

Focus Areas:
- Application dependency mapping
- Network topology visualization
- Banking regulatory compliance (PCI-DSS, GLBA, FFIEC)
- Banking-specific segmentation patterns
- Financial services security best practices

Author: Network Security Team
Version: 1.0
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

logger = logging.getLogger(__name__)


class TopologyNetworkAnalysisDocument:
    """Generates topology and network analysis documents with banking focus"""

    # Banking-specific zone classifications
    BANKING_ZONES = {
        'CARD_PROCESSING': 'Cardholder Data Environment (CDE) - PCI-DSS Scope',
        'ONLINE_BANKING': 'Customer-facing digital banking channels',
        'CORE_BANKING': 'Core banking system and general ledger',
        'ATM_BRANCH': 'ATM network and branch systems',
        'TREASURY_TRADING': 'Treasury management and trading systems',
        'CUSTOMER_DATA': 'Customer PII and account information',
        'REPORTING_ANALYTICS': 'Business intelligence and regulatory reporting',
        'THIRD_PARTY': 'External vendor and partner integrations'
    }

    def __init__(self, app_name: str, app_data: Optional[Dict] = None):
        """Initialize document generator

        Args:
            app_name: Application name
            app_data: Application topology and analysis data
        """
        self.app_name = app_name
        self.app_data = app_data or {}
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles"""
        styles = self.doc.styles

        try:
            # Heading styles - blue theme for topology docs
            h1_style = styles['Heading 1']
            h1_style.font.size = Pt(18)
            h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            h1_style.font.bold = True

            h2_style = styles['Heading 2']
            h2_style.font.size = Pt(14)
            h2_style.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue

            h3_style = styles['Heading 3']
            h3_style.font.size = Pt(12)
            h3_style.font.color.rgb = RGBColor(51, 153, 255)  # Light blue
        except:
            pass

    def generate_document(self, output_path: str):
        """Generate topology and network analysis document

        Args:
            output_path: Output path for Word document
        """
        logger.info(f"Generating topology and network analysis document for: {self.app_name}")

        # Cover page
        self._add_cover_page()
        self.doc.add_page_break()

        # Executive summary
        self._add_executive_summary()
        self.doc.add_page_break()

        # Network topology analysis
        self._add_topology_analysis()
        self.doc.add_page_break()

        # Application dependency mapping
        self._add_dependency_mapping()
        self.doc.add_page_break()

        # Banking regulatory framework
        self._add_banking_regulatory_framework()
        self.doc.add_page_break()

        # Banking-specific segmentation patterns
        self._add_banking_segmentation_patterns()
        self.doc.add_page_break()

        # Network security zones for banking
        self._add_banking_security_zones()
        self.doc.add_page_break()

        # Compliance and audit considerations
        self._add_compliance_audit()
        self.doc.add_page_break()

        # Implementation recommendations
        self._add_implementation_recommendations()

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output_file))
        logger.info(f"✓ Topology and network analysis document saved: {output_path}")

    def _add_cover_page(self):
        """Add professional cover page"""
        # Title
        title = self.doc.add_heading(
            f'Network Topology Analysis\n&\nBanking Segmentation Strategy\n\n{self.app_name}',
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Classification
        classification = self.doc.add_paragraph()
        classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classification.add_run('🏦 CONFIDENTIAL - BANKING NETWORK ANALYSIS 🏦')
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)

        for _ in range(2):
            self.doc.add_paragraph()

        # Document info
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run('Version: 1.0\n').bold = True
        info_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
        info_para.add_run('Classification: Confidential - Financial Services\n')
        info_para.add_run(f'Application: {self.app_name}\n')

        zone = self.app_data.get('security_zone', 'Unknown')
        info_para.add_run(f'Security Zone: {zone}\n')

        for _ in range(3):
            self.doc.add_paragraph()

        # Footer
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Prepared by: Network Architecture & Security Team\n')
        footer.add_run('Financial Services Network Segmentation Program\n')
        footer.add_run('Auto-generated by Network Segmentation Analyzer v3.0\n')

    def _add_executive_summary(self):
        """Add executive summary"""
        self.doc.add_heading('Executive Summary', level=1)

        intro = ("This document provides a comprehensive analysis of network topology, application dependencies, "
                "and banking-specific segmentation requirements for " + self.app_name + ". The analysis is designed "
                "to support regulatory compliance, operational resilience, and secure banking operations.")
        self.doc.add_paragraph(intro)

        self.doc.add_paragraph()

        # Key metrics
        self.doc.add_heading('Network Topology Metrics', level=2)

        dependencies = self.app_data.get('predicted_dependencies', [])
        zone = self.app_data.get('security_zone', 'APP_TIER')

        metrics_table = self.doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Medium Shading 1 Accent 1'

        metrics = [
            ('Application Name', self.app_name),
            ('Security Zone', zone),
            ('Total Dependencies', str(len(dependencies))),
            ('Banking Classification', self._get_banking_classification(zone)),
            ('Regulatory Scope', self._get_regulatory_scope(zone))
        ]

        for idx, (metric, value) in enumerate(metrics):
            metrics_table.rows[idx].cells[0].text = metric
            metrics_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            metrics_table.rows[idx].cells[1].text = value

        self.doc.add_paragraph()

        # Key findings
        findings_para = self.doc.add_paragraph()
        findings_para.add_run('Key Findings:').bold = True

        findings = [
            f'Application operates in {zone} with {len(dependencies)} identified dependencies',
            f'Banking classification: {self._get_banking_classification(zone)}',
            f'Primary regulatory framework: {self._get_regulatory_scope(zone)}',
            'Segmentation recommendations align with banking industry best practices'
        ]

        for finding in findings:
            self.doc.add_paragraph(f'• {finding}')

    def _add_topology_analysis(self):
        """Add network topology analysis"""
        self.doc.add_heading('Network Topology Analysis', level=1)

        topo_intro = ("Network topology analysis examines the structure of application dependencies, "
                     "communication patterns, and security zone boundaries. Understanding topology is "
                     "essential for effective segmentation in banking environments.")
        self.doc.add_paragraph(topo_intro)

        # Topology overview
        self.doc.add_heading('Topology Overview', level=2)

        dependencies = self.app_data.get('predicted_dependencies', [])
        zone = self.app_data.get('security_zone', 'APP_TIER')

        overview_text = (f"The {self.app_name} application is deployed in the {zone} security zone. "
                        f"Analysis has identified {len(dependencies)} network dependencies representing "
                        "connections to databases, caches, messaging systems, and external services.")
        self.doc.add_paragraph(overview_text)

        # Dependency breakdown by type
        self.doc.add_paragraph()
        self.doc.add_heading('Dependency Breakdown', level=3)

        dep_by_type = defaultdict(int)
        for dep in dependencies:
            dep_type = dep.get('type', 'unknown')
            dep_by_type[dep_type] += 1

        if dep_by_type:
            dep_table = self.doc.add_table(rows=len(dep_by_type)+1, cols=2)
            dep_table.style = 'Light Grid Accent 1'

            dep_table.rows[0].cells[0].text = 'Dependency Type'
            dep_table.rows[0].cells[1].text = 'Count'
            for cell in dep_table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True

            for idx, (dep_type, count) in enumerate(sorted(dep_by_type.items()), 1):
                dep_table.rows[idx].cells[0].text = dep_type.title()
                dep_table.rows[idx].cells[1].text = str(count)
        else:
            self.doc.add_paragraph('No dependencies identified in current analysis.')

        # Network flow patterns
        self.doc.add_paragraph()
        self.doc.add_heading('Network Flow Patterns', level=2)

        flow_text = ("Network flow analysis reveals communication patterns between application components. "
                    "In banking environments, strict flow controls prevent unauthorized lateral movement "
                    "and contain security incidents.")
        self.doc.add_paragraph(flow_text)

        flow_patterns = [
            ('North-South Traffic', f'{zone} ↔ External services and APIs', 'Controlled egress through proxies'),
            ('East-West Traffic', f'{zone} ↔ Other internal tiers', 'Micro-segmentation with explicit allow rules'),
            ('Management Traffic', f'Management tier → {zone}', 'Privileged access via bastion hosts'),
            ('Monitoring Traffic', f'{zone} → SIEM/Monitoring', 'Continuous security event collection')
        ]

        for pattern, description, control in flow_patterns:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(pattern + ': ').bold = True
            para.add_run(f'{description} - ')
            para.add_run(control).italic = True

    def _add_dependency_mapping(self):
        """Add application dependency mapping"""
        self.doc.add_heading('Application Dependency Mapping', level=1)

        dep_intro = ("Dependency mapping identifies all network connections required for application operation. "
                    "Complete dependency maps are critical for banking segmentation to ensure business "
                    "continuity while maintaining security controls.")
        self.doc.add_paragraph(dep_intro)

        dependencies = self.app_data.get('predicted_dependencies', [])

        if len(dependencies) > 0:
            self.doc.add_heading('Identified Dependencies', level=2)

            # Create dependency table (show up to 20)
            dep_table = self.doc.add_table(rows=min(len(dependencies), 20)+1, cols=3)
            dep_table.style = 'Light List Accent 1'

            dep_table.rows[0].cells[0].text = 'Dependency Name'
            dep_table.rows[0].cells[1].text = 'Type'
            dep_table.rows[0].cells[2].text = 'Purpose'
            for cell in dep_table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True

            for idx, dep in enumerate(dependencies[:20], 1):
                dep_name = dep.get('name', 'Unknown')
                dep_type = dep.get('type', 'unknown')
                purpose = self._get_dependency_purpose(dep_type)

                dep_table.rows[idx].cells[0].text = dep_name
                dep_table.rows[idx].cells[1].text = dep_type.title()
                dep_table.rows[idx].cells[2].text = purpose

            if len(dependencies) > 20:
                self.doc.add_paragraph()
                note = self.doc.add_paragraph()
                note.add_run(f'Note: ').bold = True
                note.add_run(f'Showing first 20 of {len(dependencies)} total dependencies. '
                           'See topology JSON file for complete listing.')
        else:
            self.doc.add_paragraph('No dependencies identified. This may indicate:')
            self.doc.add_paragraph('• Standalone application with no external dependencies', style='List Bullet')
            self.doc.add_paragraph('• Incomplete network flow analysis', style='List Bullet')
            self.doc.add_paragraph('• Application not yet deployed or monitored', style='List Bullet')

        # Dependency security considerations
        self.doc.add_paragraph()
        self.doc.add_heading('Dependency Security Considerations', level=2)

        security_considerations = [
            ('Database Dependencies', 'Require encryption in transit (TLS), connection pooling limits, database activity monitoring (DAM)'),
            ('Cache Dependencies', 'Isolate cache tier, implement cache poisoning protections, encrypt sensitive cached data'),
            ('External APIs', 'API gateway with authentication, rate limiting, egress filtering, certificate pinning'),
            ('Internal Services', 'Service mesh with mTLS, mutual authentication, least-privilege access controls')
        ]

        for consideration, description in security_considerations:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(consideration + ': ').bold = True
            para.add_run(description)

    def _add_banking_regulatory_framework(self):
        """Add banking regulatory framework section"""
        self.doc.add_heading('Banking Regulatory Framework', level=1)

        reg_intro = ("Financial institutions operate under comprehensive regulatory frameworks that mandate "
                    "specific network segmentation and security controls. This section maps regulatory "
                    "requirements to segmentation strategies.")
        self.doc.add_paragraph(reg_intro)

        # PCI-DSS for banking
        self.doc.add_heading('PCI-DSS: Payment Card Industry Data Security Standard', level=2)

        pci_text = ("PCI-DSS is MANDATORY for any organization that stores, processes, or transmits cardholder data. "
                   "Network segmentation is explicitly required under Requirement 1.2.1 to reduce CDE scope.")
        self.doc.add_paragraph(pci_text)

        pci_requirements = [
            ('Requirement 1: Firewalls', 'Install and maintain firewall configuration to protect cardholder data', 'CRITICAL'),
            ('Requirement 1.2.1: CDE Isolation', 'Restrict inbound/outbound traffic to that necessary for CDE', 'CRITICAL'),
            ('Requirement 2: No Defaults', 'Change vendor-supplied defaults, disable unnecessary services', 'HIGH'),
            ('Requirement 3: Data Protection', 'Protect stored cardholder data with encryption', 'CRITICAL'),
            ('Requirement 11.3.4: Pen Testing', 'Verify segmentation methods are operational and effective', 'HIGH')
        ]

        pci_table = self.doc.add_table(rows=len(pci_requirements)+1, cols=3)
        pci_table.style = 'Medium Shading 1 Accent 1'

        pci_table.rows[0].cells[0].text = 'Requirement'
        pci_table.rows[0].cells[1].text = 'Description'
        pci_table.rows[0].cells[2].text = 'Priority'
        for cell in pci_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (req, desc, priority) in enumerate(pci_requirements, 1):
            pci_table.rows[idx].cells[0].text = req
            pci_table.rows[idx].cells[1].text = desc
            pci_table.rows[idx].cells[2].text = priority

        self.doc.add_paragraph()

        # GLBA
        self.doc.add_heading('GLBA: Gramm-Leach-Bliley Act', level=2)

        glba_text = ("GLBA requires financial institutions to protect customer financial information. "
                    "The Safeguards Rule mandates administrative, technical, and physical safeguards "
                    "including network segmentation.")
        self.doc.add_paragraph(glba_text)

        glba_requirements = [
            'Designate employee(s) to coordinate information security program',
            'Identify and assess risks to customer information in each business area',
            'Design and implement safeguards to control identified risks',
            'Regularly monitor and test the effectiveness of safeguards',
            'Select service providers that maintain appropriate safeguards'
        ]

        for req in glba_requirements:
            self.doc.add_paragraph(req, style='List Bullet')

        self.doc.add_paragraph()
        glba_seg = self.doc.add_paragraph()
        glba_seg.add_run('Segmentation Benefit: ').bold = True
        glba_seg.add_run('Network segmentation addresses GLBA by limiting access to customer information '
                        'systems, enabling better risk assessment, and facilitating monitoring.')

        # FFIEC
        self.doc.add_paragraph()
        self.doc.add_heading('FFIEC: Federal Financial Institutions Examination Council', level=2)

        ffiec_text = ("FFIEC provides examination standards for financial institutions. The Cybersecurity "
                     "Assessment Tool (CAT) evaluates network segmentation as a key control.")
        self.doc.add_paragraph(ffiec_text)

        ffiec_domains = [
            ('Cyber Risk Management', 'Risk assessment and mitigation strategies including segmentation'),
            ('Threat Intelligence', 'Integration of threat intelligence to inform network controls'),
            ('Cybersecurity Controls', 'Implementation of defense-in-depth including network segmentation'),
            ('External Dependency Management', 'Controls for third-party connections and vendor access'),
            ('Incident Response', 'Capability to contain incidents through segmentation boundaries')
        ]

        for domain, description in ffiec_domains:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(domain + ': ').bold = True
            para.add_run(description)

        # SOX for banking
        self.doc.add_paragraph()
        self.doc.add_heading('SOX: Sarbanes-Oxley Act (Financial Reporting)', level=2)

        sox_text = ("SOX Section 404 requires assessment of internal controls over financial reporting. "
                   "For banks, this includes controls over general ledger systems, financial data warehouses, "
                   "and reporting applications.")
        self.doc.add_paragraph(sox_text)

        sox_controls = [
            'Segregation of Duties: Network segmentation supports separation between development, testing, and production financial systems',
            'Access Controls: Limit network access to financial reporting systems to authorized personnel only',
            'Change Management: Segmentation enables strict change control and approval workflows for production systems',
            'Audit Trails: Network segmentation facilitates comprehensive logging and monitoring of financial system access'
        ]

        for control in sox_controls:
            self.doc.add_paragraph(control, style='List Bullet')

    def _add_banking_segmentation_patterns(self):
        """Add banking-specific segmentation patterns"""
        self.doc.add_heading('Banking-Specific Segmentation Patterns', level=1)

        pattern_intro = ("Banking and financial services have unique segmentation requirements driven by "
                        "regulatory mandates, operational risk, and the criticality of financial data. "
                        "This section provides proven segmentation patterns for banking environments.")
        self.doc.add_paragraph(pattern_intro)

        # Pattern 1: Card Processing Environment
        self.doc.add_heading('Pattern 1: Card Processing Environment (CDE) Isolation', level=2)

        cde_text = ("The Cardholder Data Environment must be strictly isolated per PCI-DSS requirements. "
                   "This is the HIGHEST PRIORITY segmentation pattern for banks processing payment cards.")
        self.doc.add_paragraph(cde_text)

        self.doc.add_heading('CDE Segmentation Architecture:', level=3)

        cde_zones = [
            ('CDE-WEB', 'Payment gateway and customer-facing payment pages', 'Internet → HTTPS/443 only'),
            ('CDE-APP', 'Payment processing application servers', 'CDE-WEB only'),
            ('CDE-DATA', 'Cardholder data storage (encrypted)', 'CDE-APP only'),
            ('OUT-OF-SCOPE', 'All other systems (explicitly NOT in CDE)', 'No direct connectivity to CDE')
        ]

        cde_table = self.doc.add_table(rows=len(cde_zones)+1, cols=3)
        cde_table.style = 'Light Grid Accent 1'

        cde_table.rows[0].cells[0].text = 'Zone'
        cde_table.rows[0].cells[1].text = 'Function'
        cde_table.rows[0].cells[2].text = 'Allowed Ingress'
        for cell in cde_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (zone, function, ingress) in enumerate(cde_zones, 1):
            cde_table.rows[idx].cells[0].text = zone
            cde_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            cde_table.rows[idx].cells[1].text = function
            cde_table.rows[idx].cells[2].text = ingress

        self.doc.add_paragraph()
        cde_critical = self.doc.add_paragraph()
        cde_critical.add_run('CRITICAL: ').bold = True
        cde_critical.add_run('Annual penetration testing (PCI Requirement 11.3.4) must verify segmentation '
                            'effectively isolates CDE from out-of-scope systems.')

        # Pattern 2: Online Banking DMZ
        self.doc.add_paragraph()
        self.doc.add_heading('Pattern 2: Online Banking DMZ', level=2)

        ob_text = ("Online banking channels require defense-in-depth with multiple security zones between "
                  "the internet and core banking systems.")
        self.doc.add_paragraph(ob_text)

        ob_layers = [
            ('Internet Edge', 'DDoS mitigation, WAF, bot detection', 'First line of defense'),
            ('DMZ - Web Tier', 'Load balancers, web servers, reverse proxies', 'Public-facing components'),
            ('DMZ - App Tier', 'Online banking application servers', 'Business logic (no direct internet access)'),
            ('Internal - API Gateway', 'API gateway and authentication services', 'Mediates access to core banking'),
            ('Internal - Core Banking', 'Core banking system and account databases', 'No direct external connectivity')
        ]

        for layer, components, purpose in ob_layers:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(f'{layer}: ').bold = True
            para.add_run(f'{components} - ')
            para.add_run(purpose).italic = True

        # Pattern 3: Core Banking Isolation
        self.doc.add_paragraph()
        self.doc.add_heading('Pattern 3: Core Banking System Isolation', level=2)

        core_text = ("Core banking systems (general ledger, account management, transaction processing) "
                    "represent the heart of banking operations and require maximum security.")
        self.doc.add_paragraph(core_text)

        core_controls = [
            'Physical Isolation: Deploy core banking in dedicated network segments with no direct internet connectivity',
            'API Gateway Pattern: All external access routed through authenticated API gateway with rate limiting',
            'Database Firewall: Layer 7 firewall inspecting all database queries for anomalies and injection attacks',
            'Privileged Access: Administrative access only via PAM solution with session recording and MFA',
            'Change Control: All changes to core banking network controls require change board approval'
        ]

        for control in core_controls:
            self.doc.add_paragraph(control, style='List Bullet')

        # Pattern 4: ATM/Branch Network
        self.doc.add_paragraph()
        self.doc.add_heading('Pattern 4: ATM and Branch Network Segmentation', level=2)

        atm_text = ("ATM and branch networks connect remote locations to core systems. Compromise of these "
                   "networks poses significant operational risk.")
        self.doc.add_paragraph(atm_text)

        atm_controls = [
            ('Network Isolation', 'Dedicated VLAN/MPLS for ATM traffic, no connectivity to corporate networks'),
            ('Encryption', 'End-to-end encryption for all ATM transactions (PIN, card data, transaction data)'),
            ('Device Authentication', 'Mutual certificate-based authentication between ATMs and host systems'),
            ('Monitoring', 'Real-time monitoring for anomalous transaction patterns and network behavior'),
            ('Incident Response', 'Automated blocking of compromised ATMs with minimal impact to other devices')
        ]

        for control, description in atm_controls:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(control + ': ').bold = True
            para.add_run(description)

    def _add_banking_security_zones(self):
        """Add banking security zones"""
        self.doc.add_heading('Network Security Zones for Banking', level=1)

        zones_intro = ("Banking environments require specialized security zones beyond standard enterprise "
                      "segmentation. This section defines banking-specific zones and their requirements.")
        self.doc.add_paragraph(zones_intro)

        # Banking zone matrix
        self.doc.add_heading('Banking Security Zone Matrix', level=2)

        banking_zones = [
            ('CARD_PROCESSING', 'Cardholder Data Environment (CDE)', 'PCI-DSS In-Scope', 'CRITICAL', 'Quarterly ASV scan, Annual pen test'),
            ('ONLINE_BANKING', 'Customer digital channels', 'Customer PII', 'HIGH', 'FFIEC CAT, GLBA Safeguards'),
            ('CORE_BANKING', 'General ledger, accounts', 'Financial transactions', 'CRITICAL', 'SOX controls, Operational risk'),
            ('ATM_BRANCH', 'ATM and branch systems', 'Remote transaction processing', 'HIGH', 'Physical security, Device authentication'),
            ('TREASURY_TRADING', 'Treasury and trading platforms', 'Market data, Trading positions', 'HIGH', 'Market abuse regulations'),
            ('CUSTOMER_DATA', 'CRM, Customer databases', 'Customer PII, Account info', 'HIGH', 'GLBA, State privacy laws'),
            ('REPORTING', 'BI, Regulatory reporting', 'Financial reports, Metrics', 'MEDIUM', 'SOX, Regulatory reporting'),
            ('THIRD_PARTY', 'Vendor integrations', 'External partner connections', 'MEDIUM', 'Third-party risk management')
        ]

        zone_table = self.doc.add_table(rows=len(banking_zones)+1, cols=5)
        zone_table.style = 'Medium Shading 1 Accent 1'

        headers = ['Zone', 'Description', 'Data Classification', 'Risk Level', 'Compliance']
        for idx, header in enumerate(headers):
            zone_table.rows[0].cells[idx].text = header
            zone_table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

        for idx, (zone, desc, data_class, risk, compliance) in enumerate(banking_zones, 1):
            zone_table.rows[idx].cells[0].text = zone
            zone_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            zone_table.rows[idx].cells[1].text = desc
            zone_table.rows[idx].cells[2].text = data_class
            zone_table.rows[idx].cells[3].text = risk
            zone_table.rows[idx].cells[4].text = compliance

        # Inter-zone firewall rules
        self.doc.add_paragraph()
        self.doc.add_heading('Inter-Zone Firewall Rules', level=2)

        firewall_intro = ("All traffic between banking security zones must traverse firewalls with explicit "
                         "allow rules. Default deny is mandatory.")
        self.doc.add_paragraph(firewall_intro)

        zone_rules = [
            ('ONLINE_BANKING → CORE_BANKING', 'HTTPS/443 via API Gateway only', 'Authenticated API calls with rate limiting'),
            ('CARD_PROCESSING → THIRD_PARTY', 'HTTPS/443 to payment processors', 'Certificate pinning, egress whitelist'),
            ('ATM_BRANCH → CORE_BANKING', 'Proprietary protocol (encrypted)', 'Mutual certificate authentication'),
            ('REPORTING → CUSTOMER_DATA', 'Database queries (read-only)', 'Service account with minimal SELECT privileges'),
            ('ALL ZONES → MANAGEMENT', 'Syslog (TCP/514), Monitoring', 'One-way traffic to SIEM')
        ]

        for rule, protocol, control in zone_rules:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(rule).bold = True
            para.add_run(f'\n  Protocol: {protocol}')
            para.add_run(f'\n  Control: ').italic = True
            para.add_run(control).italic = True

    def _add_compliance_audit(self):
        """Add compliance and audit considerations"""
        self.doc.add_heading('Compliance and Audit Considerations', level=1)

        audit_intro = ("Banking network segmentation must withstand regulatory examination and internal audit. "
                      "This section provides guidance for demonstrating compliance.")
        self.doc.add_paragraph(audit_intro)

        # Evidence collection
        self.doc.add_heading('Evidence Collection for Auditors', level=2)

        evidence_items = [
            ('Network Diagrams', 'Current-state architecture diagrams showing all security zones and boundaries', 'Updated quarterly or after major changes'),
            ('Firewall Rule Documentation', 'Complete firewall ruleset with business justification for each allow rule', 'Reviewed annually minimum'),
            ('Penetration Test Reports', 'Annual penetration testing validating segmentation effectiveness (PCI 11.3.4)', 'Required for PCI compliance'),
            ('Change Management Records', 'Approval records for all changes to network segmentation controls', 'Retained per retention policy'),
            ('Access Control Lists', 'Documentation of who has access to configure network security controls', 'Reviewed quarterly'),
            ('Monitoring and Alerting', 'Evidence of continuous monitoring and alerting on segmentation violations', 'Logs retained 90+ days')
        ]

        evidence_table = self.doc.add_table(rows=len(evidence_items)+1, cols=3)
        evidence_table.style = 'Light List Accent 1'

        evidence_table.rows[0].cells[0].text = 'Evidence Type'
        evidence_table.rows[0].cells[1].text = 'Description'
        evidence_table.rows[0].cells[2].text = 'Frequency'
        for cell in evidence_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for idx, (evidence, desc, freq) in enumerate(evidence_items, 1):
            evidence_table.rows[idx].cells[0].text = evidence
            evidence_table.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
            evidence_table.rows[idx].cells[1].text = desc
            evidence_table.rows[idx].cells[2].text = freq

        # Common audit findings
        self.doc.add_paragraph()
        self.doc.add_heading('Common Audit Findings (and How to Avoid Them)', level=2)

        findings = [
            ('Finding: Overly Permissive Firewall Rules', 'Remediation: Implement deny-all default with explicit allows, document business justification for all rules'),
            ('Finding: Lack of Segmentation Testing', 'Remediation: Conduct annual penetration testing specifically targeting segmentation (PCI 11.3.4)'),
            ('Finding: Inadequate Monitoring', 'Remediation: Implement SIEM with alerts for cross-zone traffic violations and anomalies'),
            ('Finding: Undocumented Network Architecture', 'Remediation: Maintain current network diagrams, update within 30 days of changes'),
            ('Finding: Ineffective Scope Reduction', 'Remediation: Implement flat file replacement, tokenization, or point-to-point encryption to remove systems from PCI scope'),
            ('Finding: Missing Change Approvals', 'Remediation: Enforce change management process with approval workflow for all network security changes')
        ]

        for finding, remediation in findings:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(finding).bold = True
            para.add_run(f'\n  {remediation}')

    def _add_implementation_recommendations(self):
        """Add implementation recommendations"""
        self.doc.add_heading('Implementation Recommendations', level=1)

        impl_intro = ("Implementing banking network segmentation requires careful planning and phased execution "
                     "to avoid operational disruption while achieving compliance and security objectives.")
        self.doc.add_paragraph(impl_intro)

        # Phased approach
        self.doc.add_heading('Phased Implementation Approach', level=2)

        phases = [
            {
                'phase': 'Phase 1: Assessment and Scoping (Weeks 1-4)',
                'activities': [
                    'Conduct comprehensive asset inventory and data flow mapping',
                    'Identify regulatory requirements applicable to each system',
                    'Define security zone boundaries and classify all systems',
                    'Document current state and design target state architecture',
                    'Obtain stakeholder buy-in and executive sponsorship'
                ]
            },
            {
                'phase': 'Phase 2: CDE Isolation (Weeks 5-12) - PRIORITY 1',
                'activities': [
                    'Isolate Cardholder Data Environment per PCI-DSS Requirement 1.2.1',
                    'Deploy firewalls between CDE and out-of-scope systems',
                    'Implement deny-all default with explicit allow rules',
                    'Configure logging and monitoring for all CDE boundary traffic',
                    'Conduct segmentation penetration test to validate isolation'
                ]
            },
            {
                'phase': 'Phase 3: Online Banking DMZ (Weeks 13-20)',
                'activities': [
                    'Establish DMZ architecture for customer-facing systems',
                    'Deploy WAF and DDoS protection at internet edge',
                    'Implement API gateway for mediated access to core banking',
                    'Enable TLS 1.3 for all customer connections',
                    'Configure rate limiting and fraud detection'
                ]
            },
            {
                'phase': 'Phase 4: Core Banking Isolation (Weeks 21-28)',
                'activities': [
                    'Isolate core banking systems from direct external access',
                    'Deploy database firewalls with query inspection',
                    'Implement PAM for privileged access to core systems',
                    'Establish strict change management for core banking network',
                    'Enable comprehensive audit logging and SIEM integration'
                ]
            },
            {
                'phase': 'Phase 5: Remaining Zones and Optimization (Weeks 29+)',
                'activities': [
                    'Complete segmentation for ATM/Branch, Treasury, Reporting zones',
                    'Fine-tune firewall rules to minimize false positives',
                    'Conduct enterprise-wide segmentation validation testing',
                    'Develop runbooks and train operations staff',
                    'Establish continuous monitoring and improvement process'
                ]
            }
        ]

        for phase_info in phases:
            self.doc.add_heading(phase_info['phase'], level=3)
            for activity in phase_info['activities']:
                self.doc.add_paragraph(activity, style='List Bullet')
            self.doc.add_paragraph()

        # Success criteria
        self.doc.add_heading('Success Criteria', level=2)

        success_criteria = [
            'PCI-DSS Compliance: CDE successfully scoped and isolated, penetration testing validates segmentation',
            'GLBA Safeguards: Customer financial information systems protected with appropriate network controls',
            'FFIEC CAT Maturity: Network segmentation controls rated "Innovative" or "Evolving" in FFIEC assessment',
            'SOX Controls: Effective segregation of duties for financial reporting systems with documented evidence',
            'Operational Resilience: Zero unplanned outages caused by segmentation implementation',
            'Audit Readiness: All required evidence available and documentation current for regulatory examination'
        ]

        for criteria in success_criteria:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run('✓ ').bold = True
            para.add_run(criteria)

        # Final recommendations
        self.doc.add_paragraph()
        self.doc.add_heading('Final Recommendations', level=2)

        final_recs = [
            ('Prioritize CDE Isolation', 'If you process payment cards, CDE isolation is the highest priority and is explicitly required by PCI-DSS.'),
            ('Engage Stakeholders Early', 'Network segmentation affects multiple business units. Obtain buy-in from application owners, operations, and compliance.'),
            ('Plan for Testing', 'Budget time and resources for comprehensive testing including penetration testing to validate segmentation effectiveness.'),
            ('Automate Compliance', 'Implement automated compliance monitoring and alerting rather than manual periodic reviews.'),
            ('Document Everything', 'Maintain comprehensive documentation of architecture, firewall rules, and change approvals for audit purposes.'),
            ('Continuous Improvement', 'Segmentation is not a one-time project. Establish ongoing review and optimization processes.')
        ]

        for rec, description in final_recs:
            para = self.doc.add_paragraph(style='List Bullet')
            para.add_run(rec + ': ').bold = True
            para.add_run(description)

    # Helper methods
    def _get_banking_classification(self, zone: str) -> str:
        """Get banking classification for a zone"""
        banking_map = {
            'WEB_TIER': 'Online Banking',
            'APP_TIER': 'Application Services',
            'DATA_TIER': 'Core Banking / Customer Data',
            'CACHE_TIER': 'Performance Tier',
            'MESSAGING_TIER': 'Integration Tier',
            'MANAGEMENT_TIER': 'Operations & Monitoring',
            'EXTERNAL': 'Third-Party Integration'
        }
        return banking_map.get(zone, 'General Purpose')

    def _get_regulatory_scope(self, zone: str) -> str:
        """Get regulatory scope for a zone"""
        regulatory_map = {
            'WEB_TIER': 'GLBA, FFIEC',
            'APP_TIER': 'GLBA, FFIEC',
            'DATA_TIER': 'PCI-DSS, GLBA, SOX',
            'CACHE_TIER': 'GLBA',
            'MESSAGING_TIER': 'GLBA',
            'MANAGEMENT_TIER': 'SOX, FFIEC',
            'EXTERNAL': 'Third-Party Risk Management'
        }
        return regulatory_map.get(zone, 'GLBA, FFIEC (General)')

    def _get_dependency_purpose(self, dep_type: str) -> str:
        """Get purpose description for dependency type"""
        purpose_map = {
            'database': 'Data persistence and retrieval',
            'cache': 'Performance optimization and session management',
            'queue': 'Asynchronous messaging and event processing',
            'downstream_app': 'Integration with external services',
            'api': 'External API integration',
            'unknown': 'Requires classification'
        }
        return purpose_map.get(dep_type, 'Application dependency')


def generate_topology_network_analysis_document(
    app_name: str,
    app_data: Dict,
    output_path: str
):
    """Generate topology and network analysis document

    Args:
        app_name: Application name
        app_data: Application topology data
        output_path: Output path for Word document

    Returns:
        Path to generated document
    """
    doc_gen = TopologyNetworkAnalysisDocument(app_name, app_data)
    doc_gen.generate_document(output_path)

    return output_path

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Test SVG Embedding in Word Documents
=====================================
Tests whether python-docx properly embeds SVG files
"""

import sys
import os
from pathlib import Path

# Force UTF-8 encoding (Windows fix)
if sys.platform == 'win32':
    import codecs
    if hasattr(sys.stdout, 'buffer'):
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    if hasattr(sys.stderr, 'buffer'):
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

from docx import Document
from docx.shared import Inches

def test_svg_embedding():
    """Test SVG vs PNG embedding"""
    print("="*80)
    print("SVG EMBEDDING TEST")
    print("="*80)
    print()

    # Find an existing SVG file
    svg_files = list(Path('outputs_final/diagrams').glob('*.svg'))

    if not svg_files:
        print("[ERROR] No SVG files found. Generate one first:")
        print("  python generate_pngs_python.py --apps ACDA --format svg")
        return 1

    svg_path = svg_files[0]
    png_path = svg_path.with_suffix('.png')

    print(f"Test files:")
    print(f"  SVG: {svg_path}")
    print(f"  PNG: {png_path} (exists: {png_path.exists()})")
    print()

    # Create test Word document with SVG
    print("Creating test Word document with SVG...")
    doc_svg = Document()
    doc_svg.add_heading('SVG Embedding Test', 0)

    try:
        doc_svg.add_paragraph("This diagram is embedded as SVG (vector format):")
        doc_svg.add_picture(str(svg_path), width=Inches(6))

        svg_output = Path('test_output_svg.docx')
        doc_svg.save(str(svg_output))

        svg_size = svg_output.stat().st_size
        print(f"[OK] SVG document created: {svg_output}")
        print(f"     File size: {svg_size / 1024:.1f} KB")
        print()

    except Exception as e:
        print(f"[ERROR] Failed to embed SVG: {e}")
        print()
        return 1

    # Create test Word document with PNG for comparison
    if png_path.exists():
        print("Creating test Word document with PNG for comparison...")
        doc_png = Document()
        doc_png.add_heading('PNG Embedding Test', 0)

        try:
            doc_png.add_paragraph("This diagram is embedded as PNG (raster format):")
            doc_png.add_picture(str(png_path), width=Inches(6))

            png_output = Path('test_output_png.docx')
            doc_png.save(str(png_output))

            png_size = png_output.stat().st_size
            print(f"[OK] PNG document created: {png_output}")
            print(f"     File size: {png_size / 1024:.1f} KB")
            print()

            # Compare file sizes
            print("Comparison:")
            print(f"  SVG document: {svg_size / 1024:.1f} KB")
            print(f"  PNG document: {png_size / 1024:.1f} KB")

            if svg_size < png_size:
                print(f"  [OK] SVG is {(png_size - svg_size) / 1024:.1f} KB smaller!")
            elif svg_size > png_size:
                print(f"  [WARNING] SVG is {(svg_size - png_size) / 1024:.1f} KB larger")
                print(f"  This might indicate SVG was converted to PNG internally")
            else:
                print(f"  [WARNING] Same size - SVG might be converted to PNG")
            print()

        except Exception as e:
            print(f"[ERROR] Failed to embed PNG: {e}")
            print()

    # Inspect the Word document to see what's actually embedded
    print("Inspecting SVG document internals...")
    try:
        from zipfile import ZipFile
        import xml.etree.ElementTree as ET

        with ZipFile(svg_output, 'r') as docx_zip:
            # Check what image files are in the document
            image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

            print(f"Embedded media files: {len(image_files)}")
            for img in image_files:
                info = docx_zip.getinfo(img)
                ext = Path(img).suffix
                print(f"  - {img} ({info.file_size / 1024:.1f} KB) - Type: {ext}")

            # Check content types
            content_types = docx_zip.read('[Content_Types].xml').decode('utf-8')

            if 'svg' in content_types.lower():
                print("\n[SUCCESS] Document contains SVG references!")
                print("SVG is embedded as vector format - infinite zoom!")
            elif 'png' in content_types.lower() or 'jpeg' in content_types.lower():
                print("\n[WARNING] Document contains PNG/JPEG, not SVG")
                print("python-docx may have converted SVG to raster format")
            else:
                print("\n[INFO] Could not determine image format from content types")

    except Exception as e:
        print(f"[ERROR] Failed to inspect document: {e}")

    print()
    print("="*80)
    print("TEST COMPLETE")
    print("="*80)
    print()
    print("Next steps:")
    print("  1. Open test_output_svg.docx in Word")
    print("  2. Zoom in on the diagram (200%, 400%, etc.)")
    print("  3. Check if text remains crisp or becomes pixelated")
    print()
    print("Expected results:")
    print("  - SVG (vector): Crisp at any zoom level")
    print("  - PNG (raster): Pixelated when zoomed")
    print()

    return 0

if __name__ == '__main__':
    sys.exit(test_svg_embedding())
